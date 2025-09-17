# pages/dashboard.py
import streamlit as st
from logo_component import display_logo, display_sidebar_logo
from dashboard_exporter import export_dashboard_to_html
from auth_manager import is_authenticated, check_persistent_auth # Import authentication functions
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.io as pio
from io import StringIO
import base64
import scipy
import json
from datetime import datetime
from scipy import stats
import statsmodels.api as sm
import joypy

def sync_global_api_settings():
    """Sync global API settings for compatibility"""
    # Ensure global settings are available
    if 'global_api_key' not in st.session_state:
        st.session_state.global_api_key = ""
    if 'global_ai_provider' not in st.session_state:
        st.session_state.global_ai_provider = "DeepSeek"
    if 'global_model_name' not in st.session_state:
        st.session_state.global_model_name = "deepseek-chat"
    
    # Sync with local keys for backward compatibility
    st.session_state.api_key = st.session_state.global_api_key
    st.session_state.ai_provider = st.session_state.global_ai_provider
    st.session_state.model_name = st.session_state.global_model_name

# --- 1. Persistent Authentication Check ---
# Check for persistent authentication first (handles page refresh)
if not st.session_state.get('user'):
    check_persistent_auth()
    # Function to encode your local image
    def get_base64_of_bin_file(bin_file):
        with open(bin_file, 'rb') as f:
            data = f.read()
        return base64.b64encode(data).decode()

    def set_background_image(image_path):
        bin_str = get_base64_of_bin_file(image_path)
        st.markdown(
            f"""
            <style>
            .stApp {{
                background-image: url("data:image/jpg;base64,{bin_str}");
                background-size: cover;
                background-position: center;
                background-attachment: fixed;
            }}
            </style>
            """,
            unsafe_allow_html=True
        )

    # Call this function with your image path
    set_background_image('light_splash.avif') # Make sure 'mesh.JPG' exists

    col1, col2, col3 = st.columns([1, 6, 1])
    with col1:
        if st.button("🏠 Back to Login", type="primary", use_container_width=True):
            st.switch_page("analysis_agent.py")


    # Hero Section
    st.markdown("""
        <div style="text-align: center; margin: 40px 0 50px 0;">
            <h1 style="background: linear-gradient(90deg, #667eea, #764ba2);
                       -webkit-background-clip: text;
                       -webkit-text-fill-color: transparent;
                       font-weight: 800;
                       font-size: 3.2rem;
                       margin-bottom: 15px;
                       text-shadow: 0 4px 8px rgba(0,0,0,0.1);">
                📊 Interactive Dashboard Preview
            </h1>
            <p style="color: #666; font-size: 1.3rem; max-width: 900px; margin: 0 auto; line-height: 1.6;">
                Transform raw data into stunning visualizations and actionable insights with our AI-powered dashboard platform
            </p>
        </div>
    """, unsafe_allow_html=True)
    
    # Key Features Section
    st.markdown("""
        <div style="text-align: center; margin: 40px 0 30px 0;">
            <h2 style="color: #2d3748; font-size: 2.2rem; font-weight: 700; margin-bottom: 20px;">
                🎯 Dashboard Capabilities
            </h2>
        </div>
    """, unsafe_allow_html=True)
    
    # Enhanced Feature Cards
    col1, col2, col3 = st.columns(3, gap="large")
    
    with col1:
        st.markdown("""
            <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        color: white;
                        border-radius: 20px;
                        padding: 30px;
                        text-align: center;
                        box-shadow: 0 15px 35px rgba(102, 126, 234, 0.3);
                        height: 100%;
                        transform: translateY(0);
                        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);">
                <div style="font-size: 4rem; margin-bottom: 20px; filter: drop-shadow(0 4px 8px rgba(0,0,0,0.3));">📊</div>
                <h3 style="margin-bottom: 20px; color: white; font-size: 1.6rem; font-weight: 700;">Interactive Visualizations</h3>
                <ul style="text-align: left; color: rgba(255, 255, 255, 0.95); font-size: 1.1rem; padding-left: 25px; line-height: 1.8;">
                    <li><strong>15+ Chart Types:</strong> Scatter, Line, Bar, Heatmaps, 3D plots</li>
                    <li><strong>Real-time Filtering:</strong> Dynamic data exploration</li>
                    <li><strong>Custom Styling:</strong> Colors, themes, annotations</li>
                    <li><strong>Interactive Elements:</strong> Zoom, pan, hover details</li>
                </ul>
                <div style="margin-top: 25px; padding: 15px; background: rgba(255,255,255,0.1); border-radius: 10px; font-style: italic; color: rgba(255, 255, 255, 0.9);">
                    "Turn complex data into clear visual stories"
                </div>
            </div>
        """, unsafe_allow_html=True)
        
    with col2:
        st.markdown("""
            <div style="background: linear-gradient(135deg, #ff6b6b 0%, #ffa36c 100%);
                        color: white;
                        border-radius: 20px;
                        padding: 30px;
                        text-align: center;
                        box-shadow: 0 15px 35px rgba(255, 107, 107, 0.3);
                        height: 100%;
                        transform: translateY(0);
                        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);">
                <div style="font-size: 4rem; margin-bottom: 20px; filter: drop-shadow(0 4px 8px rgba(0,0,0,0.3));">🤖</div>
                <h3 style="margin-bottom: 20px; color: white; font-size: 1.6rem; font-weight: 700;">AI-Powered Analytics</h3>
                <ul style="text-align: left; color: rgba(255, 255, 255, 0.95); font-size: 1.1rem; padding-left: 25px; line-height: 1.8;">
                    <li><strong>Smart Insights:</strong> Automated pattern detection</li>
                    <li><strong>Correlation Analysis:</strong> Find hidden relationships</li>
                    <li><strong>Anomaly Detection:</strong> Spot outliers instantly</li>
                    <li><strong>Natural Language:</strong> Ask questions about your data</li>
                </ul>
                <div style="margin-top: 25px; padding: 15px; background: rgba(255,255,255,0.1); border-radius: 10px; font-style: italic; color: rgba(255, 255, 255, 0.9);">
                    "Let AI discover insights you might miss"
                </div>
            </div>
        """, unsafe_allow_html=True)
        
    with col3:
        st.markdown("""
            <div style="background: linear-gradient(135deg, #1d976c 0%, #93f9b9 100%);
                        color: white;
                        border-radius: 20px;
                        padding: 30px;
                        text-align: center;
                        box-shadow: 0 15px 35px rgba(29, 151, 108, 0.3);
                        height: 100%;
                        transform: translateY(0);
                        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);">
                <div style="font-size: 4rem; margin-bottom: 20px; filter: drop-shadow(0 4px 8px rgba(0,0,0,0.3));">🚀</div>
                <h3 style="margin-bottom: 20px; color: white; font-size: 1.6rem; font-weight: 700;">Export & Share</h3>
                <ul style="text-align: left; color: rgba(255, 255, 255, 0.95); font-size: 1.1rem; padding-left: 25px; line-height: 1.8;">
                    <li><strong>Multiple Formats:</strong> PNG, WORD, HTML, SVG, JSON</li>
                    <li><strong>Interactive Reports:</strong> Embed live dashboards</li>
                    <li><strong>Scheduled Exports:</strong> Automated reporting</li>
                    <li><strong>Team Collaboration:</strong> Share insights securely</li>
                </ul>
                <div style="margin-top: 25px; padding: 15px; background: rgba(255,255,255,0.1); border-radius: 10px; font-style: italic; color: rgba(255, 255, 255, 0.9);">
                    "Share insights that drive decisions"
                </div>
            </div>
        """, unsafe_allow_html=True)
    
    # Sample Visualizations Section
    st.markdown("""
        <div style="text-align: center; margin: 60px 0 40px 0;">
            <h2 style="color: #2d3748; font-size: 2.2rem; font-weight: 700; margin-bottom: 20px;">
                📈 Sample Visualizations
            </h2>
            <p style="color: #666; font-size: 1.2rem; max-width: 800px; margin: 0 auto;">
                See what your data could look like with our powerful visualization engine
            </p>
        </div>
    """, unsafe_allow_html=True)
    
    # Sample Charts
    col1, col2 = st.columns(2, gap="large")
    
    with col1:
        st.markdown("### 📊 Revenue Trends Analysis")
        import numpy as np
        sample_data = np.random.normal(100, 15, 30).cumsum()
        st.line_chart(sample_data, height=300)
        st.info("**Interactive Line Charts** - Track trends over time with hover details, zoom controls, and custom styling")
        
        st.markdown("### 🎯 Performance Metrics")
        metrics_data = {"Q1": 85, "Q2": 92, "Q3": 78, "Q4": 96}
        st.bar_chart(metrics_data, height=250)
        st.info("**Dynamic Bar Charts** - Compare categories with animated transitions and drill-down capabilities")
    
    with col2:
        st.markdown("### 🔥 Correlation Heatmap")
        correlation_data = np.random.rand(5, 5)
        st.plotly_chart({
            "data": [{
                "z": correlation_data.tolist(),
                "type": "heatmap",
                "colorscale": "Viridis"
            }],
            "layout": {"height": 300, "margin": {"t": 20, "b": 20, "l": 20, "r": 20}}
        }, use_container_width=True)
        st.info("**Advanced Heatmaps** - Discover relationships between variables with interactive color mapping")
        
        st.markdown("### 📈 Distribution Analysis")
        dist_data = np.random.normal(50, 15, 1000)
        st.plotly_chart({
            "data": [{
                "x": dist_data.tolist(),
                "type": "histogram",
                "nbinsx": 30,
                "marker": {"color": "#667eea"}
            }],
            "layout": {"height": 250, "margin": {"t": 20, "b": 20, "l": 20, "r": 20}}
        }, use_container_width=True)
        st.info("**Statistical Distributions** - Understand your data patterns with histograms and density plots")
    
    # Call to Action
    st.markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border-radius: 20px;
                    padding: 40px;
                    text-align: center;
                    margin: 50px 0;
                    box-shadow: 0 20px 40px rgba(102, 126, 234, 0.3);">
            <h2 style="color: white; font-size: 2.5rem; font-weight: 800; margin-bottom: 20px;">
                🚀 Ready to Transform Your Data?
            </h2>
            <p style="color: rgba(255, 255, 255, 0.9); font-size: 1.3rem; margin-bottom: 30px; line-height: 1.6;">
                Join thousands of data professionals who trust InsightNav AI to turn their data into actionable insights
            </p>
            <div style="display: flex; justify-content: center; gap: 20px; flex-wrap: wrap;">
                <div style="background: rgba(255,255,255,0.1); padding: 20px; border-radius: 15px; min-width: 200px;">
                    <div style="font-size: 2.5rem; font-weight: 800;">10K+</div>
                    <div style="font-size: 1.1rem;">Active Users</div>
                </div>
                <div style="background: rgba(255,255,255,0.1); padding: 20px; border-radius: 15px; min-width: 200px;">
                    <div style="font-size: 2.5rem; font-weight: 800;">1M+</div>
                    <div style="font-size: 1.1rem;">Charts Created</div>
                </div>
                <div style="background: rgba(255,255,255,0.1); padding: 20px; border-radius: 15px; min-width: 200px;">
                    <div style="font-size: 2.5rem; font-weight: 800;">99.9%</div>
                    <div style="font-size: 1.1rem;">Uptime</div>
                </div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    # Login CTA
    col1, col2, col3 = st.columns([2, 3, 2])
    with col2:
        if st.button("🔐 Login to Access Full Dashboard", type="primary", use_container_width=True):
            st.switch_page("analysis_agent.py")
    

def get_aggregated_value(df, column, agg_func):
    """Helper function to calculate aggregated values"""
    if column not in df.columns:
        return 0
    
    if agg_func == "mean":
        return df[column].mean()
    elif agg_func == "sum":
        return df[column].sum()
    elif agg_func == "count":
        return df[column].count()
    elif agg_func == "min":
        return df[column].min()
    elif agg_func == "max":
        return df[column].max()
    elif agg_func == "median":
        return df[column].median()
    elif agg_func == "std":
        return df[column].std()
    elif agg_func == "var":
        return df[column].var()
    else:
        return df[column].mean()



# CRITICAL: Initialize auth system and check for persistent login FIRST
check_persistent_auth() # Check persistent auth directly
if not is_authenticated():
    st.warning("⚠️ Please log in to access the Dashboard.")
    st.stop()

# ENHANCED GLASSMORPHISM STYLING - Beautiful, modern dashboard design
st.markdown("""

<style>
/* ===== BEAUTIFUL GLASSMORPHISM DASHBOARD STYLING ===== */
/* Modern, aesthetic dashboard with glassmorphism effects and enhanced navigation */

/* Main dashboard background */
.main .block-container {
    background: linear-gradient(135deg, 
        rgba(102, 126, 234, 0.03) 0%, 
        rgba(118, 75, 162, 0.05) 50%, 
        rgba(102, 126, 234, 0.03) 100%) !important;
    backdrop-filter: blur(10px) !important;
    padding: 2rem 1rem !important;
}

/* Main sidebar container */
.css-1d391kg, .css-1lcbmhc, .css-1outpf7, .css-17lntkn, .css-1y4p8pa {
    background: linear-gradient(180deg, 
        rgba(102, 126, 234, 0.1) 0%, 
        rgba(118, 75, 162, 0.1) 50%, 
        rgba(102, 126, 234, 0.05) 100%) !important;
    border-right: 2px solid rgba(102, 126, 234, 0.2) !important;
    backdrop-filter: blur(10px) !important;
}

/* Dashboard page title styling */
.main h1 {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
    text-align: center !important;
    font-weight: 700 !important;
    margin-bottom: 2rem !important;
    font-size: 2.5rem !important;
}

/* Enhanced navigation buttons styling */
.stButton > button {
    width: 100% !important;
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 0.75rem 1rem !important;
    font-weight: 600 !important;
    font-size: 0.95rem !important;
    margin: 0.25rem 0 !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3) !important;
    text-align: center !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}

.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 25px rgba(102, 126, 234, 0.4) !important;
    background: linear-gradient(135deg, #5a6fd8 0%, #6a4190 100%) !important;
}

/* Sidebar headers styling */
.css-1lcbmhc h1, .css-1lcbmhc h2, .css-1lcbmhc h3 {
    color: #2d3748 !important;
    text-align: center !important;
    margin-bottom: 1rem !important;
    font-weight: 700 !important;
}

.css-1lcbmhc h2 {
    font-size: 1.4rem !important;
    background: linear-gradient(135deg, #667eea, #764ba2) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
}

/* Sidebar form elements */
.css-1lcbmhc .stSelectbox > div > div {
    background: rgba(255, 255, 255, 0.9) !important;
    border: 2px solid rgba(102, 126, 234, 0.2) !important;
    border-radius: 8px !important;
    transition: all 0.3s ease !important;
}

.css-1lcbmhc .stSelectbox > div > div:hover {
    border-color: rgba(102, 126, 234, 0.5) !important;
    box-shadow: 0 4px 12px rgba(102, 126, 234, 0.15) !important;
}

/* Modern AI Badge - Inspired by top-tier apps like Notion, Linear, GitHub Copilot */
.ai-powered-title {
    display: inline-flex !important;
    align-items: center !important;
    justify-content: center !important;
    margin: 0.75rem auto 1.25rem auto !important;
    padding: 0.5rem 1rem !important;
    background: linear-gradient(135deg, 
        rgba(128, 0, 128, 0.08) 0%, 
        rgba(221, 160, 221, 0.12) 50%, 
        rgba(255, 160, 122, 0.08) 100%) !important;
    border-radius: 50px !important;
    border: 1px solid rgba(128, 0, 128, 0.15) !important;
    backdrop-filter: blur(12px) !important;
    box-shadow: 
        0 2px 8px rgba(128, 0, 128, 0.1),
        0 1px 2px rgba(0, 0, 0, 0.05),
        inset 0 1px 0 rgba(255, 255, 255, 0.1) !important;
    position: relative !important;
    overflow: hidden !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    cursor: default !important;
    max-width: fit-content !important;
}

.ai-powered-title:hover {
    transform: translateY(-1px) !important;
    box-shadow: 
        0 4px 12px rgba(128, 0, 128, 0.15),
        0 2px 4px rgba(0, 0, 0, 0.08),
        inset 0 1px 0 rgba(255, 255, 255, 0.15) !important;
    border-color: rgba(128, 0, 128, 0.25) !important;
}

@keyframes modernColorFlow {
    0%, 100% { 
        color: #800080;
        text-shadow: 0 0 6px rgba(128, 0, 128, 0.3);
    }
    50% { 
        color: #DDA0DD;
        text-shadow: 0 0 8px rgba(221, 160, 221, 0.4);
    }
}

@keyframes sparkle {
    0%, 100% { 
        transform: scale(1) rotate(0deg);
        opacity: 0.8;
    }
    25% { 
        transform: scale(1.1) rotate(90deg);
        opacity: 1;
    }
    50% { 
        transform: scale(0.9) rotate(180deg);
        opacity: 0.9;
    }
    75% { 
        transform: scale(1.05) rotate(270deg);
        opacity: 1;
    }
}

.ai-powered-text {
    font-size: 0.875rem !important;
    font-weight: 600 !important;
    animation: modernColorFlow 6s ease-in-out infinite !important;
    letter-spacing: 0.5px !important;
    position: relative !important;
    display: inline-flex !important;
    align-items: center !important;
    gap: 0.375rem !important;
    text-transform: uppercase !important;
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', system-ui, sans-serif !important;
}

.ai-powered-text::before {
    content: '֎' !important;
    font-size: 1rem !important;
    animation: sparkle 3s ease-in-out infinite !important;
    filter: drop-shadow(0 0 4px rgba(128, 0, 128, 0.3)) !important;
}

/* Navigation button container centering */
.nav-buttons-container {
    display: flex !important;
    flex-direction: column !important;
    align-items: center !important;
    gap: 0.5rem !important;
    margin: 1rem 0 !important;
    padding: 1rem !important;
    background: rgba(255, 255, 255, 0.1) !important;
    border-radius: 12px !important;
    border: 1px solid rgba(102, 126, 234, 0.2) !important;
}

/* Sidebar expanders */
.css-1lcbmhc .streamlit-expanderHeader {
    background: rgba(255, 255, 255, 0.7) !important;
    border-radius: 8px !important;
    border: 1px solid rgba(102, 126, 234, 0.2) !important;
    margin: 0.5rem 0 !important;
    transition: all 0.3s ease !important;
}

.css-1lcbmhc .streamlit-expanderHeader:hover {
    background: rgba(102, 126, 234, 0.1) !important;
    border-color: rgba(102, 126, 234, 0.4) !important;
}

/* Beautiful scrollbar for sidebar */
.css-1lcbmhc::-webkit-scrollbar {
    width: 8px !important;
}

.css-1lcbmhc::-webkit-scrollbar-track {
    background: rgba(102, 126, 234, 0.1) !important;
    border-radius: 4px !important;
}

.css-1lcbmhc::-webkit-scrollbar-thumb {
    background: linear-gradient(180deg, #667eea, #764ba2) !important;
    border-radius: 4px !important;
}

/* ===== CURSOR POINTER STYLING FOR SELECTION BOXES ===== */
/* Change cursor for dropdown controls */
div[data-baseweb="select"] > div:first-child {
    cursor: pointer !important;
}

/* Change cursor for multiselect controls */
div[data-baseweb="select"] div[role="combobox"] {
    cursor: pointer !important;
}

/* Change cursor for slider controls */
div[data-baseweb="slider"] > div > div > div {
    cursor: pointer !important;
}

/* Change cursor for buttons */
button {
    cursor: pointer !important;
}

/* Change cursor for checkbox */
label[data-baseweb="checkbox"] {
    cursor: pointer !important;
}

/* Change cursor for radio buttons */
label[data-baseweb="radio"] {
    cursor: pointer !important;
}

/* Change cursor for number input controls */
div[data-baseweb="input"] input[type="number"] {
    cursor: pointer !important;
}

/* Change cursor for text input controls */
div[data-baseweb="input"] input[type="text"] {
    cursor: text !important;
}

/* Change cursor for file uploader */
div[data-testid="stFileUploader"] {
    cursor: pointer !important;
}
</style>
""", unsafe_allow_html=True)
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import json
from io import StringIO
import base64
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
import time
import openai  # Added for AI functionality
# --- LANGCHAIN IMPORTS ---
from langchain_experimental.agents.agent_toolkits import create_pandas_dataframe_agent
from langchain_openai import ChatOpenAI
from langchain.agents import AgentType
from langchain.memory import ConversationBufferMemory
from langchain.schema import HumanMessage, AIMessage
from langchain.tools import BaseTool
from langchain.agents import AgentExecutor
from langchain.schema import BaseMessage
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_core.agents import AgentAction, AgentFinish
from langchain.schema import OutputParserException

# Initialize global API settings
if 'global_api_key' not in st.session_state:
    st.session_state.global_api_key = ""
if 'global_api_key' not in st.session_state:
    st.session_state.global_api_key = ""
if 'global_ai_provider' not in st.session_state:
    st.session_state.global_ai_provider = "DeepSeek"
if 'global_model_name' not in st.session_state:
    st.session_state.global_model_name = "deepseek-chat"    

# ===== OPTIMIZED DASHBOARD AI HELPER FUNCTIONS =====

# Initialize conversation management for dashboard
if 'dashboard_chat_history' not in st.session_state:
    st.session_state.dashboard_chat_history = []

def manage_dashboard_conversation_history(max_turns=8):
    """Limit dashboard conversation history to prevent token overflow"""
    if len(st.session_state.dashboard_chat_history) > max_turns * 2:
        # Keep system context and recent conversations
        st.session_state.dashboard_chat_history = st.session_state.dashboard_chat_history[-max_turns*2:]

@st.cache_data(ttl=300)  # Cache for 5 minutes
def build_optimized_dashboard_context(df_hash, columns_list, sample_rows, chart_count=0):
    """Cache expensive dashboard context building operations"""
    return {
        'columns': columns_list,
        'sample_size': sample_rows,
        'column_count': len(columns_list),
        'chart_count': chart_count
    }

@st.cache_data(ttl=300)  # Cache for 5 minutes  
def get_dashboard_data_summary(df_hash):
    """Cache dashboard data summary to avoid repeated computation"""
    if current_df is not None:
        numeric_cols = current_df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = current_df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        return {
            'total_rows': len(current_df),
            'numeric_columns': numeric_cols[:5],  # Limit to first 5
            'categorical_columns': categorical_cols[:5],  # Limit to first 5
            'missing_data': current_df.isnull().sum().sum(),
            'data_types': dict(current_df.dtypes.value_counts())
        }
    return {}

def get_comprehensive_chart_knowledge():
    """
    Comprehensive knowledge base of all supported chart types with detailed specifications
    """
    return {
        # === BASIC CHART TYPES ===
        "scatter": {
            "name": "Scatter Plot",
            "icon": "🔍",
            "category": "Basic",
            "description": "Shows relationships between two continuous variables with optional color and size encoding",
            "use_cases": ["Correlation analysis", "Outlier detection", "Pattern identification", "Multi-dimensional analysis"],
            "required_params": ["x_col", "y_col"],
            "optional_params": ["color_col", "size_col", "hover_data", "color_palette"],
            "data_requirements": {
                "min_numeric_cols": 2,
                "min_rows": 10,
                "optimal_rows": "100-10000"
            },
            "insights_provided": ["Correlation strength", "Outliers", "Clusters", "Data distribution patterns"],
            "business_value": "Identify relationships between KPIs, detect anomalies, segment customers"
        },
        
        "bar": {
            "name": "Bar Chart",
            "icon": "📊",
            "category": "Basic",
            "description": "Compares values across categories with optional grouping and aggregation",
            "use_cases": ["Category comparison", "Performance metrics", "Survey results", "Sales by region"],
            "required_params": ["x_col", "y_col"],
            "optional_params": ["color_col", "agg_func", "barmode", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "optimal_categories": "3-20"
            },
            "insights_provided": ["Top performers", "Category rankings", "Group comparisons", "Trend identification"],
            "business_value": "Compare sales performance, analyze market segments, track KPIs"
        },
        
        "stacked_bar": {
            "name": "Stacked Bar Chart",
            "icon": "📊",
            "category": "Basic",
            "description": "Shows part-to-whole relationships within categories",
            "use_cases": ["Composition analysis", "Budget breakdown", "Market share", "Resource allocation"],
            "required_params": ["x_col", "y_col", "color_col"],
            "optional_params": ["agg_func", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 2,
                "min_numeric_cols": 1,
                "optimal_categories": "3-15"
            },
            "insights_provided": ["Composition breakdown", "Relative proportions", "Category contributions", "Trend patterns"],
            "business_value": "Analyze revenue composition, track project progress, understand market dynamics"
        },
        
        "line": {
            "name": "Line Chart",
            "icon": "📈",
            "category": "Basic",
            "description": "Shows trends and changes over time with optional multiple series",
            "use_cases": ["Time series analysis", "Trend tracking", "Performance monitoring", "Forecasting"],
            "required_params": ["x_col", "y_col"],
            "optional_params": ["color_col_primary", "agg_func_primary", "show_markers", "color_palette"],
            "data_requirements": {
                "min_time_points": 5,
                "optimal_time_points": "10-1000",
                "time_column_preferred": True
            },
            "insights_provided": ["Trends", "Seasonality", "Growth rates", "Comparative performance"],
            "business_value": "Track sales growth, monitor KPIs over time, identify seasonal patterns"
        },
        
        "histogram": {
            "name": "Histogram",
            "icon": "📏",
            "category": "Basic",
            "description": "Shows distribution of continuous data with customizable bins",
            "use_cases": ["Distribution analysis", "Data quality assessment", "Statistical analysis", "Outlier detection"],
            "required_params": ["col"],
            "optional_params": ["bins", "color_col", "marginal", "cumulative", "color_palette"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "min_rows": 20,
                "optimal_rows": "100-10000"
            },
            "insights_provided": ["Distribution shape", "Central tendency", "Spread", "Skewness", "Outliers"],
            "business_value": "Understand customer behavior patterns, analyze performance distributions, quality control"
        },
        
        "pie": {
            "name": "Pie Chart",
            "icon": "🥧",
            "category": "Basic",
            "description": "Shows proportional relationships in categorical data",
            "use_cases": ["Market share", "Budget allocation", "Survey responses", "Composition analysis"],
            "required_params": ["col"],
            "optional_params": ["max_categories", "hole", "color_palette", "mode"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "optimal_categories": "3-8",
                "max_recommended_categories": 10
            },
            "insights_provided": ["Proportional breakdown", "Dominant categories", "Market share", "Resource allocation"],
            "business_value": "Visualize market share, show budget allocation, display survey results"
        },
        
        "box": {
            "name": "Box Plot",
            "icon": "📦",
            "category": "Basic",
            "description": "Shows statistical distribution with quartiles, median, and outliers",
            "use_cases": ["Statistical analysis", "Outlier detection", "Group comparison", "Data quality assessment"],
            "required_params": ["num_col", "cat_col"],
            "optional_params": ["color_col", "orientation", "color_palette"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "min_categorical_cols": 1,
                "min_rows": 20
            },
            "insights_provided": ["Median values", "Quartiles", "Outliers", "Distribution spread", "Group differences"],
            "business_value": "Compare performance across regions, identify outliers, assess data quality"
        },
        
        "correlation_heatmap": {
            "name": "Correlation Heatmap",
            "icon": "🧪",
            "category": "Basic",
            "description": "Shows correlation matrix between multiple numeric variables",
            "use_cases": ["Feature selection", "Multicollinearity detection", "Relationship analysis", "Data exploration"],
            "required_params": ["columns"],
            "optional_params": ["color_scale", "text_auto"],
            "data_requirements": {
                "min_numeric_cols": 3,
                "optimal_numeric_cols": "5-20",
                "min_rows": 30
            },
            "insights_provided": ["Variable relationships", "Correlation strength", "Multicollinearity", "Feature importance"],
            "business_value": "Identify key performance drivers, optimize feature selection, understand data relationships"
        },
        
        # === INDUSTRY-STANDARD CHART TYPES ===
        "waterfall": {
            "name": "Waterfall Chart",
            "icon": "💹",
            "category": "Industry-Standard",
            "description": "Shows cumulative effect of sequential positive and negative values",
            "use_cases": ["Financial analysis", "Budget variance", "Performance breakdown", "Process analysis"],
            "required_params": ["category_col", "value_col"],
            "optional_params": ["show_total", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "optimal_categories": "3-15"
            },
            "insights_provided": ["Cumulative impact", "Contributing factors", "Net effect", "Sequential changes"],
            "business_value": "Analyze profit/loss drivers, track budget changes, understand process impacts"
        },
        
        "gauge": {
            "name": "Gauge Chart",
            "icon": "🎯",
            "category": "Industry-Standard",
            "description": "Shows single metric performance against targets with visual indicators",
            "use_cases": ["KPI monitoring", "Performance dashboards", "Goal tracking", "Status indicators"],
            "required_params": ["value_col"],
            "optional_params": ["min_val", "max_val", "target_val", "gauge_type"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "single_value_preferred": True
            },
            "insights_provided": ["Performance level", "Target achievement", "Status indication", "Progress tracking"],
            "business_value": "Monitor KPIs, track goal achievement, display performance status"
        },
        
        "funnel": {
            "name": "Funnel Chart",
            "icon": "📊",
            "category": "Industry-Standard",
            "description": "Shows progressive reduction through process stages",
            "use_cases": ["Sales funnel", "Conversion analysis", "Process optimization", "Customer journey"],
            "required_params": ["stage_col", "value_col"],
            "optional_params": ["show_percentages", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "sequential_stages_preferred": True
            },
            "insights_provided": ["Conversion rates", "Drop-off points", "Process efficiency", "Stage performance"],
            "business_value": "Optimize sales process, improve conversion rates, identify bottlenecks"
        },
        
        "area": {
            "name": "Area Chart",
            "icon": "🌊",
            "category": "Industry-Standard",
            "description": "Shows trends over time with filled areas for emphasis",
            "use_cases": ["Trend analysis", "Cumulative metrics", "Volume tracking", "Time series comparison"],
            "required_params": ["x_col", "y_col"],
            "optional_params": ["color_col", "fill_mode", "color_palette"],
            "data_requirements": {
                "min_time_points": 5,
                "optimal_time_points": "10-500",
                "time_column_preferred": True
            },
            "insights_provided": ["Trend patterns", "Volume changes", "Cumulative growth", "Comparative trends"],
            "business_value": "Track cumulative sales, monitor growth trends, analyze volume patterns"
        },
        
        "treemap": {
            "name": "Treemap",
            "icon": "🌳",
            "category": "Industry-Standard",
            "description": "Shows hierarchical data with nested rectangles proportional to values",
            "use_cases": ["Hierarchical analysis", "Portfolio visualization", "Resource allocation", "Market analysis"],
            "required_params": ["labels_col", "values_col"],
            "optional_params": ["parents_col", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "hierarchical_data_preferred": True
            },
            "insights_provided": ["Hierarchical relationships", "Proportional sizes", "Category dominance", "Structure analysis"],
            "business_value": "Visualize portfolio composition, analyze market segments, show organizational structure"
        },
        
        "violin": {
            "name": "Violin Plot",
            "icon": "🔔",
            "category": "Industry-Standard",
            "description": "Shows distribution shape and statistical summary for groups",
            "use_cases": ["Distribution comparison", "Statistical analysis", "Group analysis", "Data exploration"],
            "required_params": ["x_col", "y_col"],
            "optional_params": ["color_col", "color_palette"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "min_categorical_cols": 1,
                "min_rows": 30
            },
            "insights_provided": ["Distribution shapes", "Group differences", "Statistical summaries", "Data patterns"],
            "business_value": "Compare performance distributions, analyze customer segments, assess data quality"
        },
        
        "bullet": {
            "name": "Bullet Chart",
            "icon": "🎯",
            "category": "Industry-Standard",
            "description": "Compares performance against targets with contextual ranges",
            "use_cases": ["Performance monitoring", "Goal tracking", "Benchmark comparison", "KPI dashboards"],
            "required_params": ["value_col"],
            "optional_params": ["target_col"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "target_values_preferred": True
            },
            "insights_provided": ["Performance vs target", "Achievement levels", "Benchmark comparison", "Goal progress"],
            "business_value": "Monitor KPI achievement, track goals, compare against benchmarks"
        },

        "grouped_bar": {
            "name": "Grouped Bar Chart",
            "icon": "📊",
            "category": "Basic",
            "description": "Shows grouped comparisons across categories with side-by-side bars",
            "use_cases": ["Multi-category comparison", "Performance analysis", "Survey results", "A/B testing"],
            "required_params": ["x_col", "y_col", "color_col"],
            "optional_params": ["agg_func", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 2,
                "min_numeric_cols": 1,
                "optimal_categories": "3-10"
            },
            "insights_provided": ["Group comparisons", "Category performance", "Relative differences", "Pattern identification"],
            "business_value": "Compare performance across multiple dimensions, analyze segmented data, identify patterns"
        },
        
        "dual_axis": {
            "name": "Dual-Axis Chart",
            "icon": "📈",
            "category": "Industry-Standard",
            "description": "Shows two different metrics on separate Y-axes with a shared X-axis",
            "use_cases": ["Correlation analysis", "Multi-metric comparison", "Relationship visualization", "Performance tracking"],
            "required_params": ["x_col", "y_col_primary", "y_col_secondary"],
            "optional_params": ["color_col_primary", "color_col_secondary"],
            "data_requirements": {
                "min_numeric_cols": 2,
                "shared_dimension": True
            },
            "insights_provided": ["Metric relationships", "Correlation patterns", "Comparative trends", "Multi-dimensional analysis"],
            "business_value": "Analyze relationships between different metrics, compare trends, identify correlations"
        },
        
        "pareto": {
            "name": "Pareto Chart",
            "icon": "📊",
            "category": "Industry-Standard",
            "description": "Shows the 80/20 principle with bars representing values and a line showing cumulative percentage",
            "use_cases": ["Priority analysis", "Problem identification", "Root cause analysis", "Quality control"],
            "required_params": ["category_col", "value_col"],
            "optional_params": ["cumulative_percentage", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "optimal_categories": "5-20"
            },
            "insights_provided": ["Vital few vs trivial many", "Priority areas", "Problem significance", "Improvement opportunities"],
            "business_value": "Identify most significant factors, prioritize improvement efforts, focus on high-impact areas"
        },

        "sunburst": {
            "name": "Sunburst Chart",
            "icon": "🌐",
            "category": "Industry-Standard",
            "description": "Shows hierarchical data with concentric rings representing different levels of the hierarchy",
            "use_cases": ["Hierarchical analysis", "Proportional breakdown", "Multi-level composition", "Organizational structure"],
            "required_params": ["labels_col", "values_col"],
            "optional_params": ["parents_col", "max_depth", "branchvalues", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "hierarchical_data_preferred": True
            },
            "insights_provided": ["Hierarchical relationships", "Proportional composition", "Multi-level breakdown", "Structural analysis"],
            "business_value": "Visualize organizational structures, analyze multi-level compositions, show hierarchical relationships"
        },

        "sankey": {
            "name": "Sankey Diagram",
            "icon": "📊",
            "category": "Industry-Standard",
            "description": "Shows flow and relationships between entities with proportional arrow widths",
            "use_cases": ["Flow analysis", "Process mapping", "Resource allocation", "Energy flows"],
            "required_params": ["source_col", "target_col", "value_col"],
            "optional_params": ["node_pad", "node_thickness", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 2,
                "min_numeric_cols": 1,
                "flow_data_required": True
            },
            "insights_provided": ["Flow patterns", "Relationship strength", "Process bottlenecks", "Resource distribution"],
            "business_value": "Analyze process flows, identify bottlenecks, visualize resource allocation, track energy flows"
        },

        "strip": {
            "name": "Strip Chart",
            "icon": "📊",
            "category": "Statistical",
            "description": "Shows individual data points along an axis with optional jitter to avoid overplotting",
            "use_cases": ["Distribution visualization", "Outlier detection", "Data point density", "Category comparison"],
            "required_params": ["x_col", "y_col"],
            "optional_params": ["color_col", "jitter", "stripmode", "color_palette"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "optimal_data_points": "50-1000"
            },
            "insights_provided": ["Data distribution", "Outliers", "Category spread", "Data density"],
            "business_value": "Visualize data distribution, identify outliers, compare category spreads, analyze data density"
        },

        "qq_plot": {
            "name": "Q-Q Plot",
            "icon": "📈",
            "category": "Statistical",
            "description": "Compares data distribution to theoretical distribution to assess normality",
            "use_cases": ["Normality testing", "Distribution comparison", "Statistical validation", "Model diagnostics"],
            "required_params": ["data_col"],
            "optional_params": ["dist", "line"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "min_rows": 20,
                "normality_testing": True
            },
            "insights_provided": ["Distribution fit", "Normality assessment", "Statistical validation", "Model diagnostics"],
            "business_value": "Validate statistical assumptions, assess data normality, compare distributions, diagnostic modeling"
        },

        "density": {
            "name": "Density Plot",
            "icon": "📊",
            "category": "Statistical",
            "description": "Shows the probability density function of continuous data with smooth curves",
            "use_cases": ["Distribution visualization", "Probability estimation", "Data smoothness", "Comparative densities"],
            "required_params": ["data_col"],
            "optional_params": ["group_col", "bandwidth", "cumulative", "fill", "color_palette"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "min_rows": 20,
                "continuous_data": True
            },
            "insights_provided": ["Probability distribution", "Data smoothness", "Comparative densities", "Distribution shape"],
            "business_value": "Understand probability distributions, compare data smoothness, analyze continuous data patterns"
        },

        "ridge": {
            "name": "Ridge Plot",
            "icon": "📊",
            "category": "Statistical",
            "description": "Shows multiple density plots stacked vertically for easy comparison across categories",
            "use_cases": ["Multi-distribution comparison", "Category density analysis", "Visual comparison", "Statistical overview"],
            "required_params": ["data_col", "category_col"],
            "optional_params": ["overlap", "bandwidth", "color_palette"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "min_categorical_cols": 1,
                "min_rows": 30
            },
            "insights_provided": ["Multi-distribution comparison", "Category differences", "Visual patterns", "Statistical overview"],
            "business_value": "Compare distributions across categories, identify pattern differences, visualize multi-category data"
        },

        "timeseries": {
            "name": "Time Series Chart",
            "icon": "📈",
            "category": "Industry-Standard",
            "description": "Shows data points over time with trendlines, confidence intervals, and grouping options",
            "use_cases": ["Trend analysis", "Temporal patterns", "Seasonality detection", "Forecasting preparation"],
            "required_params": ["date_col", "value_col"],
            "optional_params": ["group_col", "agg_func", "show_confidence", "trendline", "color_palette"],
            "data_requirements": {
                "min_datetime_cols": 1,
                "min_numeric_cols": 1,
                "time_series_data": True
            },
            "insights_provided": ["Temporal trends", "Seasonal patterns", "Growth rates", "Anomaly detection"],
            "business_value": "Track performance over time, identify seasonal patterns, detect anomalies, prepare for forecasting"
        },

        "forecast": {
                    "name": "Forecast Chart",
                    "icon": "📈",
                    "category": "Time Series",
                    "description": "Shows historical data with future predictions using statistical models",
                    "use_cases": ["Trend prediction", "Demand forecasting", "Financial planning", "Resource allocation"],
                    "required_params": ["date_col", "value_col"],
                    "optional_params": ["periods", "model_type", "show_confidence"],
                    "data_requirements": {
                        "min_datetime_cols": 1,
                        "min_numeric_cols": 1,
                        "min_rows": 20,
                        "time_series_data": True
                    },
                    "insights_provided": ["Future trends", "Seasonal patterns", "Growth projections", "Confidence intervals"],
                    "business_value": "Predict future performance, plan resources, identify growth opportunities"
                },
                
                "moving_average": {
                    "name": "Moving Average Chart",
                    "icon": "📊",
                    "category": "Time Series",
                    "description": "Shows smoothed data trends using simple or exponential moving averages",
                    "use_cases": ["Trend identification", "Noise reduction", "Signal extraction", "Technical analysis"],
                    "required_params": ["date_col", "value_col"],
                    "optional_params": ["window", "ma_type", "show_original"],
                    "data_requirements": {
                        "min_datetime_cols": 1,
                        "min_numeric_cols": 1,
                        "min_rows": 15,
                        "time_series_data": True
                    },
                    "insights_provided": ["Trend direction", "Smoothing patterns", "Volatility reduction", "Signal clarity"],
                    "business_value": "Identify underlying trends, reduce noise in data, improve signal clarity"
                },
                
                "ranking": {
                    "name": "Ranking Chart",
                    "icon": "🏆",
                    "category": "Comparative",
                    "description": "Shows top or bottom items ranked by a specific metric",
                    "use_cases": ["Performance ranking", "Top performers", "Benchmarking", "Competitive analysis"],
                    "required_params": ["category_col", "value_col"],
                    "optional_params": ["top_n", "sort_order", "orientation", "color_palette"],
                    "data_requirements": {
                        "min_categorical_cols": 1,
                        "min_numeric_cols": 1,
                        "optimal_categories": "10-50"
                    },
                    "insights_provided": ["Top performers", "Rank order", "Performance gaps", "Comparative analysis"],
                    "business_value": "Identify best/worst performers, benchmark performance, prioritize resources"
                },
        "seasonal": {
            "name": "Seasonal Decomposition",
            "icon": "📊",
            "category": "Time Series",
            "description": "Decomposes time series data into trend, seasonal, and residual components",
            "use_cases": ["Seasonality analysis", "Trend identification", "Time series diagnostics", "Pattern detection"],
            "required_params": ["date_col", "value_col"],
            "optional_params": ["model_type", "period"],
            "data_requirements": {
                "min_datetime_cols": 1,
                "min_numeric_cols": 1,
                "min_rows": 24,
                "time_series_data": True
            },
            "insights_provided": ["Seasonal patterns", "Trend direction", "Residual analysis", "Time series components"],
            "business_value": "Identify seasonal patterns, understand underlying trends, detect anomalies in time series data"
        },

        "comparison": {
            "name": "Comparison Chart",
            "icon": "📈",
            "category": "Comparative",
            "description": "Compares multiple metrics across categories using bars, lines, or areas",
            "use_cases": ["Multi-metric comparison", "Performance benchmarking", "Feature comparison", "KPI analysis"],
            "required_params": ["category_col", "value_cols"],
            "optional_params": ["chart_type", "normalize"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 2,
                "optimal_categories": "3-15"
            },
            "insights_provided": ["Relative performance", "Metric comparisons", "Benchmarking insights", "Multi-dimensional analysis"],
            "business_value": "Compare multiple KPIs, benchmark performance, analyze relative strengths and weaknesses"
        },

        "slope": {
            "name": "Slope Chart",
            "icon": "📉",
            "category": "Comparative",
            "description": "Shows changes between two or more time points with connecting lines",
            "use_cases": ["Change over time", "Progress tracking", "Before-after analysis", "Trend comparison"],
            "required_params": ["category_col", "time_col", "value_col"],
            "optional_params": [],
            "data_requirements": {
                "min_categorical_cols": 2,
                "min_numeric_cols": 1,
                "time_periods": 2
            },
            "insights_provided": ["Change magnitude", "Direction of change", "Relative performance shifts", "Progress indicators"],
            "business_value": "Track progress over time, compare changes across categories, visualize performance shifts"
        },

        "dot_plot": {
            "name": "Dot Plot",
            "icon": "🔵",
            "category": "Comparative",
            "description": "Shows individual data points along an axis with optional grouping",
            "use_cases": ["Distribution visualization", "Group comparisons", "Value distribution", "Statistical analysis"],
            "required_params": ["category_col", "value_col"],
            "optional_params": ["group_col", "orientation"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "optimal_data_points": "20-500"
            },
            "insights_provided": ["Value distribution", "Group differences", "Central tendency", "Data spread"],
            "business_value": "Compare distributions across groups, visualize value ranges, analyze data spread"
        },        

        
        # === SPECIAL COMPONENTS ===
        "kpi": {
            "name": "Advanced KPI Dashboard",
            "icon": "✅",
            "category": "Performance",
            "description": "Modern KPI dashboard with multiple metrics, comparisons, and visual enhancements",
            "use_cases": ["Executive dashboards", "Performance monitoring", "Real-time metrics", "Business reporting"],
            "required_params": ["metrics"],
            "optional_params": ["agg_funcs", "show_comparison", "comparison_type", "layout_style", "show_sparklines"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "optimal_numeric_cols": "3-8",
                "time_series_support": True
            },
            "insights_provided": ["Key performance indicators", "Trend analysis", "Goal tracking", "Performance benchmarks"],
            "business_value": "Monitor business health, track KPIs, executive reporting, performance management"
        },
        
        "table": {
            "name": "Data Table",
            "icon": "📋",
            "category": "Special",
            "description": "Displays structured data in tabular format",
            "use_cases": ["Detailed data view", "Record listing", "Data exploration", "Reference information"],
            "required_params": ["columns"],
            "optional_params": ["max_rows"],
            "data_requirements": {
                "flexible_columns": True,
                "any_data_type": True
            },
            "insights_provided": ["Detailed records", "Data structure", "Individual values", "Complete information"],
            "business_value": "Show detailed records, provide data reference, enable data exploration"
        },

        "traffic_light": {
            "name": "Advanced Traffic Light Dashboard",
            "icon": "🚦",
            "category": "Performance",
            "description": "Multi-KPI traffic light system with configurable thresholds and visual indicators for at-a-glance performance monitoring.",
            "use_cases": ["KPI monitoring", "Performance dashboards", "Status reporting", "Executive summaries", "Operational health checks"],
            "required_params": ["kpi_metrics"],
            "optional_params": ["red_thresholds", "yellow_thresholds", "green_thresholds", "layout_style", "show_sparklines", "show_values"],
            "data_requirements": {
                "min_numeric_cols": 1,
                "optimal_numeric_cols": "3-8",
                "categorical_cols": 0
            },
            "insights_provided": ["Performance status vs. targets", "Threshold compliance", "Multi-metric overview", "Alert conditions"],
            "business_value": "Enables quick, visual assessment of multiple key performance indicators against predefined targets, facilitating rapid decision-making and executive reporting."
        },
        
        "choropleth": {
            "name": "Advanced Choropleth Map",
            "icon": "🗺️",
            "category": "Geographic",
            "description": "Interactive geographic map that colors regions based on data values, with advanced styling and animation options. Supports standard maps (countries, states) and custom GeoJSON for cities, counties, etc.",
            "use_cases": ["Regional sales analysis", "Geographic distribution of customers", "Market penetration analysis", "Territory management"],
            "required_params": ["location_col", "value_col"],
            "optional_params": ["locationmode", "scope", "color_scale", "use_custom_geojson", "featureidkey"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "geographic_data": True # This is a key requirement
            },
            "insights_provided": ["Geographic patterns and hotspots", "Regional performance comparisons", "Spatial distribution of metrics", "Territory performance"],
            "business_value": "Visualizes data geographically to uncover regional trends, optimize territory management, and inform location-based business strategies."
        },
        
        "donut": {
            "name": "Multi-Layer Donut Chart",
            "icon": "🍩",
            "category": "Composition",
            "description": "Hierarchical donut chart that displays proportional data across multiple levels with a central KPI display.",
            "use_cases": ["Multi-level composition analysis", "Hierarchical data breakdown", "Budget allocation by department and sub-department", "Product category analysis"],
            "required_params": ["levels", "level_0_col", "level_0_value"],
            "optional_params": ["level_X_col", "level_X_value", "center_metric", "hole_size", "show_percentages"],
            "data_requirements": {
                "min_categorical_cols": 1,
                "min_numeric_cols": 1,
                "hierarchical_data": True # Ideal for this chart
            },
            "insights_provided": ["Hierarchical composition", "Proportional analysis across levels", "Part-to-whole relationships", "Central KPI context"],
            "business_value": "Clearly visualizes complex hierarchical data, showing how smaller components contribute to larger totals, ideal for financial and organizational analysis."
        }
    }

def get_chart_recommendations(df, user_intent=""):
    """
    Provide intelligent chart recommendations based on data characteristics and user intent
    """
    chart_knowledge = get_comprehensive_chart_knowledge()
    
    # Analyze data characteristics
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
    datetime_cols = df.select_dtypes(include=['datetime64']).columns.tolist()
    
    recommendations = []
    
    # Basic recommendations based on data structure
    if len(numeric_cols) >= 2:
        recommendations.append({
            "chart_type": "scatter",
            "reason": f"Perfect for exploring relationships between {numeric_cols[0]} and {numeric_cols[1]}",
            "confidence": 0.9
        })
        
        recommendations.append({
            "chart_type": "correlation_heatmap", 
            "reason": f"Analyze correlations between all {len(numeric_cols)} numeric variables",
            "confidence": 0.8
        })
    
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "bar",
            "reason": f"Compare {numeric_cols[0]} across {categorical_cols[0]} categories",
            "confidence": 0.9
        })
        
        recommendations.append({
            "chart_type": "box",
            "reason": f"Analyze {numeric_cols[0]} distribution by {categorical_cols[0]}",
            "confidence": 0.7
        })
    
    if len(datetime_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "line",
            "reason": f"Track {numeric_cols[0]} trends over {datetime_cols[0]}",
            "confidence": 0.95
        })
        
        recommendations.append({
            "chart_type": "area",
            "reason": f"Visualize {numeric_cols[0]} volume changes over time",
            "confidence": 0.8
        })
    
    # Intent-based recommendations
    intent_lower = user_intent.lower()
    if any(word in intent_lower for word in ['trend', 'time', 'over time', 'growth']):
        recommendations.append({
            "chart_type": "line",
            "reason": "Line charts are ideal for trend analysis",
            "confidence": 0.9
        })
    
    if any(word in intent_lower for word in ['compare', 'comparison', 'versus', 'vs']):
        recommendations.append({
            "chart_type": "bar",
            "reason": "Bar charts excel at category comparisons",
            "confidence": 0.9
        })
    
    if any(word in intent_lower for word in ['distribution', 'spread', 'histogram']):
        recommendations.append({
            "chart_type": "histogram",
            "reason": "Histograms show data distribution patterns",
            "confidence": 0.9
        })
    if len(categorical_cols) >= 2 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "grouped_bar",
            "reason": f"Compare {numeric_cols[0]} across {categorical_cols[0]} and {categorical_cols[1]}",
            "confidence": 0.8
        })
    
    if len(numeric_cols) >= 2:
        recommendations.append({
            "chart_type": "dual_axis",
            "reason": f"Compare relationship between {numeric_cols[0]} and {numeric_cols[1]}",
            "confidence": 0.7
        })
    
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "pareto",
            "reason": f"Identify most significant categories in {categorical_cols[0]} by {numeric_cols[0]}",
            "confidence": 0.75
        })
    # For sunburst chart
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "sunburst",
            "reason": f"Show hierarchical relationships between {categorical_cols[0]} and {numeric_cols[0]}",
            "confidence": 0.7
        }) 
    # For sankey diagram
    if len(categorical_cols) >= 2 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "sankey",
            "reason": f"Show flow relationships between {categorical_cols[0]} and {categorical_cols[1]} with {numeric_cols[0]}",
            "confidence": 0.7
        })

    # For strip chart
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "strip",
            "reason": f"Show individual data points of {numeric_cols[0]} across {categorical_cols[0]} categories",
            "confidence": 0.75
        })

    # For Q-Q plot
    if len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "qq_plot",
            "reason": f"Check if {numeric_cols[0]} follows a normal distribution",
            "confidence": 0.65
        })

    # For density plot
    if len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "density",
            "reason": f"Visualize the probability distribution of {numeric_cols[0]}",
            "confidence": 0.7
        })

    # For ridge plot
    if len(numeric_cols) >= 1 and len(categorical_cols) >= 1:
        recommendations.append({
            "chart_type": "ridge",
            "reason": f"Compare distributions of {numeric_cols[0]} across {categorical_cols[0]} categories",
            "confidence": 0.75
        })

    # For time series
    datetime_cols = [col for col in df.columns if pd.api.types.is_datetime64_any_dtype(df[col])]
    if datetime_cols and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "timeseries",
            "reason": f"Track {numeric_cols[0]} over {datetime_cols[0]} with trend analysis",
            "confidence": 0.9
        }) 

    datetime_cols = [col for col in df.columns if pd.api.types.is_datetime64_any_dtype(df[col])]
    if datetime_cols and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "forecast",
            "reason": f"Forecast future values of {numeric_cols[0]} based on {datetime_cols[0]}",
            "confidence": 0.8
        })
        
        recommendations.append({
            "chart_type": "moving_average", 
            "reason": f"Smooth out noise in {numeric_cols[0]} over {datetime_cols[0]}",
            "confidence": 0.7
        })
    
    # For ranking chart
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "ranking",
            "reason": f"Rank {categorical_cols[0]} by {numeric_cols[0]}",
            "confidence": 0.9
        })

    # For seasonal decomposition
    if len(datetime_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "seasonal",
            "reason": f"Analyze seasonality and trends in {numeric_cols[0]} over time",
            "confidence": 0.8
        })

    # For comparison chart
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 2:
        recommendations.append({
            "chart_type": "comparison",
            "reason": f"Compare multiple metrics ({', '.join(numeric_cols[:3])}) across {categorical_cols[0]} categories",
            "confidence": 0.7
        })

    # For slope chart
    if len(categorical_cols) >= 2 and len(numeric_cols) >= 1:
        # Check if we have at least 2 time periods
        time_col_candidate = categorical_cols[0]  # Assume first categorical could be time
        if df[time_col_candidate].nunique() >= 2:
            recommendations.append({
                "chart_type": "slope",
                "reason": f"Track changes in {numeric_cols[0]} across time periods in {categorical_cols[1]}",
                "confidence": 0.6
            })

    # For dot plot
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "dot_plot",
            "reason": f"Show distribution of {numeric_cols[0]} across {categorical_cols[0]} categories",
            "confidence": 0.7
        }) 

    if len(numeric_cols) >= 3:
        recommendations.append({
            "chart_type": "kpi",
            "reason": f"Perfect for monitoring key metrics like {', '.join(numeric_cols[:3])}",
            "confidence": 0.8
        })

    if len(numeric_cols) >= 3:
        recommendations.append({
            "chart_type": "traffic_light",
            "reason": f"Perfect for monitoring multiple KPIs like {', '.join(numeric_cols[:3])} with traffic light status.",
            "confidence": 0.7
        })
    
    # Check for geographic data
    geographic_keywords = ['country', 'state', 'region', 'city', 'zip', 'postal', 'latitude', 'longitude', 'iso_alpha', 'iso_code']
    geographic_cols = [col for col in df.columns if any(keyword in col.lower() for keyword in geographic_keywords)]
    
    if geographic_cols and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "choropleth",
            "reason": f"Visualize geographic distribution of {numeric_cols[0]} across the '{geographic_cols[0]}' column.",
            "confidence": 0.85 # Higher confidence if geo data is likely present
        })
    
    if len(categorical_cols) >= 2 and len(numeric_cols) >= 1:
        recommendations.append({
            "chart_type": "donut",
            "reason": f"Show a hierarchical relationship between '{categorical_cols[0]}' and '{categorical_cols[1]}' using '{numeric_cols[0]}' as the value.",
            "confidence": 0.65
        })                             
    
    # Sort by confidence and return top recommendations
    recommendations.sort(key=lambda x: x['confidence'], reverse=True)
    return recommendations[:5]

def get_validated_dashboard_llm():
    """Single source of truth for Dashboard LLM initialization with global API key"""
    api_key = st.session_state.get('global_api_key')
    if not api_key:
        raise ValueError("Please enter your API key in the Analysis Agent page to enable the AI Assistant.")
    
    provider = st.session_state.get('global_ai_provider', 'DeepSeek')
    model_name = st.session_state.get('global_model_name', 'deepseek-chat')
    
    # Map providers to their respective classes (matching analysis_agent.py exactly)
    try:
        if provider == "DeepSeek":
            from langchain_openai import ChatOpenAI
            return ChatOpenAI(
                api_key=api_key,
                model="deepseek-chat",
                base_url="https://api.deepseek.com/v1",
                temperature=0.7,
                max_tokens=4000
            )
        elif provider == "OpenAI":
            from langchain_openai import ChatOpenAI
            return ChatOpenAI(
                api_key=api_key,
                model=model_name,
                temperature=0.7,
                max_tokens=4000
            )
        elif provider == "Anthropic":
            from langchain_anthropic import ChatAnthropic
            return ChatAnthropic(
                api_key=api_key,
                model_name=model_name,
                temperature=0.7,
                max_tokens=4000
            )
        elif provider == "Google":
            from langchain_google_genai import ChatGoogleGenerativeAI
            return ChatGoogleGenerativeAI(
                model=model_name,
                google_api_key=api_key,
                temperature=0.7,
                max_output_tokens=4000
            )
        else:
            # Fallback to DeepSeek (most cost-effective)
            from langchain_openai import ChatOpenAI
            return ChatOpenAI(
                api_key=api_key,
                model="deepseek-chat",
                base_url="https://api.deepseek.com/v1",
                temperature=0.7,
                max_tokens=4000
            )
    except ImportError as e:
        # If specific provider library is not available, fallback to OpenAI format
        st.warning(f"⚠️ {provider} library not available, using OpenAI format. Install required packages for full support.")
        from langchain_openai import ChatOpenAI
        return ChatOpenAI(api_key=api_key, model="gpt-3.5-turbo", temperature=0.7)

def get_dashboard_summary(dashboard_config):
    """Create a summary of dashboard configuration"""
    if not dashboard_config:
        return "No dashboard created yet"
    
    charts = dashboard_config.get('charts', [])
    chart_types = [chart.get('type', 'unknown') for chart in charts]
    
    return f"Dashboard with {len(charts)} charts: {', '.join(set(chart_types))}"

def cleanup_dashboard_session_state():
    """Clean up unused dashboard objects from session state"""
    cleanup_keys = ['temp_dashboard_agent', 'old_dashboard_contexts', 'cached_dashboard_responses']
    for key in cleanup_keys:
        if key in st.session_state:
            del st.session_state[key]
    
    # Force garbage collection
    import gc
    gc.collect()

# ===== DASHBOARD EXPERT ENHANCEMENT FUNCTIONS =====

def detect_dashboard_context_and_suggest():
    """Detect current dashboard context and provide smart suggestions"""
    
    context = {
        'stage': 'unknown',
        'suggestions': [],
        'warnings': [],
        'next_steps': []
    }
    
    # Detect current stage
    if not st.session_state.get('dashboard'):
        context['stage'] = 'dashboard_creation'
        context['suggestions'].append("🎯 Start by creating your first dashboard")
        context['next_steps'].append("Describe what you want to visualize in the AI prompt")
    
    else:
        dashboard = st.session_state.dashboard
        charts = dashboard.get('charts', [])
        
        if len(charts) == 0:
            context['stage'] = 'empty_dashboard'
            context['warnings'].append("📊 Your dashboard has no charts")
            context['next_steps'].append("Add charts to visualize your data")
        
        elif len(charts) < 3:
            context['stage'] = 'basic_dashboard'
            # Simplified - removed cluttered suggestions
        
        else:
            context['stage'] = 'complete_dashboard'
            # Simplified - removed cluttered suggestions and warnings
    
    # Data-specific suggestions
    if current_df is not None:
        missing_data = current_df.isnull().sum().sum()
        if missing_data > 0:
            context['warnings'].append(f"⚠️ Your data has {missing_data} missing values")
        
        # Check for time series potential
        date_cols = [col for col in current_df.columns if 'date' in col.lower() or 'time' in col.lower()]
        current_dashboard = st.session_state.get('dashboard')
        if date_cols and current_dashboard and not any('line' in chart.get('type', '') for chart in current_dashboard.get('charts', [])):
            context['suggestions'].append("📈 Consider a line chart for time series data")
    
    return context

def get_smart_dashboard_questions():
    """Generate context-aware dashboard questions"""
    questions = []
    
    # Base questions always available
    base_questions = [
        "🎯 What chart type best shows my data relationships?",
        "🎨 How can I improve my dashboard's visual design?",
        "📊 What story does my data tell?",
        "⚡ How can I make my dashboard more interactive?",
        "🌈 What color scheme works best for my data?",
        "📱 How do I make my dashboard mobile-friendly?",
        "🚀 What's the best way to share my dashboard?",
        "📈 How do I highlight key insights?"
    ]
    
    # Enhanced comprehensive analysis questions
    comprehensive_questions = [
        "🔬 Generate comprehensive dashboard analysis report",
        "📊 Analyze each chart individually with insights",
        "🧠 What cross-chart relationships exist in my dashboard?",
        "📈 What trends and patterns can you identify?",
        "💼 What business insights emerge from my dashboard?",
        "🎯 How effective is my current dashboard design?",
        "📋 What optimization recommendations do you have?",
        "🔍 What hidden patterns exist in my data visualizations?"
    ]
    
    # Context-aware questions
    if not st.session_state.get('dashboard'):
        questions.extend([
            "🏁 How do I get started with dashboard creation?",
            "📋 What makes a good dashboard?",
            "🎲 Should I start with simple or complex visualizations?"
        ])
    else:
        dashboard = st.session_state.dashboard
        charts = dashboard.get('charts', [])
        
        if len(charts) < 3:
            questions.extend([
                "📊 What additional charts should I add?",
                "🔍 How do I identify missing insights?",
                "📈 What chart types complement each other?"
            ])
        else:
            questions.extend([
                "🎨 How do I optimize my dashboard layout?",
                "📖 How do I create a compelling data story?",
                "⚡ How do I improve dashboard performance?"
            ])
            # Add comprehensive analysis questions when dashboard exists
            questions.extend(comprehensive_questions)
        
        # Chart-specific questions
        chart_types = [chart.get('type', 'unknown') for chart in charts]
        if 'scatter' in chart_types:
            questions.extend([
                "🔍 How do I interpret scatter plot patterns?",
                "📊 What correlations does my scatter plot reveal?"
            ])
        if 'bar' in chart_types:
            questions.extend([
                "📊 How do I make my bar charts more effective?",
                "📈 What comparisons does my bar chart highlight?"
            ])
        if 'line' in chart_types:
            questions.extend([
                "📈 What trends does my line chart show?",
                "🔮 Can you forecast future trends from my line chart?"
            ])
        if 'correlation_heatmap' in chart_types:
            questions.extend([
                "🔥 What does my correlation heatmap reveal?",
                "🧠 Which variables are most strongly related?"
            ])
        if 'histogram' in chart_types:
            questions.extend([
                "📊 What does my data distribution tell me?",
                "🔍 Are there any outliers in my histogram?"
            ])
    
    return base_questions + questions

def analyze_data_for_chart_recommendations(df):
    """Analyze data characteristics for intelligent chart recommendations"""
    
    analysis = {
        'numeric_columns': df.select_dtypes(include=[np.number]).columns.tolist(),
        'categorical_columns': df.select_dtypes(include=['object', 'category']).columns.tolist(),
        'datetime_columns': [],
        'data_size': len(df),
        'missing_data': df.isnull().sum().sum(),
        'recommendations': []
    }
    
    # Detect potential datetime columns
    for col in df.columns:
        if 'date' in col.lower() or 'time' in col.lower():
            analysis['datetime_columns'].append(col)
    
    # Generate recommendations based on data characteristics
    numeric_cols = analysis['numeric_columns']
    categorical_cols = analysis['categorical_columns']
    datetime_cols = analysis['datetime_columns']
    
    # Scatter plot recommendations
    if len(numeric_cols) >= 2:
        analysis['recommendations'].append({
            'type': 'scatter',
            'title': 'Scatter Plot - Explore Relationships',
            'reasoning': f'You have {len(numeric_cols)} numeric columns. Scatter plots reveal correlations and patterns.',
            'columns_suggested': numeric_cols[:2],
            'priority': 'high'
        })
    
    # Bar chart recommendations
    if len(categorical_cols) >= 1 and len(numeric_cols) >= 1:
        analysis['recommendations'].append({
            'type': 'bar',
            'title': 'Bar Chart - Compare Categories',
            'reasoning': f'Perfect for comparing {categorical_cols[0]} across different values.',
            'columns_suggested': [categorical_cols[0], numeric_cols[0]],
            'priority': 'high'
        })
    
    # Line chart recommendations
    if datetime_cols and numeric_cols:
        analysis['recommendations'].append({
            'type': 'line',
            'title': 'Line Chart - Show Trends Over Time',
            'reasoning': f'Ideal for showing how {numeric_cols[0]} changes over time.',
            'columns_suggested': [datetime_cols[0], numeric_cols[0]],
            'priority': 'high'
        })
    
    # Histogram recommendations
    if numeric_cols:
        analysis['recommendations'].append({
            'type': 'histogram',
            'title': 'Histogram - Show Distribution',
            'reasoning': f'Understand the distribution of {numeric_cols[0]} values.',
            'columns_suggested': [numeric_cols[0]],
            'priority': 'medium'
        })
    
    # Pie chart recommendations (with caution)
    if categorical_cols:
        unique_values = df[categorical_cols[0]].nunique()
        if unique_values <= 6:
            analysis['recommendations'].append({
                'type': 'pie',
                'title': 'Pie Chart - Show Proportions',
                'reasoning': f'Good for showing proportions of {categorical_cols[0]} (has {unique_values} categories).',
                'columns_suggested': [categorical_cols[0]],
                'priority': 'low',
                'warning': 'Use sparingly - bar charts often work better'
            })
    
    # Correlation heatmap recommendations
    if len(numeric_cols) >= 3:
        analysis['recommendations'].append({
            'type': 'correlation_heatmap',
            'title': 'Correlation Heatmap - Find Relationships',
            'reasoning': f'Discover correlations between your {len(numeric_cols)} numeric variables.',
            'columns_suggested': numeric_cols,
            'priority': 'medium'
        })
    
    return analysis

def analyze_dashboard_design(dashboard):
    """Simplified dashboard design analysis - returns minimal feedback to reduce clutter"""
    
    # Return minimal analysis to prevent clutter
    analysis = {
        'strengths': [],
        'improvements': [],
        'recommendations': []
    }
    
    # Only add a simple strength if dashboard exists
    if dashboard and dashboard.get('charts'):
        chart_count = len(dashboard.get('charts', []))
        if chart_count > 0:
            analysis['strengths'].append(f"Dashboard created with {chart_count} chart{'s' if chart_count != 1 else ''}")
    
    return analysis

# ===== END DASHBOARD EXPERT ENHANCEMENT FUNCTIONS =====

# --- DASHBOARD EXPERT TOOL ---
class DashboardExpertTool:
    """LangChain Tool for providing Dashboard expertise."""

    def __init__(self, name="Dashboard_Expert"):
        self.name = name

    def run(self, query: str) -> str:
        """
        Provide Dashboard expertise based on the query and current session state.
        Args:
            query (str): The user's question about dashboards.
        Returns:
            str: The expert's response.
        """
        try:
            response_parts = [f"📊 Response from {self.name}:"]
            
            # --- Access Data and Dashboard Config from Session State ---
            if 'df_sample' not in st.session_state or st.session_state.df_sample is None:
                response_parts.append("❌ No data sample found in session state. Please create a dashboard in the Dashboard page first.")
                return "\n".join(response_parts)

            df_sample = st.session_state.df_sample
            response_parts.append(f"📈 **Data Context:** Sampled data with {len(df_sample)} rows and {len(df_sample.columns)} columns.")

            if 'dashboard' not in st.session_state or not st.session_state.dashboard:
                response_parts.append("ℹ️ No dashboard configuration found in session state. Please create a dashboard in the Dashboard page first.")
                response_parts.append("💡 I can still provide general dashboard design advice:")
                response_parts.append(self._provide_general_dashboard_advice(query))
                return "\n".join(response_parts)

            dashboard_config = st.session_state.dashboard
            response_parts.append(f"🧱 **Dashboard Context:** Configuration loaded with {len(dashboard_config.get('charts', []))} charts.")

            # --- Provide Dashboard Interpretation ---
            charts = dashboard_config.get("charts", [])
            if charts:
                response_parts.append("\n🖼️ **Dashboard Interpretation:**")
                for i, chart_config in enumerate(charts[:3]): # Limit to first 3 for brevity
                    chart_type = chart_config.get("type", "unknown")
                    params = chart_config.get("params", {})
                    title = (chart_config.get("title") or
                             params.get("custom_title", "") or
                             chart_type.replace('_', ' ').title())
                    response_parts.append(f"  - Chart {i+1} ({title}): Type '{chart_type}'.")
                    # Add brief interpretation based on chart type
                    if "scatter" in chart_type.lower():
                        response_parts.append("    ➤ Shows relationship between two numerical variables.")
                    elif "bar" in chart_type.lower():
                        response_parts.append("    ➤ Compares quantities across categories.")
                    elif "line" in chart_type.lower():
                        response_parts.append("    ➤ Shows trends over time or ordered categories.")
                    elif "pie" in chart_type.lower():
                        response_parts.append("    ➤ Shows proportions of a whole.")
                    elif "heatmap" in chart_type.lower():
                        response_parts.append("    ➤ Shows correlation or intensity in a matrix format.")
                    elif "box" in chart_type.lower():
                        response_parts.append("    ➤ Shows distribution and outliers for numerical data.")
                    elif "histogram" in chart_type.lower():
                        response_parts.append("    ➤ Shows frequency distribution of a numerical variable.")
                    elif "kpi" in chart_type.lower():
                        response_parts.append("    ➤ Displays a key performance indicator.")
                    elif "table" in chart_type.lower():
                        response_parts.append("    ➤ Shows raw data in a tabular format.")
                    # Add interpretations for other chart types as needed
            
            if len(charts) > 3:
                response_parts.append(f"... (and {len(charts) - 3} more charts)")

            # --- Respond to Specific Query with Real AI ---
            try:
                llm = get_validated_dashboard_llm()
                context_summary = "\n".join(response_parts)
                
                prompt = f"""
                Dashboard Context: {context_summary}
                User Query: {query}
                
                Provide specific, actionable advice about dashboard design and data visualization.
                Focus on practical recommendations and insights.
                """
                
                ai_response = llm.invoke(prompt).content
                return ai_response
                
            except Exception as e:
                response_parts.append(f"\n🔍 **Regarding your query:** '{query}'")
                response_parts.append("I've analyzed the dashboard configuration and provided a summary above.")
                response_parts.append(f"Note: AI response unavailable ({str(e)})")
                return "\n".join(response_parts)

        except Exception as e:
            return f"❌ Error in {self.name}.run: {e}"

    def _provide_general_dashboard_advice(self, query: str) -> str:
        """Provide general Dashboard advice if no specific config is found."""
        # This could be a simple lookup or a call to a general-purpose LLM
        # For now, a placeholder
        return f"General Dashboard advice for query: '{query}' is not yet implemented in this tool."

# --- END DASHBOARD EXPERT TOOL ---

# Function to encode your local image
def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

# Set your background image
def set_background_image(image_path):
    bin_str = get_base64_of_bin_file(image_path)
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-image: url("data:image/jpg;base64,{bin_str}");
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

# Call this function with your image path
set_background_image('light_splash.avif') # Make sure 'mesh.JPG' exists

# --- 1. Persistent Authentication Check ---
# Check for persistent authentication first (handles page refresh)
if not st.session_state.get('user'):
    check_persistent_auth()

# --- 2. Authentication Check ---
if not is_authenticated():
    # Back to Login button at top left
    col1, col2, col3 = st.columns([1, 6, 1])
    with col1:
        if st.button("🏠 Back to Login", type="primary", use_container_width=True):
            st.switch_page("analysis_agent.py")

# Set page config
st.set_page_config(
    page_title="InsightNav AI - Dashboard",
    layout="wide",
    page_icon="📊"
)

# Display logo in top left corner
display_logo()

# Display logo in sidebar
display_sidebar_logo()

# --- END Page Configuration ---

# --- LOAD DATA WITH CORRECT PRIORITY (UPDATED) ---
# Priority 1: Use explicitly sampled dataset
current_df = None
df_sample = None

# --- 1. HIGHEST Priority: Use cleaned data from advanced cleaning (df_cleaned or cleaned_df) ---
df_cleaned = st.session_state.get('df_cleaned')
cleaned_df = st.session_state.get('cleaned_df')
cleaned_data = df_cleaned if df_cleaned is not None else cleaned_df
if cleaned_data is not None:
    current_df = cleaned_data
    # Create a proper sample from cleaned data using analysis_agent sample size
    sample_size = st.session_state.get('sample_size_slider', st.session_state.get('sample_size', min(2000, len(current_df))))
    df_sample = current_df.sample(n=sample_size, random_state=42) if len(current_df) > sample_size else current_df.copy()
    # Update session state to maintain consistency
    st.session_state.df_sample = df_sample
    data_loaded = True
    # st.info(f"ℹ️ Using cleaned data from Analysis Agent (Sample: {len(df_sample):,} rows, Full: {len(current_df):,} rows).")

# --- 2. Priority: Use the specific sample created by analysis_agent.py ---
elif 'df_sample' in st.session_state and st.session_state.df_sample is not None:
    # Use the actual sample from Analysis Agent
    df_sample = st.session_state.df_sample
    current_df = df_sample  # Use sample as current_df if no cleaned data available
    data_loaded = True
    # st.info(f"ℹ️ Using specific data sample from Analysis Agent (Sample: {len(df_sample):,} rows).")

# --- 3. Fallback: Use original uploaded data ---
elif 'current_df' in st.session_state and st.session_state.current_df is not None:
    current_df = st.session_state.current_df
    # Create a proper sample from original data using analysis_agent sample size
    sample_size = st.session_state.get('sample_size_slider', st.session_state.get('sample_size', min(2000, len(current_df))))
    df_sample = current_df.sample(n=sample_size, random_state=42) if len(current_df) > sample_size else current_df.copy()
    # Update session state to maintain consistency
    st.session_state.df_sample = df_sample
    data_loaded = True
    # st.info(f"ℹ️ Using original uploaded data (Sample: {len(df_sample):,} rows, Full: {len(current_df):,} rows).")

# --- 4. Final check: If still no data, show informative content ---
if current_df is None or df_sample is None:
    st.warning("📊 No data found from the main Analysis Agent page or previous ML page upload.")
    
    # --- MODERN NAVIGATION BUTTONS (TOP OF SIDEBAR) ---
   
    # Enhanced centered navigation button
    st.sidebar.markdown('<div class="nav-buttons-container">', unsafe_allow_html=True)
    if st.sidebar.button("🏠 Home", use_container_width=True, help="Return to main Analysis Agent"):
        st.switch_page("analysis_agent.py")
    st.sidebar.markdown('</div>', unsafe_allow_html=True)
    
    st.sidebar.markdown("---") # Visual separator
    # --- END MODERN NAVIGATION BUTTONS ---
    
    # --- MODERN CARDS & INFOGRAPHICS FOR DASHBOARD ---
    st.markdown("""
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="background: linear-gradient(90deg, #667eea, #764ba2);
                       -webkit-background-clip: text;
                       -webkit-text-fill-color: transparent;
                       font-weight: 800;
                       font-size: 2.8rem;
                       margin-bottom: 10px;">
                📊 InsightNav AI - Interactive Dashboard
            </h1>
            <p style="color: #666; font-size: 1.2rem; max-width: 800px; margin: 0 auto;">
                Create stunning interactive visualizations and gain insights from your data.
                Connect to data from the Analysis Agent or upload directly.
            </p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # --- Dashboard Info Cards ---
    # --- ENHANCED Dashboard Info Cards ---
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
            <div style="background: linear-gradient(135deg, #3a7bd5 0%, #00d2ff 100%);
                        color: white;
                        border-radius: 16px;
                        padding: 25px;
                        text-align: center;
                        box-shadow: 0 10px 20px rgba(0, 0, 0, 0.15);
                        height: 100%;
                        transition: transform 0.3s ease, box-shadow 0.3s ease;">
                <div style="font-size: 3rem; margin-bottom: 15px; filter: drop-shadow(0 4px 6px rgba(0,0,0,0.2));">📈</div>
                <h3 style="margin-bottom: 15px; color: white; font-size: 1.5rem; font-weight: 700;">Visualize Your Data</h3>
                <ul style="text-align: left; color: rgba(255, 255, 255, 0.95); font-size: 1rem; padding-left: 20px;">
                    <li>Interactive charts & dashboards</li>
                    <li>10+ visualization types</li>
                    <li>Real-time data exploration</li>
                    <li>Customizable chart parameters</li>
                </ul>
                <div style="margin-top: 20px; font-style: italic; color: rgba(255, 255, 255, 0.85);">
                    See patterns emerge from complex datasets
                </div>
            </div>
        """, unsafe_allow_html=True)
        
    with col2:
        st.markdown("""
            <div style="background: linear-gradient(135deg, #ff6b6b 0%, #ffa36c 100%);
                        color: white;
                        border-radius: 16px;
                        padding: 25px;
                        text-align: center;
                        box-shadow: 0 10px 20px rgba(0, 0, 0, 0.15);
                        height: 100%;
                        transition: transform 0.3s ease, box-shadow 0.3s ease;">
                <div style="font-size: 3rem; margin-bottom: 15px; filter: drop-shadow(0 4px 6px rgba(0,0,0,0.2));">🔍</div>
                <h3 style="margin-bottom: 15px; color: white; font-size: 1.5rem; font-weight: 700;">Discover Insights</h3>
                <ul style="text-align: left; color: rgba(255, 255, 255, 0.95); font-size: 1rem; padding-left: 20px;">
                    <li>AI-powered data analysis</li>
                    <li>Drill-down capabilities</li>
                    <li>Correlation detection</li>
                    <li>Automated pattern recognition</li>
                </ul>
                <div style="margin-top: 20px; font-style: italic; color: rgba(255, 255, 255, 0.85);">
                    Unlock hidden opportunities in your data
                </div>
            </div>
        """, unsafe_allow_html=True)
        
    with col3:
        st.markdown("""
            <div style="background: linear-gradient(135deg, #1d976c 0%, #93f9b9 100%);
                        color: white;
                        border-radius: 16px;
                        padding: 25px;
                        text-align: center;
                        box-shadow: 0 10px 20px rgba(0, 0, 0, 0.15);
                        height: 100%;
                        transition: transform 0.3s ease, box-shadow 0.3s ease;">
                <div style="font-size: 3rem; margin-bottom: 15px; filter: drop-shadow(0 4px 6px rgba(0,0,0,0.2));">🚀</div>
                <h3 style="margin-bottom: 15px; color: white; font-size: 1.5rem; font-weight: 700;">Share & Collaborate</h3>
                <ul style="text-align: left; color: rgba(255, 255, 255, 0.95); font-size: 1rem; padding-left: 20px;">
                    <li>Export interactive dashboards</li>
                    <li>Generate PDF reports</li>
                    <li>Share secure links</li>
                    <li>Embed in presentations</li>
                </ul>
                <div style="margin-top: 20px; font-style: italic; color: rgba(255, 255, 255, 0.85);">
                    Turn insights into action across your organization
                </div>
            </div>
        """, unsafe_allow_html=True)
    
    # Add hover effects with CSS
    st.markdown("""
        <style>
            div[data-testid="column"] > div > div:hover {
                transform: translateY(-5px) scale(1.02);
                box-shadow: 0 15px 30px rgba(0, 0, 0, 0.2) !important;
                z-index: 100;
            }
        </style>
    """, unsafe_allow_html=True)
    # --- END Dashboard Info Cards ---
    
    st.markdown("---")
    
    # --- Call to Action ---
    st.subheader("🚀 Get Started")
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.info("ℹ️ To begin creating dashboards, please upload data in the Analysis Agent page first.")
      
    # --- Final check: If still no data, stop execution ---
    if current_df is None or df_sample is None:
        #st.error("❌ Unable to load data. Please upload a file directly or ensure data is available from the Analysis Agent page.")
        st.page_link("analysis_agent.py", label="Go to Analysis Agent to Upload/Clean Data", icon="🏠")
        st.stop()

    st.stop() # Stop execution as no data is loaded yet for analysis

# --- END Final check: If still no data ---
# --- END of Critical Data Loading Section ---
# --- END DATA LOADING ---

# Page title with modern styling
st.markdown("""
    <div style="text-align: center; margin-bottom: 30px;">
        <div style="
            display: inline-flex;
            align-items: center;
            gap: 1rem;
            margin-bottom: 1rem;
        ">
            <div style="
                background: linear-gradient(135deg, #800080, #DDA0DD);
                border-radius: 16px;
                padding: 1rem;
                font-size: 2rem;
                box-shadow: 0 4px 12px rgba(128, 0, 128, 0.2);
            ">📊</div>
            <h1 style="
                background: linear-gradient(135deg, #800080, #DDA0DD, #FFA07A);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                font-weight: 700;
                font-size: 2.2rem;
                margin: 0;
                letter-spacing: -0.5px;
            ">
                Interactive Dashboard
            </h1>
        </div>
        <p style="
            color: #666;
            font-size: 1rem;
            max-width: 500px;
            margin: 0 auto 1.5rem auto;
            line-height: 1.5;
        ">
            Transform your data into compelling visual stories with drag-and-drop simplicity
        </p>
        <div style="
            display: flex;
            justify-content: center;
            gap: 1rem;
            flex-wrap: wrap;
        ">
            <div style="
                background: rgba(128, 0, 128, 0.1);
                padding: 0.5rem 1rem;
                border-radius: 20px;
                font-size: 0.85rem;
                color: #800080;
                font-weight: 600;
            ">✨ Real-time Updates</div>
            <div style="
                background: rgba(221, 160, 221, 0.1);
                padding: 0.5rem 1rem;
                border-radius: 20px;
                font-size: 0.85rem;
                color: #800080;
                font-weight: 600;
            ">🎨 Custom Styling</div>
            <div style="
                background: rgba(255, 160, 122, 0.1);
                padding: 0.5rem 1rem;
                border-radius: 20px;
                font-size: 0.85rem;
                color: #800080;
                font-weight: 600;
            ">📱 Export Ready</div>
        </div>
    </div>
""", unsafe_allow_html=True)

# ===== DASHBOARD CONFIGURATION =====

# --- MODERN NAVIGATION BUTTONS (TOP OF SIDEBAR) ---

# Enhanced centered navigation buttons
#st.sidebar.markdown('<div class="nav-buttons-container">', unsafe_allow_html=True)

col1, col2 = st.sidebar.columns(2)
with col1:
    if st.button("🏠 Home", use_container_width=True, help="Return to main Analysis Agent"):
        st.switch_page("analysis_agent.py")
with col2:
    if st.button("🧰 ML Lab", use_container_width=True, help="Advanced machine learning analysis"):
        st.switch_page("pages/ml_analysis.py")

st.sidebar.markdown('</div>', unsafe_allow_html=True)
st.sidebar.markdown("---") # Visual separator
# --- END MODERN NAVIGATION BUTTONS ---

st.sidebar.header("⚙️ Dashboard Configuration")

with st.sidebar.expander("📊 Data Overview", expanded=False):
    # --- Display Full Dataset Info ---
    # Determine the actual full dataset size
    full_dataset_size = 0
    if 'df_cleaned' in st.session_state and st.session_state.df_cleaned is not None:
        full_dataset_size = len(st.session_state.df_cleaned)
        st.write(f"**Full Dataset Rows:** {full_dataset_size:,}")
    elif 'current_df' in st.session_state and st.session_state.current_df is not None:
        full_dataset_size = len(st.session_state.current_df)
        st.write(f"**Full Dataset Rows:** {full_dataset_size:,}")
    else:
        st.write(f"**Full Dataset Rows:** {len(current_df):,}")
    
    st.write(f"**Columns:** {len(current_df.columns)}")
    # --- END Display Full Dataset Info ---

    # --- CRITICAL: Accurately Display Sample Info from analysis_agent.py ---
    # Priority 1: Use the specific sample created by analysis_agent.py (if it exists)
    if 'df_sample' in st.session_state and st.session_state.df_sample is not None:
        df_sample_from_session = st.session_state.df_sample
        st.write(f"**🎯 Sample Rows (from Analysis Agent):** {len(df_sample_from_session):,}")
        st.info("ℹ️ This sample size reflects the exact data sample chosen in the Analysis Agent page.")
    # Priority 2: Fallback to the df_sample used locally in this Dashboard page
    elif df_sample is not None:
        st.write(f"**🎯 Sample Rows (Local):** {len(df_sample):,}")
        st.info("ℹ️ This sample size reflects the data sample used locally in this Dashboard page.")
    else:
        # This case should ideally not be reached if data loading is successful
        st.write("**🎯 Sample Rows:** N/A")
        st.warning("⚠️ Sample data information not available.")
    # --- END CRITICAL: Accurately Display Sample Info ---

    # --- MODIFIED: Robust Reload Logic ---
    # Generate a unique key for the button to ensure it's stateless
    reload_button_key = f"reload_data_button_{int(time.time() * 1000) % 100000}"

    # Option 1: Manual Reload Button (Improved)
    if st.button("🔄 Reload Data", key=reload_button_key, use_container_width=True, help="Manually reload data from Analysis Agent."):
        with st.spinner("🔄 Reloading data from Analysis Agent..."):
            try:
                # --- CRITICAL: Signal for Forced Reload ---
                # Set a flag in session state to indicate the next data load should bypass cache
                st.session_state['dashboard_force_data_reload'] = True
                # --- END CRITICAL: Signal for Forced Reload ---

                # --- CRITICAL: Clear Specific Caches and Session State ---
                # 1. Clear the cache for the specific data loading function
                # This is now handled by the force_reload flag, but clearing can be a belt-and-suspenders approach
                

                # 2. Clear session state keys related to the *loaded data* in *this* Dashboard page
                # This ensures the data loading logic uses the freshest data from st.session_state
                keys_to_clear_on_reload = [
                    'current_df', 'df_sample', 'df_cleaned', 
                    'data_source_type', 'file_data_uploaded', 'data_loaded',
                    'consistent_dataset', 'sample_info', 'data_source_dashboard_upload',
                    # Add keys related to dashboard state if needed
                    'dashboard', 'configured_charts', 'dashboard_results',
                    'cluster_data', 'kmeans_model', 'X_processed_for_kmeans', 
                    'feature_columns_for_kmeans', 'elbow_data', 'elbow_optimal_k',
                    'reduced_data', 'transformed_data'
                    # Add any other keys your app uses that are specific to the loaded dataset or dashboard state
                ]
                for key in keys_to_clear_on_reload:
                    st.session_state.pop(key, None) # Use pop to avoid KeyError

                # --- CRITICAL: Re-run the Script ---
                # This forces the entire script to re-execute from the top.
                # On the next run, the 'dashboard_force_data_reload' flag will be True,
                # causing load_dashboard_data_from_session to bypass its cache.
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ Error reloading  {e}")
                import traceback
                st.code(traceback.format_exc()) # Show traceback for debugging
    # --- END MODIFIED: Robust Reload Logic ---
    # --- END Data Overview Expander ---

# Layout options
dashboard_num_cols = st.sidebar.selectbox("📐 Layout Columns", [1, 2, 3, 4], index=1)


# Add reorder mode toggle to the sidebar
st.sidebar.markdown("---")
reorder_mode = st.sidebar.checkbox("🔀 Enable Chart Reordering", 
                                   help="Enable this to reorder charts using arrow buttons")

# Store reorder mode in session state so render_dashboard can access it
st.session_state.reorder_mode = reorder_mode

# Debug: Show reorder mode status
if st.sidebar.button("🔍 Debug Reorder Mode"):
    st.sidebar.write(f"Checkbox value: {reorder_mode}")
    st.sidebar.write(f"Session state: {st.session_state.get('reorder_mode', 'Not set')}")

if reorder_mode:
    st.sidebar.markdown("💡 **Tip**: Use the arrow buttons above each chart to reorder them")


# Industry-Standard Chart Type Selection
chart_type_options = {
    # Core Statistical Charts
    "📈 Scatter Plot": "scatter",
    "📉 Line Chart": "line",
    "📊 Bar Chart": "bar",
    "쌓 Stacked Bar": "stacked_bar",
    "📶 Histogram": "histogram",
    "📦 Box Plot": "box",
    
    # Business Intelligence Charts
    "💹 Waterfall Chart": "waterfall",
    "🎯 Gauge Chart": "gauge",
    "📊 Funnel Chart": "funnel",
    "🌊 Area Chart": "area",
    "📈 Dual-Axis Chart": "dual_axis",
    "📊 Grouped Bar": "grouped_bar",
    
    # Advanced Analytics Charts
    "🌡️ Correlation Heatmap": "correlation_heatmap",
    "🎯 Bullet Chart": "bullet",
    "📊 Pareto Chart": "pareto",
    "🌳 Treemap": "treemap",
    "🌐 Sunburst Chart": "sunburst",
    "📊 Sankey Diagram": "sankey",
    
    # Distribution & Statistical Charts
    "🔔 Violin Plot": "violin",
    "📊 Strip Chart": "strip",
    "📈 Q-Q Plot": "qq_plot",
    "📊 Density Plot": "density",
    "📈 Ridge Plot": "ridge",
    
    # Time Series & Forecasting
    "📈 Time Series": "timeseries",
    "🔮 Forecast Chart": "forecast",
    "📊 Seasonal Decomposition": "seasonal",
    "📈 Moving Average": "moving_average",
    
    # Comparison & Ranking
    "🏆 Ranking Chart": "ranking",
    "⚖️ Comparison Chart": "comparison",
    "📊 Slope Chart": "slope",
    "🎯 Dot Plot": "dot_plot",
    
    # Composition Charts
    "🥧 Pie Chart": "pie",
    "🍩 Donut Chart": "donut",
    "📊 Stacked Area": "stacked_area",
    "📊 100% Stacked Bar": "percent_stacked_bar",
    
    # Geographic & Spatial
    "🗺️ Choropleth Map": "choropleth",
    "📍 Scatter Map": "scatter_map",
    "🌍 Bubble Map": "bubble_map",
    
    # Performance & KPI
    "🎯 KPI Metric": "kpi",
    "📊 Scorecard": "scorecard",
    "🚦 Traffic Light": "traffic_light",
    "📈 Sparkline": "sparkline",
    
    # Data Tables & Grids
    "📋 Data Table": "table",
    "📊 Pivot Table": "pivot_table",
    "📈 Cross-Tab": "crosstab",
}
display_name_to_type = {v: k for k, v in chart_type_options.items()}
available_chart_types_internal = list(chart_type_options.values())


# Chart Configuration Persistence System
if "chart_configurations" not in st.session_state:
    st.session_state.chart_configurations = {}  # Store individual chart configs by unique ID

if "current_chart_id" not in st.session_state:
    st.session_state.current_chart_id = None  # Track which chart is being edited

if "chart_counter" not in st.session_state:
    st.session_state.chart_counter = 0  # Counter for generating unique chart IDs

# Initialize session state for charts

if "configured_charts" not in st.session_state:
    st.session_state.configured_charts = []

if "temp_chart_config" not in st.session_state:
    st.session_state.temp_chart_config = {
        "type": None,
        "params": {},
        "index": None
    }

if "dashboard_edit_mode" not in st.session_state:
    st.session_state.dashboard_edit_mode = False
# --- END NEW FLAG ---

# --- Function to Reset Temporary Config ---
def reset_temp_config():
    st.session_state.temp_chart_config = {"type": None, "params": {}, "index": None}


def generate_chart_id():
    """Generate unique chart ID"""
    if "chart_counter" not in st.session_state:
        st.session_state.chart_counter = 0
    st.session_state.chart_counter += 1
    return f"chart_{st.session_state.chart_counter}"


def save_chart_configuration(chart_config, chart_id=None, update_dashboard=True):
    """Save chart configuration with persistence - ENHANCED for all chart types"""
    if chart_id is None:
        chart_id = generate_chart_id()
    
    # FIX: Ensure we're copying the params correctly
    params_copy = chart_config.get("params", {}).copy()
    
    # FIX: For KPI charts, ensure metrics are properly saved
    if chart_config.get("type") == "kpi":
        metrics = params_copy.get("metrics", [])
        if metrics is None:
            params_copy["metrics"] = []
        elif not isinstance(metrics, list):
            params_copy["metrics"] = [metrics]
    
    # Store the complete configuration for ALL chart types
    st.session_state.chart_configurations[chart_id] = {
        "type": chart_config.get("type"),
        "params": params_copy,
        "title": chart_config.get("title", ""),
        "description": chart_config.get("description", "")
    }
    
    # Only update dashboard charts if requested (to avoid duplicates)
    if update_dashboard:
        # Also update the dashboard's charts list
        if 'dashboard' not in st.session_state:
            st.session_state.dashboard = {"charts": [], "config": {}}
        
        # Check if this chart already exists in dashboard
        chart_index = next((i for i, chart in enumerate(st.session_state.dashboard["charts"])
                           if chart.get("id") == chart_id), -1)
        
        if chart_index >= 0:
            # Update existing chart
            st.session_state.dashboard["charts"][chart_index] = {
                "id": chart_id,
                "type": chart_config.get("type"),
                "params": params_copy,
                "title": chart_config.get("title", ""),
                "description": chart_config.get("description", "")
            }
        else:
            # Add new chart
            st.session_state.dashboard["charts"].append({
                "id": chart_id,
                "type": chart_config.get("type"),
                "params": params_copy,
                "title": chart_config.get("title", ""),
                "description": chart_config.get("description", "")
            })
    
    return chart_id

def sync_configured_charts_to_dashboard():
    """Sync configured charts to dashboard for generation"""
    if 'dashboard' not in st.session_state:
        st.session_state.dashboard = {"charts": [], "config": {}}
    
    # Clear existing dashboard charts to avoid duplicates
    st.session_state.dashboard["charts"] = []
    
    # Add all configured charts to dashboard
    for chart in st.session_state.configured_charts:
        chart_id = chart.get("id")
        if chart_id:
            st.session_state.dashboard["charts"].append({
                "id": chart_id,
                "type": chart.get("type"),
                "params": chart.get("params", {}).copy(),
                "title": chart.get("title", ""),
                "description": chart.get("description", "")
            })




# Update the move_chart functions to properly swap positions
def move_chart_up(chart_index):
    """Move a chart up in the dashboard"""
    if chart_index > 0:
        charts = st.session_state.configured_charts.copy()
        # Swap positions
        charts[chart_index], charts[chart_index-1] = charts[chart_index-1], charts[chart_index]
        st.session_state.configured_charts = charts
        
        # Also update dashboard if it exists
        if 'dashboard' in st.session_state and st.session_state.dashboard:
            dashboard_charts = st.session_state.dashboard.get('charts', []).copy()
            if chart_index < len(dashboard_charts) and chart_index-1 >= 0:
                dashboard_charts[chart_index], dashboard_charts[chart_index-1] = dashboard_charts[chart_index-1], dashboard_charts[chart_index]
                st.session_state.dashboard['charts'] = dashboard_charts
        
        st.success(f"✅ Moved chart from position {chart_index+1} to {chart_index}")
        st.rerun()

def move_chart_down(chart_index):
    """Move a chart down in the dashboard"""
    charts = st.session_state.configured_charts.copy()
    if chart_index < len(charts) - 1:
        # Swap positions
        charts[chart_index], charts[chart_index+1] = charts[chart_index+1], charts[chart_index]
        st.session_state.configured_charts = charts
        
        # Also update dashboard if it exists
        if 'dashboard' in st.session_state and st.session_state.dashboard:
            dashboard_charts = st.session_state.dashboard.get('charts', []).copy()
            if chart_index < len(dashboard_charts) and chart_index+1 < len(dashboard_charts):
                dashboard_charts[chart_index], dashboard_charts[chart_index+1] = dashboard_charts[chart_index+1], dashboard_charts[chart_index]
                st.session_state.dashboard['charts'] = dashboard_charts
        
        st.success(f"✅ Moved chart from position {chart_index+1} to {chart_index+2}")
        st.rerun()

def move_chart_left(chart_index, num_cols):
    """Move a chart left in the dashboard"""
    if chart_index % num_cols != 0:  # Not in first column
        charts = st.session_state.configured_charts.copy()
        # Swap positions
        charts[chart_index], charts[chart_index-1] = charts[chart_index-1], charts[chart_index]
        st.session_state.configured_charts = charts
        
        # Also update dashboard if it exists
        if 'dashboard' in st.session_state and st.session_state.dashboard:
            dashboard_charts = st.session_state.dashboard.get('charts', []).copy()
            if chart_index < len(dashboard_charts) and chart_index-1 >= 0:
                dashboard_charts[chart_index], dashboard_charts[chart_index-1] = dashboard_charts[chart_index-1], dashboard_charts[chart_index]
                st.session_state.dashboard['charts'] = dashboard_charts
        
        st.success(f"✅ Moved chart left from position {chart_index+1} to {chart_index}")
        st.rerun()

def move_chart_right(chart_index, num_cols):
    """Move a chart right in the dashboard"""
    charts = st.session_state.configured_charts.copy()
    if chart_index % num_cols != num_cols - 1 and chart_index < len(charts) - 1:  # Not in last column and not last chart
        # Swap positions
        charts[chart_index], charts[chart_index+1] = charts[chart_index+1], charts[chart_index]
        st.session_state.configured_charts = charts
        
        # Also update dashboard if it exists
        if 'dashboard' in st.session_state and st.session_state.dashboard:
            dashboard_charts = st.session_state.dashboard.get('charts', []).copy()
            if chart_index < len(dashboard_charts) and chart_index+1 < len(dashboard_charts):
                dashboard_charts[chart_index], dashboard_charts[chart_index+1] = dashboard_charts[chart_index+1], dashboard_charts[chart_index]
                st.session_state.dashboard['charts'] = dashboard_charts
        
        st.success(f"✅ Moved chart right from position {chart_index+1} to {chart_index+2}")
        st.rerun()

def render_chart_with_persistence(chart_config, df):
    """Render chart using saved configuration"""
    # Get the chart ID and retrieve the full configuration
    chart_id = chart_config.get("id")
    if chart_id and chart_id in st.session_state.chart_configurations:
        full_config = st.session_state.chart_configurations[chart_id]
        chart_type = full_config.get("type", "unknown")
        params = full_config.get("params", {})
    else:
        # Fallback to the passed config
        chart_type = chart_config.get("type", "unknown")
        params = chart_config.get("params", {})
    
    # Use the existing render_chart logic with the retrieved configuration
    return render_chart({"type": chart_type, "params": params}, df)


def ensure_chart_ids():
    """Ensure all charts have unique IDs - ENHANCED for all chart types"""
    if 'configured_charts' in st.session_state:
        for i, chart in enumerate(st.session_state.configured_charts):
            if 'id' not in chart:
                chart_id = generate_chart_id()
                chart['id'] = chart_id
                # Save to persistent storage for ALL chart types
                save_chart_configuration(chart, chart_id)
    
    # Also ensure dashboard charts have IDs
    if 'dashboard' in st.session_state:
        for i, chart in enumerate(st.session_state.dashboard.get("charts", [])):
            if 'id' not in chart:
                chart_id = generate_chart_id()
                chart['id'] = chart_id
                # Save to persistent storage
                save_chart_configuration(chart, chart_id)


def load_chart_configuration_for_editing(chart_id):
    """Load chart configuration for editing with all parameters"""
    if chart_id and chart_id in st.session_state.chart_configurations:
        return st.session_state.chart_configurations[chart_id]
    return None

def update_config_with_saved_values(config, saved_config):
    """Update config dictionary with saved values"""
    if saved_config and "params" in saved_config:
        for key, value in saved_config["params"].items():
            if key not in config or config[key] is None:
                config[key] = value
    return config

# Initialize chart IDs for existing charts
ensure_chart_ids()

# --- END TEMPORARY STATE ---

# --- HELPER FUNCTION FOR COLOR PALETTE SELECTION ---
def render_color_palette_selector(config_key_prefix, current_palette=None):
    """Render color palette selector for chart configuration"""
    try:
        color_sequences = [seq for seq in dir(px.colors.sequential) if not seq.startswith('_') and isinstance(getattr(px.colors.sequential, seq), list)]
        
        # Add some popular discrete color sequences as well
        discrete_sequences = [seq for seq in dir(px.colors.qualitative) if not seq.startswith('_') and isinstance(getattr(px.colors.qualitative, seq), list)]
        
        # Combine both sequential and qualitative palettes with None option first
        all_palettes = ["--None--"] + color_sequences + [f"Qualitative_{seq}" for seq in discrete_sequences[:10]]  # Limit discrete to avoid clutter
        
        default_index = 0  # Default to "--None--"
        if current_palette:
            if current_palette in all_palettes:
                default_index = all_palettes.index(current_palette)
            elif f"Qualitative_{current_palette}" in all_palettes:
                default_index = all_palettes.index(f"Qualitative_{current_palette}")
            
        selected_palette = st.selectbox(
            "🎨 Color Palette",
            options=all_palettes,
            index=default_index,
            key=f"{config_key_prefix}_color_palette",
            help="Choose color palette for this visualization. Select '--None--' to use default colors."
        )
        
        # Return None if "--None--" is selected
        return None if selected_palette == "--None--" else selected_palette
    except Exception as e:
        st.warning(f"Could not load color palettes: {e}")
        return None

def apply_color_palette_to_chart(fig, palette_name, chart_has_color_column=False):
    """Apply color palette to a plotly figure"""
    if not palette_name or palette_name == "--None--":
        return fig  # Use default colors
    
    try:
        # Handle qualitative palettes
        if palette_name.startswith("Qualitative_"):
            palette_name = palette_name.replace("Qualitative_", "")
            if hasattr(px.colors.qualitative, palette_name):
                color_palette = getattr(px.colors.qualitative, palette_name)
                if chart_has_color_column:
                    fig.update_traces(marker=dict(colorscale=None))
                    # For discrete colors, update the color discrete sequence
                    fig.update_layout(colorway=color_palette)
                else:
                    # Apply to all traces
                    for i, trace in enumerate(fig.data):
                        if i < len(color_palette):
                            fig.data[i].marker.color = color_palette[i % len(color_palette)]
        else:
            # Handle sequential palettes
            if hasattr(px.colors.sequential, palette_name):
                color_palette = getattr(px.colors.sequential, palette_name)
                if chart_has_color_column:
                    # For continuous color scales
                    fig.update_traces(marker=dict(colorscale=color_palette))
                else:
                    # For discrete colors, sample from the palette
                    sampled_colors = [color_palette[int(i * (len(color_palette) - 1) / max(1, len(fig.data) - 1))] 
                                    for i in range(len(fig.data))]
                    for i, trace in enumerate(fig.data):
                        if i < len(sampled_colors):
                            fig.data[i].marker.color = sampled_colors[i]
                            
    except Exception as e:
        st.warning(f"Could not apply color palette '{palette_name}': {e}")
    
    return fig

# --- DYNAMIC CHART CONFIGURATION UI ---
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="
    background: linear-gradient(135deg, rgba(128, 0, 128, 0.05) 0%, rgba(221, 160, 221, 0.08) 100%);
    border-radius: 12px;
    padding: 1rem;
    margin: 1rem 0;
    border: 1px solid rgba(128, 0, 128, 0.1);
">
    <h3 style="
        margin: 0 0 0.5rem 0;
        color: #800080;
        font-weight: 600;
        font-size: 1.1rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    ">
        ➕ Chart Builder
    </h3>
    <p style="
        margin: 0;
        color: #666;
        font-size: 0.85rem;
        line-height: 1.4;
    ">
        Select a chart type and configure it in the main area
    </p>
</div>
""", unsafe_allow_html=True)

# --- COLUMN IDENTIFICATION ---
# Updated to handle new features from cleaning process
numeric_cols = current_df.select_dtypes(include=['number']).columns.tolist()
categorical_cols = current_df.select_dtypes(include=['object', 'category']).columns.tolist()
datetime_cols = [col for col in current_df.columns if pd.api.types.is_datetime64_any_dtype(current_df[col])]
all_cols = current_df.columns.tolist()

if not numeric_cols:
    numeric_cols = [None]
if not categorical_cols:
    categorical_cols = [None]
if not datetime_cols:
    datetime_cols = [None]
# --- END COLUMN IDENTIFICATION ---

# --- Step 1: Select Chart Type (MODIFIED FOR EDIT) ---
# Determine the default selection for the chart type dropdown based on temp config
temp_config_type = st.session_state.temp_chart_config["type"]
default_chart_type_index = 0
if temp_config_type and temp_config_type in available_chart_types_internal:
    try:
        default_chart_type_index = available_chart_types_internal.index(temp_config_type)
    except ValueError:
        pass

# Modern chart type selection with enhanced styling
st.sidebar.markdown("**Choose Chart Type:**")
selected_chart_display_name = st.sidebar.selectbox(
    "Chart Type",
    options=list(chart_type_options.keys()),
    index=default_chart_type_index,
    key="selected_chart_type_display_key_for_edit",
    help="Select the type of visualization you want to create",
    label_visibility="collapsed"
)
# Extract the actual chart type name (remove emoji)
selected_chart_type_internal = chart_type_options[selected_chart_display_name]
st.session_state.temp_chart_config["type"] = selected_chart_type_internal
# --- END CHART TYPE SELECTION ---

# --- Step 2: Configure Selected Chart (in Main Area) ---
# Determine if the configuration UI should be shown
# Show it if:
# 1. No dashboard is currently displayed (normal flow), OR
# 2. We are in dashboard edit mode (user clicked Edit while dashboard was shown)
show_config_ui = not st.session_state.get('show_dashboard', False) or st.session_state.get('dashboard_edit_mode', False)

if show_config_ui:
    # Modern configuration header
    st.markdown(f"""
    <div style="
        background: linear-gradient(135deg, rgba(128, 0, 128, 0.03) 0%, rgba(221, 160, 221, 0.05) 100%);
        border-radius: 16px;
        padding: 1.5rem;
        margin: 1rem 0;
        border: 1px solid rgba(128, 0, 128, 0.1);
    ">
        <h3 style="
            margin: 0 0 0.5rem 0;
            color: #800080;
            font-weight: 600;
            font-size: 1.4rem;
            display: flex;
            align-items: center;
            gap: 0.5rem;
        ">
            🛠️ Configure: <span style="
                color: #FFFFFF;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                padding: 4px 12px;
                border-radius: 8px;
                font-weight: 600;
                text-shadow: 0 1px 2px rgba(0,0,0,0.3);
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                margin-left: 8px;
            ">{selected_chart_display_name}</span>
        </h3>
        <p style="
            margin: 0;
            color: #666;
            font-size: 0.9rem;
        ">
            Customize your visualization settings below
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Use the parameters from temp state if editing, otherwise default empty dict
    config = st.session_state.temp_chart_config["params"]
    # Create unique key prefix including chart index to ensure widget state updates
    editing_index = st.session_state.temp_chart_config.get("index", "new")
    config_key_prefix = f"temp_config_{selected_chart_type_internal}_{editing_index}"

    # Helper function to find index safely for selectboxes
    def find_index_safe(options_list, value):
        """Helper function to find index safely for selectboxes"""
        try:
            if value is None and "--None--" in options_list:
                return options_list.index("--None--")
            elif value is None and None in options_list:
                return options_list.index(None)
            elif value is None:
                return 0
            return options_list.index(value)
        except (ValueError, AttributeError):
            return 0  # Default to first item if not found
    # --- Step 3: Render Dynamic Configuration Form (in Main Area) ---
    # Determine if we're editing
    is_editing = st.session_state.temp_chart_config["index"] is not None
    
    # Load saved configuration if editing
    if is_editing and st.session_state.temp_chart_config.get("id"):
        saved_config = load_chart_configuration_for_editing(st.session_state.temp_chart_config["id"])
        if saved_config:
            config = update_config_with_saved_values(config, saved_config)
    
    try:
        if selected_chart_type_internal == "scatter":
            col1, col2 = st.columns(2)
            with col1:
                # Initialize with current value if exists
                current_x = config.get("x_col")
                config["x_col"] = st.selectbox(
                    "X-axis", 
                    numeric_cols, 
                    index=numeric_cols.index(current_x) if current_x in numeric_cols else 0,
                    key=f"{config_key_prefix}_x"
                )
            with col2:
                current_y = config.get("y_col")
                config["y_col"] = st.selectbox(
                    "Y-axis", 
                    numeric_cols, 
                    index=numeric_cols.index(current_y) if current_y in numeric_cols else 0,
                    key=f"{config_key_prefix}_y"
                )
            current_color = config.get("color_col")
            config["color_col"] = st.selectbox(
                "🎨 Color By", 
                [None] + categorical_cols, 
                index=([None] + categorical_cols).index(current_color) if current_color in [None] + categorical_cols else 0,
                key=f"{config_key_prefix}_color"
            )
            current_size = config.get("size_col")
            config["size_col"] = st.selectbox(
                "📏 Size By", 
                [None] + numeric_cols, 
                index=([None] + numeric_cols).index(current_size) if current_size in [None] + numeric_cols else 0,
                key=f"{config_key_prefix}_size"
            )
            current_hover = config.get("hover_data", [])
            config["hover_data"] = st.multiselect(
                "📌 Additional Hover Data", 
                all_cols, 
                default=current_hover,
                key=f"{config_key_prefix}_hover"
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "correlation_heatmap":
            current_columns = config.get("columns", [])
            default_corr_cols = current_columns if current_columns else numeric_cols[:min(10, len(numeric_cols))]
            config["columns"] = st.multiselect(
                "📈 Numerical Columns", 
                numeric_cols, 
                default=default_corr_cols,
                key=f"{config_key_prefix}_cols"
            )
            current_scale = config.get("color_scale", "RdBu")
            config["color_scale"] = st.selectbox(
                "🎨 Color Scale", 
                options=["RdBu", "Viridis", "Plasma", "Inferno", "Jet", "Hot"], 
                index=["RdBu", "Viridis", "Plasma", "Inferno", "Jet", "Hot"].index(current_scale) if current_scale in ["RdBu", "Viridis", "Plasma", "Inferno", "Jet", "Hot"] else 0,
                key=f"{config_key_prefix}_colorscale"
            )
            current_text = config.get("text_auto", True)
            config["text_auto"] = st.checkbox(
                "🔢 Show Values", 
                value=current_text, 
                key=f"{config_key_prefix}_text"
            )

        elif selected_chart_type_internal == "box":
            col1, col2 = st.columns(2)
            with col1:
                current_num = config.get("num_col")
                config["num_col"] = st.selectbox(
                    "🔢 Numerical Column", 
                    numeric_cols, 
                    index=numeric_cols.index(current_num) if current_num in numeric_cols else 0,
                    key=f"{config_key_prefix}_num"
                )
            with col2:
                current_cat = config.get("cat_col")
                config["cat_col"] = st.selectbox(
                    "🏷️ Categorical Column", 
                    categorical_cols, 
                    index=categorical_cols.index(current_cat) if current_cat in categorical_cols else 0,
                    key=f"{config_key_prefix}_cat"
                )
            current_color = config.get("color_col")
            config["color_col"] = st.selectbox(
                "🎨 Color By", 
                [None] + categorical_cols, 
                index=([None] + categorical_cols).index(current_color) if current_color in [None] + categorical_cols else 0,
                key=f"{config_key_prefix}_color_box"
            )
            current_orient = config.get("orientation", "v")
            config["orientation"] = st.radio(
                "📐 Orientation", 
                options=["v", "h"], 
                index=0 if current_orient == "v" else 1,
                key=f"{config_key_prefix}_orient", 
                horizontal=True
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "line":
            col1, col2 = st.columns(2)
            with col1:
                current_x = config.get("x_col")
                x_options = datetime_cols + categorical_cols
                config["x_col"] = st.selectbox("X-axis (Date/Category)", x_options, index=find_index_safe(x_options, current_x), key=f"{config_key_prefix}_x")
            with col2:
                current_y = config.get("y_col")
                config["y_col"] = st.selectbox("Y-axis (Primary Value)", numeric_cols, index=find_index_safe(numeric_cols, current_y), key=f"{config_key_prefix}_y_primary")

            current_color = config.get("color_col")
            color_options = [None] + categorical_cols
            config["color_col"] = st.selectbox("Group By (Color - Primary)", color_options, index=find_index_safe(color_options, current_color), key=f"{config_key_prefix}_color_primary", help="Group primary lines by a category")
            
            st.markdown("**Secondary Y-Axis (Optional)**")
            col1, col2 = st.columns(2)
            with col1:
                current_y_sec = config.get("y_col_secondary")
                y_sec_options = [None] + numeric_cols
                config["y_col_secondary"] = st.selectbox("Y-axis (Secondary Value)", y_sec_options, index=find_index_safe(y_sec_options, current_y_sec), key=f"{config_key_prefix}_y_secondary", help="Add a second metric on a different scale.")
            with col2:
                current_color_sec = config.get("color_col_secondary")
                color_sec_options = [None] + categorical_cols
                config["color_col_secondary"] = st.selectbox("Group By (Color - Secondary)", color_sec_options, index=find_index_safe(color_sec_options, current_color_sec), key=f"{config_key_prefix}_color_secondary", help="Group secondary lines by a category")

            st.markdown("**Aggregation & Time Settings**")
            col1, col2, col3 = st.columns(3)
            with col1:
                agg_options = ["mean", "sum", "count", "min", "max", "median", "std"]
                current_agg = config.get("agg_func", "mean")
                config["agg_func"] = st.selectbox("Aggregation (Primary)", agg_options, index=find_index_safe(agg_options, current_agg), key=f"{config_key_prefix}_agg_primary", help="How to aggregate primary Y values.")
            with col2:
                current_agg_sec = config.get("agg_func_secondary", "mean")
                config["agg_func_secondary"] = st.selectbox("Aggregation (Secondary)", agg_options, index=find_index_safe(agg_options, current_agg_sec), key=f"{config_key_prefix}_agg_secondary", help="How to aggregate secondary Y values.")
            with col3:
                time_agg_options = [None, 'D', 'W', 'M', 'Q', 'Y']
                current_time_agg = config.get("time_agg")
                config["time_agg"] = st.selectbox(
                    "Time Aggregation",
                    time_agg_options,
                    index=find_index_safe(time_agg_options, current_time_agg),
                    format_func=lambda x: {'D': 'Daily', 'W': 'Weekly', 'M': 'Monthly', 'Q': 'Quarterly', 'Y': 'Yearly'}.get(x, "None"),
                    key=f"{config_key_prefix}_time_agg",
                    help="Aggregate time series data. Only applies if X-axis is a datetime column."
                )

            config["show_markers"] = st.checkbox("Show Markers", value=config.get("show_markers", False), key=f"{config_key_prefix}_show_markers")
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))
            # --- END ENHANCED LINE CHART CONFIGURATION ---

        elif selected_chart_type_internal == "bar":
            col1, col2 = st.columns(2)
            with col1:
                current_x = config.get("x_col")
                x_options = categorical_cols
                config["x_col"] = st.selectbox(
                    "🏷️ X-axis (Category)",
                    x_options,
                    index=x_options.index(current_x) if current_x in x_options else 0,
                    key=f"{config_key_prefix}_x_bar"
                )
            with col2:
                current_y = config.get("y_col")
                y_options = [None] + numeric_cols
                y_index = 0
                if current_y is None:
                    y_index = 0 # Index of None in [None] + numeric_cols
                elif current_y in numeric_cols:
                    y_index = y_options.index(current_y)
                # Default to None if y_col not set or invalid
                config["y_col"] = st.selectbox(
                    "📊 Y-axis (Value)",
                    y_options,
                    index=y_index,
                    key=f"{config_key_prefix}_y_bar"
                )

            current_color = config.get("color_col")
            color_options = [None] + categorical_cols
            config["color_col"] = st.selectbox(
                "🎨 Color By",
                color_options,
                index=color_options.index(current_color) if current_color in color_options else 0,
                key=f"{config_key_prefix}_color_bar"
            )

            # Determine default aggregation based on Y value
            default_agg_index = 0 # Default to 'count'
            if config.get("y_col") is not None:
                 default_agg_index = 2 # Default to 'mean' if Y is selected
            current_agg = config.get("agg_func", "count" if config.get("y_col") is None else "mean")
            agg_options = ["count", "sum", "mean", "min", "max"]
            try:
                agg_index = agg_options.index(current_agg)
            except ValueError:
                agg_index = default_agg_index
            config["agg_func"] = st.selectbox(
                "🧮 Aggregation",
                agg_options,
                index=agg_index,
                key=f"{config_key_prefix}_agg_bar"
            )

            current_mode = config.get("barmode", "group")
            mode_options = ["group", "stack", "overlay", "relative"]
            try:
                mode_index = mode_options.index(current_mode)
            except ValueError:
                mode_index = 0 # Default to 'group'
            config["barmode"] = st.radio(
                "쌓 Bar Mode",
                options=mode_options,
                index=mode_index,
                key=f"{config_key_prefix}_barmode_bar",
                horizontal=True
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "stacked_bar": # Assuming this is handled like bar with forced stack
            # This could be merged with 'bar' with a barmode default of 'stack'
            # For now, treat as distinct
            col1, col2 = st.columns(2)
            with col1:
                current_x = config.get("x_col")
                x_options = categorical_cols
                config["x_col"] = st.selectbox(
                    "🏷️ X-axis (Category)",
                    x_options,
                    index=x_options.index(current_x) if current_x in x_options else 0,
                    key=f"{config_key_prefix}_x_stacked" # Unique key
                )
            with col2:
                current_y = config.get("y_col")
                y_options = numeric_cols
                config["y_col"] = st.selectbox(
                    "📊 Y-axis (Value)",
                    y_options,
                    index=y_options.index(current_y) if current_y in y_options else 0,
                    key=f"{config_key_prefix}_y_stacked" # Unique key
                )

            current_color = config.get("color_col")
            color_options = categorical_cols
            config["color_col"] = st.selectbox(
                "🎨 Color By (Stack Groups)",
                color_options,
                index=color_options.index(current_color) if current_color in color_options else 0,
                key=f"{config_key_prefix}_color_stacked" # Unique key
            )

            current_agg = config.get("agg_func", "mean")
            agg_options = ["mean", "sum", "count", "min", "max"]
            try:
                agg_index = agg_options.index(current_agg)
            except ValueError:
                agg_index = 0 # Default to 'mean'
            config["agg_func"] = st.selectbox(
                "🧮 Aggregation",
                agg_options,
                index=agg_index,
                key=f"{config_key_prefix}_agg_stacked" # Unique key
            )
            # Implicitly stacked, so no barmode selector needed here if it's always stack
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "histogram":
            st.markdown("**Core Configuration**")
            col1, col2 = st.columns(2)
            with col1:
                current_col = config.get("col")
                config["col"] = st.selectbox(
                    "🔢 Numerical Column",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, current_col),
                    key=f"{config_key_prefix}_num_hist"
                )
            with col2:
                current_bins = config.get("bins", 20)
                config["bins"] = st.slider(
                    "📏 Number of Bins", 5, 100, value=current_bins,
                    key=f"{config_key_prefix}_bins_hist"
                )

            st.markdown("**Grouping & Faceting**")
            col1, col2, col3 = st.columns(3)
            with col1:
                current_color = config.get("color_col")
                config["color_col"] = st.selectbox(
                    "🎨 Color By", [None] + categorical_cols,
                    index=find_index_safe([None] + categorical_cols, current_color),
                    key=f"{config_key_prefix}_color_hist"
                )
            with col2:
                current_facet_row = config.get("facet_row")
                config["facet_row"] = st.selectbox(
                    "↕️ Facet by Row", [None] + categorical_cols,
                    index=find_index_safe([None] + categorical_cols, current_facet_row),
                    key=f"{config_key_prefix}_facet_row_hist",
                    help="Create subplots for each category in a vertical layout."
                )
            with col3:
                current_facet_col = config.get("facet_col")
                config["facet_col"] = st.selectbox(
                    "↔️ Facet by Column", [None] + categorical_cols,
                    index=find_index_safe([None] + categorical_cols, current_facet_col),
                    key=f"{config_key_prefix}_facet_col_hist",
                    help="Create subplots for each category in a horizontal layout."
                )

            st.markdown("**Advanced Options**")
            col1, col2, col3 = st.columns(3)
            with col1:
                current_marginal = config.get("marginal")
                config["marginal"] = st.selectbox(
                    "📉 Marginal Plot", [None, "rug", "box", "violin"],
                    index=find_index_safe([None, "rug", "box", "violin"], current_marginal),
                    key=f"{config_key_prefix}_marginal_hist"
                )
            with col2:
                current_histnorm = config.get("histnorm")
                config["histnorm"] = st.selectbox(
                    "🧮 Normalization", [None, "percent", "probability", "density", "probability density"],
                    index=find_index_safe([None, "percent", "probability", "density", "probability density"], current_histnorm),
                    key=f"{config_key_prefix}_histnorm_hist",
                    help="Normalize bar heights (e.g., show as percentages)."
                )
            with col3:
                current_barmode = config.get("barmode", "overlay")
                config["barmode"] = st.selectbox(
                    "📊 Bar Mode", ["overlay", "stack", "group", "relative"],
                    index=find_index_safe(["overlay", "stack", "group", "relative"], current_barmode),
                    key=f"{config_key_prefix}_barmode_hist",
                    help="How to display bars when 'Color By' is used. 'Overlay' is recommended for distributions."
                )

            col1, col2, col3 = st.columns(3)
            with col1:
                config["cumulative"] = st.checkbox("📈 Cumulative", value=config.get("cumulative", False), key=f"{config_key_prefix}_cumulative_hist")
            with col2:
                config["log_x"] = st.checkbox("Log X-Axis", value=config.get("log_x", False), key=f"{config_key_prefix}_logx_hist")
            with col3:
                config["log_y"] = st.checkbox("Log Y-Axis", value=config.get("log_y", False), key=f"{config_key_prefix}_logy_hist")

            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "pie":
            st.markdown("**Configure Pie Chart Distribution(s)**")
                    
            # Option 1: Single Column Distribution (Traditional)
            # Option 2: Multiple Columns Distribution (New)
            chart_mode = st.radio(
                "Chart Mode",
                options=["Single Column Distribution", "Multi-Column Comparison"],
                index=0,
                key=f"{config_key_prefix}_mode",
                help="Choose to show distribution of one column or compare distributions across multiple columns."
            )

            if chart_mode == "Single Column Distribution":
                # --- Traditional Single Column Pie ---
                col1, col2 = st.columns(2)
                with col1:
                    config["col"] = st.selectbox(
                        "_slices",
                        categorical_cols,
                        key=f"{config_key_prefix}_cat_slices",
                        help="Select the categorical column for the pie slices."
                    )
                with col2:
                    config["value_col"] = st.selectbox(
                        "Values",
                        [None] + numeric_cols,
                        index=0,
                        key=f"{config_key_prefix}_cat_values",
                        help="Numeric column for slice sizes. If None, counts occurrences."
                    )

                grouping_col_options = [None] + [col for col in categorical_cols if col != config.get("col")]
                config["grouping_col"] = st.selectbox(
                    "Group By (Small Multiples)",
                    grouping_col_options,
                    key=f"{config_key_prefix}_grouping",
                    help="Create separate pie charts for each category in this column."
                )
                        
                # --- Single Column Specific Options ---
                col1, col2, col3 = st.columns(3)
                with col1:
                    config["max_categories"] = st.number_input(
                        "Max Slices",
                        min_value=2, max_value=50, value=8,
                        key=f"{config_key_prefix}_max",
                        help="Combine smallest categories into 'Other' if total categories exceed this."
                    )
                with col2:
                    config["hole"] = st.slider(
                        "Donut Hole",
                        0.0, 0.8, 0.0,
                        key=f"{config_key_prefix}_hole",
                        help="Size of the center hole (0 = pie, >0 = donut)."
                    )
                with col3:
                    config["sort"] = st.selectbox(
                        "Sort Slices",
                        options=["Label", "Value", "Initial"],
                        index=1,
                        key=f"{config_key_prefix}_sort",
                        help="How to order the slices."
                    )
                config["show_legend"] = st.checkbox(
                    "Show Legend",
                    value=True,
                    key=f"{config_key_prefix}_show_legend"
                )
                
                # Color palette selection for single pie chart
                config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))
                
                # Store mode in config
                config["mode"] = "single"
                # --- End Single Column Specific Options ---

            else: # chart_mode == "Multi-Column Comparison"
                # --- Multi-Column Distribution Pie ---
                config["columns_to_compare"] = st.multiselect(
                    "Columns to Compare Distributions",
                    categorical_cols,
                    default=categorical_cols[:2] if len(categorical_cols) >= 2 else categorical_cols,
                    key=f"{config_key_prefix}_cols_multi",
                    help="Select 2 or more categorical columns to compare their value distributions."
                )
                        
                config["value_col_multi"] = st.selectbox(
                    "Values (for comparison)",
                    [None] + numeric_cols,
                    index=0,
                    key=f"{config_key_prefix}_value_col_multi",
                    help="Numeric column whose values will be summed/averaged for comparison across selected columns. If None, counts occurrences."
                )
                        
                # Aggregation method for multi-column mode
                config["agg_func_multi"] = st.selectbox(
                    "Aggregation Method",
                    options=["sum", "mean", "count"],
                    index=0 if config.get("value_col_multi") else 2, # Default to count if no value col
                    key=f"{config_key_prefix}_agg_multi",
                    help="How to aggregate the 'Values' column for each category in the selected columns."
                )
                        
                # Visualization type for multi-column
                config["multi_viz_type"] = st.selectbox(
                    "Visualization Type",
                    options=["Facetted Subplots", "Side-by-Side Bars", "Radar Chart (if supported)"],
                    index=0,
                    key=f"{config_key_prefix}_multi_viz_type",
                    help="How to display the comparison. Facetted Subplots creates a pie chart for each selected column."
                )
                        
                # Common styling options for multi-mode (could be expanded)
                col1, col2 = st.columns(2)
                with col1:
                    config["hole_multi"] = st.slider(
                        "Donut Hole (Multi)",
                        0.0, 0.8, 0.0,
                        key=f"{config_key_prefix}_hole_multi",
                        help="Size of the center hole for multi-column pies."
                    )
                with col2:
                    config["max_categories_multi"] = st.number_input(
                        "Max Categories (Multi)",
                        min_value=2, max_value=20, value=5,
                        key=f"{config_key_prefix}_max_multi",
                        help="Limit categories per subplot/column for clarity."
                    )
                        
                config["show_legend_multi"] = st.checkbox(
                    "Show Legend (Multi)",
                    value=True,
                    key=f"{config_key_prefix}_show_legend_multi"
                )
                
                # Color palette selection for multi-column pie chart
                config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))
                
                        # Store mode in config
                config["mode"] = "multi"
                

                        # --- End Multi-Column Distribution Pie ---
                # --- END ENHANCED MULTI-COLUMN PIE CHART CONFIGURATION ---

        elif selected_chart_type_internal == "kpi":
            st.markdown("**Configure Advanced KPI Dashboard**")
            
            # Load saved configuration if editing
            saved_config = {}
            if is_editing and st.session_state.temp_chart_config.get("id"):
                saved_config = load_chart_configuration_for_editing(st.session_state.temp_chart_config["id"])
                if saved_config:
                    config = update_config_with_saved_values(config, saved_config)
            
            # Single metrics configuration section
            current_metrics = config.get("metrics", [])
            
            # FIX: Ensure we have a proper default (empty list instead of None)
            if current_metrics is None:
                current_metrics = []
            
            # For new charts, start with empty selection
            default_metrics = current_metrics if is_editing else []
            
            # FIX: Use a unique key for the multiselect
            selected_metrics = st.multiselect(
                "📊 Select Metrics",
                numeric_cols,
                default=default_metrics,
                key=f"{config_key_prefix}_metrics_unique"
            )
            
            # FIX: Ensure metrics are always saved as a list, even if empty
            config["metrics"] = selected_metrics
            
            # Only show aggregation options if metrics are selected
            if config["metrics"]:
                agg_options = ["mean", "sum", "count", "min", "max", "median", "std", "var"]
                
                # Initialize agg_funcs if not present
                if "agg_funcs" not in config:
                    config["agg_funcs"] = {}
                    
                st.markdown("**Aggregation Methods**")
                for i, metric in enumerate(config["metrics"]):
                    current_agg = config["agg_funcs"].get(metric, "mean")
                    selected_agg = st.selectbox(
                        f"Aggregation for {metric}",
                        agg_options,
                        index=agg_options.index(current_agg) if current_agg in agg_options else 0,
                        key=f"{config_key_prefix}_agg_{metric}_{i}"
                    )
                    config["agg_funcs"][metric] = selected_agg
            
            
            # Comparison options
            current_show_comparison = config.get("show_comparison", False)
            config["show_comparison"] = st.checkbox(
                "📈 Show Comparison Metrics",
                value=current_show_comparison,
                key=f"{config_key_prefix}_show_comp"
            )
            
            if config["show_comparison"]:
                current_comparison_type = config.get("comparison_type", "period")
                config["comparison_type"] = st.selectbox(
                    "🔍 Comparison Type",
                    ["period", "target", "benchmark"],
                    index=["period", "target", "benchmark"].index(current_comparison_type) 
                    if current_comparison_type in ["period", "target", "benchmark"] else 0,
                    key=f"{config_key_prefix}_comp_type"
                )
                
                if config["comparison_type"] == "target":
                    # Initialize target_values if not present
                    if "target_values" not in config:
                        config["target_values"] = {}
                        
                    for metric in config["metrics"]:
                        current_target = config["target_values"].get(metric, 0.0)
                        config["target_values"][metric] = st.number_input(
                            f"🎯 Target for {metric}",
                            value=current_target,
                            key=f"{config_key_prefix}_target_{metric}"
                        )
            
            # Layout options
            current_layout_style = config.get("layout_style", "--None--")
            layout_options = ["--None--", "grid", "cards", "minimal", "gauges"]  # Add this line
            config["layout_style"] = st.selectbox(
                "🎨 Layout Style",
                layout_options,
                index=find_index_safe(layout_options, current_layout_style),
                key=f"{config_key_prefix}_layout"
            )

            # Return None if "--None--" is selected
            if config["layout_style"] == "--None--":
                config["layout_style"] = None
            
            # Visual enhancements
            current_show_sparklines = config.get("show_sparklines", False)
            config["show_sparklines"] = st.checkbox(
                "📈 Show Sparklines",
                value=current_show_sparklines,
                key=f"{config_key_prefix}_sparklines"
            )
            
            current_color_scheme = config.get("color_scheme", "--None--")
            color_scheme_options = ["--None--", "corporate", "vibrant", "pastel", "mono"]
            config["color_scheme"] = st.selectbox(
                "🌈 Color Scheme",
                color_scheme_options,
                index=find_index_safe(color_scheme_options, current_color_scheme),
                key=f"{config_key_prefix}_colors"
            )

            # Return None if "--None--" is selected
            if config["color_scheme"] == "--None--":
                config["color_scheme"] = None

        elif selected_chart_type_internal == "table":
            current_columns = config.get("columns", [])
            # Default to first few columns if none selected
            default_table_cols = current_columns if current_columns else all_cols[:min(8, len(all_cols))]
            config["columns"] = st.multiselect(
                "📋 Columns",
                all_cols,
                default=default_table_cols,
                key=f"{config_key_prefix}_cols_table"
            )

            current_max_rows = config.get("max_rows", 10)
            config["max_rows"] = st.slider(
                "📏 Max Rows",
                5, 100,
                value=current_max_rows,
                key=f"{config_key_prefix}_rows_table"
            )

        # === INDUSTRY-STANDARD CHART CONFIGURATIONS ===
        
        elif selected_chart_type_internal == "waterfall":
            st.markdown("**Configure Waterfall Chart**")
            col1, col2 = st.columns(2)
            with col1:
                current_category = config.get("category_col")
                config["category_col"] = st.selectbox(
                    "🏷️ Category Column",
                    categorical_cols,
                    index=find_index_safe(categorical_cols, current_category),
                    key=f"{config_key_prefix}_cat_waterfall"
                )
            with col2:
                current_value = config.get("value_col")
                config["value_col"] = st.selectbox(
                    "📊 Value Column",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, current_value),
                    key=f"{config_key_prefix}_val_waterfall"
                )
            
            current_show_total = config.get("show_total", True)
            config["show_total"] = st.checkbox(
                "📊 Show Total",
                value=current_show_total,
                key=f"{config_key_prefix}_total_waterfall"
            )
            
            # Add color palette support for waterfall chart
            config["color_palette"] = render_color_palette_selector(
                config_key_prefix, 
                config.get("color_palette")
            )

        elif selected_chart_type_internal == "gauge":
            st.markdown("**Configure Gauge Chart**")
            config["value_col"] = st.selectbox(
                "📊 Value Column",
                numeric_cols,
                key=f"{config_key_prefix}_val_gauge"
            )
            
            col1, col2, col3 = st.columns(3)
            with col1:
                config["min_value"] = st.number_input(
                    "📉 Min Value",
                    value=0.0,
                    key=f"{config_key_prefix}_min_gauge"
                )
            with col2:
                config["max_value"] = st.number_input(
                    "📈 Max Value",
                    value=100.0,
                    key=f"{config_key_prefix}_max_gauge"
                )
            with col3:
                config["target_value"] = st.number_input(
                    "🎯 Target Value",
                    value=75.0,
                    key=f"{config_key_prefix}_target_gauge"
                )
            
            config["gauge_type"] = st.selectbox(
                "🎨 Gauge Type",
                ["Semi-Circle", "Full Circle", "Linear"],
                key=f"{config_key_prefix}_type_gauge"
            )

        elif selected_chart_type_internal == "funnel":
            st.markdown("**Configure Funnel Chart**")
            config["stage_col"] = st.selectbox(
                "🏷️ Stage Column",
                categorical_cols,
                key=f"{config_key_prefix}_stage_funnel"
            )
            config["value_col"] = st.selectbox(
                "📊 Value Column",
                numeric_cols,
                key=f"{config_key_prefix}_val_funnel"
            )
            
            config["show_percentages"] = st.checkbox(
                "📊 Show Percentages",
                value=True,
                key=f"{config_key_prefix}_pct_funnel"
            )
            
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "area":
            st.markdown("**Configure Area Chart**")
            col1, col2 = st.columns(2)
            with col1:
                config["x_col"] = st.selectbox(
                    "📊 X-axis",
                    all_cols,
                    key=f"{config_key_prefix}_x_area"
                )
            with col2:
                config["y_col"] = st.selectbox(
                    "📈 Y-axis",
                    numeric_cols,
                    key=f"{config_key_prefix}_y_area"
                )
            
            config["color_col"] = st.selectbox(
                "🎨 Color By",
                [None] + categorical_cols,
                key=f"{config_key_prefix}_color_area"
            )
            
            config["fill_mode"] = st.selectbox(
                "🎨 Fill Mode",
                ["tonexty", "tozeroy", "tonext"],
                index=1,
                key=f"{config_key_prefix}_fill_area"
            )
            
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "treemap":
            st.markdown("**Configure Treemap**")
            current_labels = config.get("labels_col")
            config["labels_col"] = st.selectbox(
                "🏷️ Labels Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, current_labels),
                key=f"{config_key_prefix}_labels_treemap"
            )
            current_values = config.get("values_col")
            config["values_col"] = st.selectbox(
                "📊 Values Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, current_values),
                key=f"{config_key_prefix}_values_treemap"
            )
            
            current_parent = config.get("parent_col")
            parent_options = [None] + categorical_cols
            config["parent_col"] = st.selectbox(
                "🌳 Parent Column (Optional)",
                parent_options,
                index=find_index_safe(parent_options, current_parent),
                key=f"{config_key_prefix}_parent_treemap"
            )
            
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "violin":
            st.markdown("**Configure Violin Plot**")
            col1, col2 = st.columns(2)
            with col1:
                config["y_col"] = st.selectbox(
                    "📊 Numerical Column",
                    numeric_cols,
                    key=f"{config_key_prefix}_y_violin"
                )
            with col2:
                config["x_col"] = st.selectbox(
                    "🏷️ Category Column",
                    categorical_cols,
                    key=f"{config_key_prefix}_x_violin"
                )
            
            config["show_box"] = st.checkbox(
                "📦 Show Box Plot",
                value=True,
                key=f"{config_key_prefix}_box_violin"
            )
            
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "bullet":
            st.markdown("**Configure Bullet Chart**")
            config["value_col"] = st.selectbox(
                "📊 Actual Value",
                numeric_cols,
                key=f"{config_key_prefix}_actual_bullet"
            )
            config["target_col"] = st.selectbox(
                "🎯 Target Value",
                numeric_cols,
                key=f"{config_key_prefix}_target_bullet"
            )
            
            config["category_col"] = st.selectbox(
                "🏷️ Category Column (Optional)",
                [None] + categorical_cols,
                key=f"{config_key_prefix}_cat_bullet"
            )
            
            config["orientation"] = st.radio(
                "📐 Orientation",
                ["Horizontal", "Vertical"],
                key=f"{config_key_prefix}_orient_bullet"
            )
        elif selected_chart_type_internal == "grouped_bar":
            st.markdown("**Configure Grouped Bar Chart**")
            col1, col2 = st.columns(2)
            with col1:
                current_x = config.get("x_col")
                config["x_col"] = st.selectbox(
                    "🏷️ X-axis (Category)",
                    categorical_cols,
                    index=find_index_safe(categorical_cols, current_x),
                    key=f"{config_key_prefix}_x_grouped"
                )
            with col2:
                current_y = config.get("y_col")
                config["y_col"] = st.selectbox(
                    "📊 Y-axis (Value)",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, current_y),
                    key=f"{config_key_prefix}_y_grouped"
                )
            
            current_color = config.get("color_col")
            config["color_col"] = st.selectbox(
                "🎨 Group By",
                categorical_cols,
                index=find_index_safe(categorical_cols, current_color),
                key=f"{config_key_prefix}_color_grouped"
            )
            
            config["barmode"] = "group"  # Force grouped mode
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "dual_axis":
            st.markdown("**Configure Dual-Axis Chart**")
            col1, col2 = st.columns(2)
            with col1:
                current_x = config.get("x_col")
                config["x_col"] = st.selectbox(
                    "📊 X-axis",
                    all_cols,
                    index=find_index_safe(all_cols, current_x),
                    key=f"{config_key_prefix}_x_dual"
                )
            with col2:
                current_y_primary = config.get("y_col_primary")
                config["y_col_primary"] = st.selectbox(
                    "📈 Primary Y-axis",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, current_y_primary),
                    key=f"{config_key_prefix}_y1_dual"
                )
            
            current_y_secondary = config.get("y_col_secondary")
            config["y_col_secondary"] = st.selectbox(
                "📈 Secondary Y-axis",
                numeric_cols,
                index=find_index_safe(numeric_cols, current_y_secondary),
                key=f"{config_key_prefix}_y2_dual"
            )
            
            current_color_primary = config.get("color_col_primary")
            config["color_col_primary"] = st.selectbox(
                "🎨 Primary Series Color",
                [None] + categorical_cols,
                index=find_index_safe([None] + categorical_cols, current_color_primary),
                key=f"{config_key_prefix}_color1_dual"
            )
            
            current_color_secondary = config.get("color_col_secondary")
            config["color_col_secondary"] = st.selectbox(
                "🎨 Secondary Series Color",
                [None] + categorical_cols,
                index=find_index_safe([None] + categorical_cols, current_color_secondary),
                key=f"{config_key_prefix}_color2_dual"
            )

        elif selected_chart_type_internal == "pareto":
            st.markdown("**Configure Pareto Chart**")
            current_category = config.get("category_col")
            config["category_col"] = st.selectbox(
                "🏷️ Category Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, current_category),
                key=f"{config_key_prefix}_cat_pareto"
            )
            
            current_value = config.get("value_col")
            config["value_col"] = st.selectbox(
                "📊 Value Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, current_value),
                key=f"{config_key_prefix}_val_pareto"
            )
            
            current_cumulative = config.get("cumulative_percentage", True)
            config["cumulative_percentage"] = st.checkbox(
                "📈 Show Cumulative Percentage",
                value=current_cumulative,
                key=f"{config_key_prefix}_cumulative_pareto"
            )
            
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette")) 

        elif selected_chart_type_internal == "sunburst":
            st.markdown("**Configure Sunburst Chart**")
            
            # Required parameters
            current_labels = config.get("labels_col")
            config["labels_col"] = st.selectbox(
                "🏷️ Labels Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, current_labels),
                key=f"{config_key_prefix}_labels_sunburst"
            )
            
            current_values = config.get("values_col")
            config["values_col"] = st.selectbox(
                "📊 Values Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, current_values),
                key=f"{config_key_prefix}_values_sunburst"
            )
            
            # Optional parent column for hierarchy
            current_parent = config.get("parents_col")
            parent_options = [None] + [col for col in categorical_cols if col != config.get("labels_col")]
            config["parents_col"] = st.selectbox(
                "🌳 Parent Column (Optional - for hierarchy)",
                parent_options,
                index=find_index_safe(parent_options, current_parent),
                key=f"{config_key_prefix}_parent_sunburst"
            )
            
            # Additional options
            config["max_depth"] = st.slider(
                "📏 Max Depth",
                min_value=1,
                max_value=5,
                value=config.get("max_depth", 2),
                key=f"{config_key_prefix}_maxdepth_sunburst"
            )
            
            config["branchvalues"] = st.radio(
                "💰 Value Calculation",
                options=["total", "remainder"],
                index=0 if config.get("branchvalues", "total") == "total" else 1,
                key=f"{config_key_prefix}_branchvalues_sunburst"
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))
        elif selected_chart_type_internal == "sankey":
            st.markdown("**Configure Sankey Diagram**")
            
            st.info("💡 Sankey diagrams require source, target, and value columns to show flow relationships")
            
            # Required parameters
            current_source = config.get("source_col")
            config["source_col"] = st.selectbox(
                "🔗 Source Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, current_source),
                key=f"{config_key_prefix}_source_sankey"
            )
            
            current_target = config.get("target_col")
            config["target_col"] = st.selectbox(
                "🔗 Target Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, current_target),
                key=f"{config_key_prefix}_target_sankey"
            )
            
            current_value = config.get("value_col")
            config["value_col"] = st.selectbox(
                "📊 Value Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, current_value),
                key=f"{config_key_prefix}_value_sankey"
            )
            
            # Additional options
            config["node_pad"] = st.slider(
                "📏 Node Padding",
                min_value=10,
                max_value=50,
                value=config.get("node_pad", 15),
                key=f"{config_key_prefix}_nodepad_sankey"
            )
            
            config["node_thickness"] = st.slider(
                "📏 Node Thickness",
                min_value=10,
                max_value=50,
                value=config.get("node_thickness", 30),
                key=f"{config_key_prefix}_nodethick_sankey"
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))
        elif selected_chart_type_internal == "strip":
            st.markdown("**Configure Strip Chart**")
            
            col1, col2 = st.columns(2)
            with col1:
                current_x = config.get("x_col")
                config["x_col"] = st.selectbox(
                    "🏷️ X-axis (Category)",
                    categorical_cols,
                    index=find_index_safe(categorical_cols, current_x),
                    key=f"{config_key_prefix}_x_strip"
                )
            with col2:
                current_y = config.get("y_col")
                config["y_col"] = st.selectbox(
                    "📊 Y-axis (Value)",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, current_y),
                    key=f"{config_key_prefix}_y_strip"
                )
            
            current_color = config.get("color_col")
            config["color_col"] = st.selectbox(
                "🎨 Color By",
                [None] + categorical_cols,
                index=find_index_safe([None] + categorical_cols, current_color),
                key=f"{config_key_prefix}_color_strip"
            )
            
            config["jitter"] = st.slider(
                "📏 Jitter Amount",
                min_value=0.0,
                max_value=1.0,
                value=config.get("jitter", 0.3),
                step=0.1,
                key=f"{config_key_prefix}_jitter_strip"
            )
            
            config["stripmode"] = st.radio(
                "📊 Strip Mode",
                options=["overlay", "group"],
                index=0 if config.get("stripmode", "overlay") == "overlay" else 1,
                key=f"{config_key_prefix}_stripmode_strip"
            )
            
            

        elif selected_chart_type_internal == "qq_plot":
            st.markdown("**Configure Q-Q Plot**")
            
            config["data_col"] = st.selectbox(
                "📊 Data Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, config.get("data_col")),
                key=f"{config_key_prefix}_data_qq"
            )
            
            config["dist"] = st.selectbox(
                "📈 Theoretical Distribution",
                options=["norm", "uniform", "expon", "logistic"],
                index=0 if config.get("dist", "norm") == "norm" else 
                       1 if config.get("dist", "norm") == "uniform" else
                       2 if config.get("dist", "norm") == "expon" else 3,
                key=f"{config_key_prefix}_dist_qq"
            )
            
            config["line"] = st.radio(
                "📐 Reference Line",
                options=["45", "s", "r", "q", "none"],
                index=0 if config.get("line", "45") == "45" else
                       1 if config.get("line", "45") == "s" else
                       2 if config.get("line", "45") == "r" else
                       3 if config.get("line", "45") == "q" else 4,
                key=f"{config_key_prefix}_line_qq"
            )
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "density":
            st.markdown("**Configure Density Plot**")
            
            config["data_col"] = st.selectbox(
                "📊 Data Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, config.get("data_col")),
                key=f"{config_key_prefix}_data_density"
            )
            
            current_group = config.get("group_col")
            config["group_col"] = st.selectbox(
                "🏷️ Group By (Optional)",
                [None] + categorical_cols,
                index=find_index_safe([None] + categorical_cols, current_group),
                key=f"{config_key_prefix}_group_density"
            )
            
            config["bandwidth"] = st.slider(
                "📏 Bandwidth",
                min_value=0.1,
                max_value=2.0,
                value=config.get("bandwidth", 0.5),
                step=0.1,
                key=f"{config_key_prefix}_bandwidth_density"
            )
            
            config["cumulative"] = st.checkbox(
                "📈 Show Cumulative Distribution",
                value=config.get("cumulative", False),
                key=f"{config_key_prefix}_cumulative_density"
            )
            
            config["fill"] = st.checkbox(
                "🎨 Fill Area Under Curve",
                value=config.get("fill", True),
                key=f"{config_key_prefix}_fill_density"
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))
            
        elif selected_chart_type_internal == "ridge":
            st.markdown("**Configure Ridge Plot**")
            
            config["data_col"] = st.selectbox(
                "📊 Data Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, config.get("data_col")),
                key=f"{config_key_prefix}_data_ridge"
            )
            
            config["category_col"] = st.selectbox(
                "🏷️ Category Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, config.get("category_col")),
                key=f"{config_key_prefix}_category_ridge"
            )
            
            config["overlap"] = st.slider(
                "📏 Overlap",
                min_value=0.5,
                max_value=2.0,
                value=config.get("overlap", 0.8),
                step=0.1,
                key=f"{config_key_prefix}_overlap_ridge"
            )
            
            config["bandwidth"] = st.slider(
                "📏 Bandwidth",
                min_value=0.1,
                max_value=2.0,
                value=config.get("bandwidth", 0.5),
                step=0.1,
                key=f"{config_key_prefix}_bandwidth_ridge"
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))
            
        elif selected_chart_type_internal == "timeseries":
            st.markdown("**Configure Time Series Chart**")
            
            # Find datetime columns
            datetime_cols = [col for col in current_df.columns if pd.api.types.is_datetime64_any_dtype(current_df[col])]
            if not datetime_cols:
                st.warning("No datetime columns found in the dataset. Please use a different chart type.")
            else:
                config["date_col"] = st.selectbox(
                    "📅 Date/Time Column",
                    datetime_cols,
                    index=find_index_safe(datetime_cols, config.get("date_col")),
                    key=f"{config_key_prefix}_date_timeseries"
                )
                
                config["value_col"] = st.selectbox(
                    "📊 Value Column",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, config.get("value_col")),
                    key=f"{config_key_prefix}_value_timeseries"
                )
                
                current_group = config.get("group_col")
                config["group_col"] = st.selectbox(
                    "🏷️ Group By (Optional)",
                    [None] + categorical_cols,
                    index=find_index_safe([None] + categorical_cols, current_group),
                    key=f"{config_key_prefix}_group_timeseries"
                )
                
                config["agg_func"] = st.selectbox(
                    "🧮 Aggregation",
                    ["mean", "sum", "count", "min", "max", "median"],
                    index=0 if config.get("agg_func", "mean") == "mean" else
                           1 if config.get("agg_func", "mean") == "sum" else
                           2 if config.get("agg_func", "mean") == "count" else
                           3 if config.get("agg_func", "mean") == "min" else
                           4 if config.get("agg_func", "mean") == "max" else 5,
                    key=f"{config_key_prefix}_agg_timeseries"
                )
                
                config["show_confidence"] = st.checkbox(
                    "📈 Show Confidence Interval",
                    value=config.get("show_confidence", False),
                    key=f"{config_key_prefix}_confidence_timeseries"
                )
                
                config["trendline"] = st.selectbox(
                    "📈 Trendline",
                    [None, "linear", "lowess", "expanding", "rolling"],
                    index=0 if config.get("trendline") is None else
                           1 if config.get("trendline") == "linear" else
                           2 if config.get("trendline") == "lowess" else
                           3 if config.get("trendline") == "expanding" else 4,
                    key=f"{config_key_prefix}_trendline_timeseries"
                )
                
                # Color palette selection
                config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "forecast":
            st.markdown("**Configure Forecast Chart**")
            
            # Find datetime columns
            datetime_cols = [col for col in current_df.columns if pd.api.types.is_datetime64_any_dtype(current_df[col])]
            if not datetime_cols:
                st.warning("No datetime columns found in the dataset. Please use a different chart type.")
            else:
                config["date_col"] = st.selectbox(
                    "📅 Date/Time Column",
                    datetime_cols,
                    index=find_index_safe(datetime_cols, config.get("date_col")),
                    key=f"{config_key_prefix}_date_forecast"
                )
                
                config["value_col"] = st.selectbox(
                    "📊 Value Column",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, config.get("value_col")),
                    key=f"{config_key_prefix}_value_forecast"
                )
                
                config["periods"] = st.number_input(
                    "🔮 Forecast Periods",
                    min_value=1,
                    max_value=100,
                    value=config.get("periods", 12),
                    key=f"{config_key_prefix}_periods_forecast"
                )
                
                config["model_type"] = st.selectbox(
                    "🧠 Model Type",
                    ["linear", "exponential", "additive", "multiplicative"],
                    index=0 if config.get("model_type", "linear") == "linear" else
                           1 if config.get("model_type", "linear") == "exponential" else
                           2 if config.get("model_type", "linear") == "additive" else 3,
                    key=f"{config_key_prefix}_model_forecast"
                )
                
                config["show_confidence"] = st.checkbox(
                    "📈 Show Confidence Interval",
                    value=config.get("show_confidence", True),
                    key=f"{config_key_prefix}_confidence_forecast"
                )
                # Color palette selection
                config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))


        elif selected_chart_type_internal == "moving_average":
            st.markdown("**Configure Moving Average Chart**")
            
            # Find datetime columns
            datetime_cols = [col for col in current_df.columns if pd.api.types.is_datetime64_any_dtype(current_df[col])]
            if not datetime_cols:
                st.warning("No datetime columns found in the dataset. Please use a different chart type.")
            else:
                config["date_col"] = st.selectbox(
                    "📅 Date/Time Column",
                    datetime_cols,
                    index=find_index_safe(datetime_cols, config.get("date_col")),
                    key=f"{config_key_prefix}_date_ma"
                )
                
                config["value_col"] = st.selectbox(
                    "📊 Value Column",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, config.get("value_col")),
                    key=f"{config_key_prefix}_value_ma"
                )
                
                config["window"] = st.slider(
                    "📏 Window Size",
                    min_value=2,
                    max_value=50,
                    value=config.get("window", 7),
                    key=f"{config_key_prefix}_window_ma"
                )
                
                config["ma_type"] = st.selectbox(
                    "📊 MA Type",
                    ["simple", "exponential"],
                    index=0 if config.get("ma_type", "simple") == "simple" else 1,
                    key=f"{config_key_prefix}_type_ma"
                )
                
                config["show_original"] = st.checkbox(
                    "📈 Show Original Data",
                    value=config.get("show_original", True),
                    key=f"{config_key_prefix}_original_ma"
                )
                # Color palette selection
                config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))


        elif selected_chart_type_internal == "ranking":
            st.markdown("**Configure Ranking Chart**")
            
            config["category_col"] = st.selectbox(
                "🏷️ Category Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, config.get("category_col")),
                key=f"{config_key_prefix}_cat_ranking"
            )
            
            config["value_col"] = st.selectbox(
                "📊 Value Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, config.get("value_col")),
                key=f"{config_key_prefix}_value_ranking"
            )
            
            config["top_n"] = st.slider(
                "🔢 Top N Items",
                min_value=5,
                max_value=50,
                value=config.get("top_n", 10),
                key=f"{config_key_prefix}_topn_ranking"
            )
            
            config["sort_order"] = st.radio(
                "📊 Sort Order",
                options=["Descending", "Ascending"],
                index=0 if config.get("sort_order", "Descending") == "Descending" else 1,
                key=f"{config_key_prefix}_order_ranking",
                horizontal=True
            )
            
            config["orientation"] = st.radio(
                "📐 Orientation",
                options=["Horizontal", "Vertical"],
                index=0 if config.get("orientation", "Horizontal") == "Horizontal" else 1,
                key=f"{config_key_prefix}_orient_ranking",
                horizontal=True
            )
            
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette")) 

        elif selected_chart_type_internal == "seasonal":
            st.markdown("**Configure Seasonal Decomposition**")
            
            # Find datetime columns
            datetime_cols = [col for col in current_df.columns if pd.api.types.is_datetime64_any_dtype(current_df[col])]
            if not datetime_cols:
                st.warning("No datetime columns found in the dataset. Please use a different chart type.")
            else:
                config["date_col"] = st.selectbox(
                    "📅 Date/Time Column",
                    datetime_cols,
                    index=find_index_safe(datetime_cols, config.get("date_col")),
                    key=f"{config_key_prefix}_date_seasonal"
                )
                
                config["value_col"] = st.selectbox(
                    "📊 Value Column",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, config.get("value_col")),
                    key=f"{config_key_prefix}_value_seasonal"
                )
                
                config["model_type"] = st.selectbox(
                    "🧠 Decomposition Model",
                    ["additive", "multiplicative"],
                    index=0 if config.get("model_type", "additive") == "additive" else 1,
                    key=f"{config_key_prefix}_model_seasonal"
                )
                
                config["period"] = st.number_input(
                    "🔄 Seasonal Period",
                    min_value=1,
                    max_value=365,
                    value=config.get("period", 12),
                    key=f"{config_key_prefix}_period_seasonal",
                    help="Number of periods in a complete seasonal cycle (e.g., 12 for monthly data with yearly seasonality)"
                )
                # Color palette selection
                config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "comparison":
            st.markdown("**Configure Comparison Chart**")
            
            config["category_col"] = st.selectbox(
                "🏷️ Category Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, config.get("category_col")),
                key=f"{config_key_prefix}_cat_comparison"
            )
            
            config["value_cols"] = st.multiselect(
                "📊 Value Columns to Compare",
                numeric_cols,
                default=numeric_cols[:min(3, len(numeric_cols))],
                key=f"{config_key_prefix}_values_comparison"
            )
            
            config["chart_type"] = st.selectbox(
                "📈 Comparison Chart Type",
                ["Bar", "Line", "Area"],
                index=0 if config.get("chart_type", "Bar") == "Bar" else
                       1 if config.get("chart_type", "Bar") == "Line" else 2,
                key=f"{config_key_prefix}_type_comparison"
            )
            
            config["normalize"] = st.checkbox(
                "📏 Normalize Values",
                value=config.get("normalize", False),
                key=f"{config_key_prefix}_normalize_comparison"
            )
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "slope":
            st.markdown("**Configure Slope Chart**")
            
            config["category_col"] = st.selectbox(
                "🏷️ Category Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, config.get("category_col")),
                key=f"{config_key_prefix}_cat_slope"
            )
            
            config["time_col"] = st.selectbox(
                "⏰ Time Period Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, config.get("time_col")),
                key=f"{config_key_prefix}_time_slope"
            )
            
            config["value_col"] = st.selectbox(
                "📊 Value Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, config.get("value_col")),
                key=f"{config_key_prefix}_value_slope"
            )
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        elif selected_chart_type_internal == "dot_plot":
            st.markdown("**Configure Dot Plot**")
            
            config["category_col"] = st.selectbox(
                "🏷️ Category Column",
                categorical_cols,
                index=find_index_safe(categorical_cols, config.get("category_col")),
                key=f"{config_key_prefix}_cat_dot"
            )
            
            config["value_col"] = st.selectbox(
                "📊 Value Column",
                numeric_cols,
                index=find_index_safe(numeric_cols, config.get("value_col")),
                key=f"{config_key_prefix}_value_dot"
            )
            
            config["group_col"] = st.selectbox(
                "👥 Group By (Optional)",
                [None] + categorical_cols,
                index=find_index_safe([None] + categorical_cols, config.get("group_col")),
                key=f"{config_key_prefix}_group_dot"
            )
            
            config["orientation"] = st.radio(
                "📐 Orientation",
                ["Horizontal", "Vertical"],
                index=0 if config.get("orientation", "Horizontal") == "Horizontal" else 1,
                key=f"{config_key_prefix}_orient_dot",
                horizontal=True
            )
            # Color palette selection
            config["color_palette"] = render_color_palette_selector(config_key_prefix, config.get("color_palette"))

        # 1. Enhanced Traffic Light Chart Configuration
        elif selected_chart_type_internal == "traffic_light":
            st.markdown("**Configure Advanced Traffic Light Dashboard**")
            
            # KPI Selection
            config["kpi_metrics"] = st.multiselect(
                "📊 Select KPI Metrics",
                numeric_cols,
                #default=numeric_cols[:min(3, len(numeric_cols))],
                default=config.get("kpi_metrics", numeric_cols[:min(3, len(numeric_cols))]),
                key=f"{config_key_prefix}_kpi_metrics"
            )
            
            # Threshold Configuration
            st.markdown("**🚦 Threshold Configuration**")
            #for metric in config.get("kpi_metrics", []):
            for metric in config.get("kpi_metrics", []): # This correctly uses the value from the multiselect    
                st.markdown(f"**Metric: `{metric}`**")
                col1, col2, col3 = st.columns(3)
                with col1:
                    config[f"{metric}_red_threshold"] = st.number_input(
                        f"🔴 Red Threshold",
                        #value=float(current_df[metric].quantile(0.25)),
                        value=config.get(f"{metric}_red_threshold", float(current_df[metric].quantile(0.25))),
                        key=f"{config_key_prefix}_red_{metric}"
                    )
                with col2:
                    config[f"{metric}_yellow_threshold"] = st.number_input(
                        f"🟡 Yellow Threshold",
                        #value=float(current_df[metric].quantile(0.50)),
                        value=config.get(f"{metric}_yellow_threshold", float(current_df[metric].quantile(0.50))),
                        key=f"{config_key_prefix}_yellow_{metric}"
                    )
                with col3:
                    config[f"{metric}_green_threshold"] = st.number_input(
                        f"🟢 Green Threshold",
                        #value=float(current_df[metric].quantile(0.75)),
                        value=config.get(f"{metric}_green_threshold", float(current_df[metric].quantile(0.75))),
                        key=f"{config_key_prefix}_green_{metric}"
                    )
            
            # Display Options
            st.markdown("**Display Options**")
            layout_options = ["grid", "circular", "linear", "dashboard"]
            config["layout_style"] = st.selectbox(
                "🎨 Layout Style",
                #["grid", "circular", "linear", "dashboard"],
                layout_options,
                index=find_index_safe(layout_options, config.get("layout_style", "grid")),
                key=f"{config_key_prefix}_traffic_layout"
            )
            
            config["show_sparklines"] = st.checkbox(
                "📈 Show Sparklines",
                #value=True,
                value=config.get("show_sparklines", True),
                key=f"{config_key_prefix}_traffic_sparklines"
            )
            
            config["show_values"] = st.checkbox(
                "🔢 Show Values",
                #value=True,
                value=config.get("show_values", True),
                key=f"{config_key_prefix}_traffic_values"
            )

       
        # 2. Enhanced Choropleth Map Configuration
        elif selected_chart_type_internal == "choropleth":
            st.markdown("**Configure Choropleth Map**")
            
            # --- Data Columns ---
            st.markdown("#### 1. Data Columns")
            col1, col2 = st.columns(2)
            with col1:
                location_col = st.selectbox(
                    "📍 Location Column",
                    all_cols,
                    index=find_index_safe(all_cols, config.get("location_col")),
                    help="Select a column with geographic names or codes (e.g., country names, ISO-3 codes, US states, custom IDs).",
                    key=f"{config_key_prefix}_location_col"
                )
                config["location_col"] = location_col
            with col2:
                value_col = st.selectbox(
                    "📊 Value Column",
                    numeric_cols,
                    index=find_index_safe(numeric_cols, config.get("value_col")),
                    help="Select the numeric column to determine the color of the regions.",
                    key=f"{config_key_prefix}_choropleth_value"
                )
                config["value_col"] = value_col

            # --- Data Aggregation ---
            st.markdown("#### 2. Data Aggregation")
            config["enable_aggregation"] = st.checkbox(
                "Aggregate data before mapping",
                value=config.get("enable_aggregation", False),
                help="Enable this to group by the location column and aggregate the value column (e.g., sum of revenue per state).",
                key=f"{config_key_prefix}_enable_aggregation"
            )

            if config.get("enable_aggregation"):
                agg_func = st.selectbox(
                    "⚙️ Aggregation Function",
                    ["sum", "mean", "median", "count", "min", "max"],
                    index=find_index_safe(["sum", "mean", "median", "count", "min", "max"], config.get("agg_func", "sum")),
                    help="The function to apply to the value column for each location.",
                    key=f"{config_key_prefix}_agg_func"
                )
                config["agg_func"] = agg_func

            # --- Map Type and Scope ---
            st.markdown("#### 3. Map Type & Scope")
            col1, col2 = st.columns(2)
            with col1:
                map_type = st.selectbox(
                    "🗺️ Map Engine",
                    ["Standard", "Mapbox (Advanced)"],
                    index=find_index_safe(["Standard", "Mapbox (Advanced)"], config.get("map_type", "Standard")),
                    help="'Standard' is good for world/country maps. 'Mapbox' offers more detail and requires a free token.",
                    key=f"{config_key_prefix}_map_type"
                )
                config["map_type"] = map_type
            with col2:
                scope = st.selectbox(
                    "🌍 Map Scope",
                    ["world", "usa", "europe", "asia", "africa", "north america", "south america"],
                    index=find_index_safe(["world", "usa", "europe", "asia", "africa", "north america", "south america"], config.get("scope", "world")),
                    help="Set the geographic scope of the map. Only applies to 'Standard' map engine.",
                    key=f"{config_key_prefix}_scope",
                    disabled=(map_type == "Mapbox (Advanced)")
                )
                config["scope"] = scope

            # --- Mapbox Token Input ---
            if map_type == "Mapbox (Advanced)":
                mapbox_token_input = st.text_input(
                    "🔑 Mapbox Access Token",
                    value=config.get("mapbox_token_input", ""),
                    placeholder="Enter your free Mapbox token here",
                    help="Required for Mapbox engine. Get a free token from mapbox.com. You can also set this in your secrets.toml as MAPBOX_TOKEN.",
                    key=f"{config_key_prefix}_mapbox_token",
                    type="password"
                )
                config["mapbox_token_input"] = mapbox_token_input

            # --- Location Data ---
            st.markdown("#### 4. Location Data Configuration")
            use_custom_geojson = st.checkbox(
                "Use Custom GeoJSON (for cities, counties, etc.)",
                value=config.get("use_custom_geojson", False),
                key=f"{config_key_prefix}_use_geojson"
            )
            config["use_custom_geojson"] = use_custom_geojson

            if use_custom_geojson:
                geojson_input_method = st.radio(
                    "GeoJSON Source",
                    ["URL", "Upload"],
                    index=0,
                    key=f"{config_key_prefix}_geojson_source",
                    horizontal=True
                )
                if geojson_input_method == "URL":
                    geojson_url = st.text_input(
                        "🔗 GeoJSON URL",
                        value=config.get("geojson_url", ""),
                        placeholder="e.g., https://example.com/data.geojson",
                        key=f"{config_key_prefix}_geojson_url"
                    )
                    config["geojson_data"] = geojson_url
                else:
                    uploaded_geojson = st.file_uploader(
                        "📤 Upload GeoJSON File",
                        type=["geojson", "json"],
                        key=f"{config_key_prefix}_geojson_upload"
                    )
                    if uploaded_geojson:
                        try:
                            config["geojson_data"] = json.load(uploaded_geojson)
                        except Exception as e:
                            st.error(f"Error reading GeoJSON file: {e}")
                
                featureidkey = st.text_input(
                    "🔑 GeoJSON Feature ID Key",
                    value=config.get("featureidkey", "properties.name"),
                    help="The key in your GeoJSON features to match with the 'Location Column'. E.g., 'properties.name' or 'id'.",
                    key=f"{config_key_prefix}_featureidkey"
                )
                config["featureidkey"] = featureidkey
            else:
                locationmode = st.selectbox(
                    "📍 Standard Location Mode",
                    ["country names", "ISO-3", "USA-states"],
                    index=find_index_safe(["country names", "ISO-3", "USA-states"], config.get("locationmode", "country names")),
                    help="The type of location identifiers in your location column.",
                    key=f"{config_key_prefix}_locationmode"
                )
                config["locationmode"] = locationmode

            # --- Styling ---
            st.markdown("#### 5. Styling")
            color_scale = st.selectbox(
                "🎨 Color Scale",
                ["Viridis", "Plasma", "Inferno", "Magma", "Cividis", "Blues", "Reds", "Greens", "RdBu"],
                index=find_index_safe(["Viridis", "Plasma", "Inferno", "Magma", "Cividis", "Blues", "Reds", "Greens", "RdBu"], config.get("color_scale", "Viridis")),
                key=f"{config_key_prefix}_choropleth_scale"
            )
            config["color_scale"] = color_scale


        # 3. Enhanced Donut Chart Configuration  
        elif selected_chart_type_internal == "donut":
            st.markdown("**Configure Multi-Layer Donut Chart**")
            
            # Multi-level configuration
            config["levels"] = st.slider(
                "📊 Number of Levels",
                1, 3, 1,
                value=config.get("levels", 1),
                help="Select the number of hierarchical rings for the donut chart.",
                key=f"{config_key_prefix}_donut_levels"
            )
            
            for level in range(config.get("levels", 1)):
                st.markdown(f"**Level {level+1} Configuration**")
                col1, col2 = st.columns(2)
                with col1:
                    config[f"level_{level}_col"] = st.selectbox(
                        f"🏷️ Category Column (Level {level+1})",
                        categorical_cols,
                        index=find_index_safe(categorical_cols, config.get(f"level_{level}_col")),
                        key=f"{config_key_prefix}_donut_cat_{level}"
                    )
                with col2:
                    config[f"level_{level}_value"] = st.selectbox(
                        f"📊 Value Column (Level {level+1})",
                        numeric_cols,
                        index=find_index_safe(numeric_cols, config.get(f"level_{level}_value")),
                        key=f"{config_key_prefix}_donut_val_{level}"
                    )
            
            # Central KPI
            center_metric_options = [None] + numeric_cols
            config["center_metric"] = st.selectbox(
                "🎯 Central KPI Metric",
                #[None] + numeric_cols,
                center_metric_options,
                index=find_index_safe(center_metric_options, config.get("center_metric")),
                help="Select a key metric to display in the center of the donut.",
                key=f"{config_key_prefix}_donut_center"
            )
            
            # Visual Options
            config["hole_size"] = st.slider(
                "🕳️ Hole Size",
                0.1, 0.8, 0.4,
                value=config.get("hole_size", 0.4),
                key=f"{config_key_prefix}_donut_hole"
            )
            
            config["show_percentages"] = st.checkbox(
                "📊 Show Percentages",
                #value=True,
                value=config.get("show_percentages", True),
                key=f"{config_key_prefix}_donut_percentages"
            )           


        # === END INDUSTRY-STANDARD CHART CONFIGURATIONS ===

        # --- Common Option for all charts ---
        current_custom_title = config.get("custom_title", "")
        config["custom_title"] = st.text_input(
            "🏷️ Custom Chart Title",
            value=current_custom_title,
            key=f"{config_key_prefix}_custom_title"
        )
        
        # Custom Chart Description
        current_custom_description = config.get("custom_description", "")
        config["custom_description"] = st.text_area(
            "📄 Custom Chart Description",
            value=current_custom_description,
            key=f"{config_key_prefix}_custom_description",
            help="Add a detailed description for this chart to provide context and insights",
            height=100
        )

    except Exception as e:
        st.error(f"Error configuring {selected_chart_display_name}: {e}")

    # --- Step 4: Action Buttons (Add/Update/Reset/Cancel Edit) ---
    st.markdown("---")
    # Modern action buttons section
    is_editing = st.session_state.temp_chart_config["index"] is not None
    button_label = "🔄 Update Chart" if is_editing else "➕ Add Chart"
    
    st.markdown("""
    <div style="
        background: rgba(255, 255, 255, 0.5);
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1.5rem 0;
        border: 1px solid rgba(128, 0, 128, 0.1);
    ">
    """, unsafe_allow_html=True)
    
    button_cols = st.columns([1, 1, 2])

    with button_cols[0]:
        if st.button(button_label, type="primary", use_container_width=True):
            # Basic validation
            is_valid = True
            if selected_chart_type_internal == "scatter" and (not config.get("x_col") or not config.get("y_col")):
                st.warning("❌ Please select both X and Y axes for the Scatter Plot.")
                is_valid = False
            elif selected_chart_type_internal == "correlation_heatmap" and (not config.get("columns") or len(config.get("columns", [])) < 2):
                 st.warning("❌ Please select at least 2 columns for the Correlation Heatmap.")
                 is_valid = False
            # Add validation for histogram
            elif selected_chart_type_internal == "histogram":
                if not config.get("col"):
                    st.warning("❌ Histogram requires a numerical column to be selected.")
                    is_valid = False
                elif config.get("col") not in numeric_cols:
                    st.warning(f"❌ Selected column '{config.get('col')}' is not numeric. Please select a numeric column.")
                    is_valid = False
            
            # Add validation for other chart types as needed...

            if is_valid:
                chart_to_save = {
                    "type": selected_chart_type_internal,
                    "params": config.copy()
                }
                
                # Save the configuration with proper ID
                if is_editing:
                    chart_id = st.session_state.temp_chart_config.get("id")
                    if not chart_id:
                        chart_id = generate_chart_id()
                    save_chart_configuration(chart_to_save, chart_id, update_dashboard=False)
                    
                    # Update existing chart in st.session_state.configured_charts
                    index_to_update = st.session_state.temp_chart_config["index"]
                    if 0 <= index_to_update < len(st.session_state.configured_charts):
                        st.session_state.configured_charts[index_to_update] = {
                            "id": chart_id,
                            **chart_to_save
                        }
                        st.success(f"✅ Updated chart #{index_to_update + 1} ({selected_chart_display_name})")
                    else:
                        st.error("❌ Error updating chart: Invalid index.")
                else:
                    chart_id = save_chart_configuration(chart_to_save, update_dashboard=False)
                    # Add new chart
                    st.session_state.configured_charts.append({
                        "id": chart_id,
                        **chart_to_save
                    })
                    st.success(f"✅ Added {selected_chart_display_name} to dashboard")

                # --- KEY CHANGE: Reset temp config AND edit mode flag ---
                reset_temp_config()
                st.session_state.dashboard_edit_mode = False # Exit edit mode
                st.rerun() # Refresh the page

    with button_cols[1]:
        # Cancel button logic
        cancel_label = "↩️ Cancel Edit" if is_editing else "↩️ Cancel"
        if st.button(cancel_label, use_container_width=True):
            reset_temp_config()
            st.session_state.dashboard_edit_mode = False # Exit edit mode
            st.rerun()

# --- END OF CONFIGURATION UI ---
else:
    # Optional: If dashboard is shown and NOT in edit mode, maybe show a message or button
    # to go back to configuration. This is optional based on your desired UX.
    pass

# --- 5. Display List of Configured Charts ---
if st.session_state.configured_charts:
    st.sidebar.markdown("---")
    st.sidebar.subheader("📋 Configured Charts")
    for i, chart in enumerate(st.session_state.configured_charts):
        chart_display_name = display_name_to_type.get(chart['type'], chart['type'].replace('_', ' ').title())
        chart_title = chart.get('params', {}).get('custom_title', '')
        display_text = f"{i+1}. {chart_display_name}"
        if chart_title:
            display_text += f" - '{chart_title}'"

        col1, col2, col3 = st.sidebar.columns([3, 1, 1])
        with col1:
            st.markdown(display_text)
        with col2:
            if st.button("✏️", key=f"edit_chart_{i}", help="Edit"):
                chart_id = chart.get("id", f"chart_{i}")
                st.session_state.current_chart_id = chart_id
                
                # Load the chart configuration into temp config
                if chart_id in st.session_state.chart_configurations:
                    chart_config = st.session_state.chart_configurations[chart_id]
                    st.session_state.temp_chart_config = {
                        "type": chart_config.get("type"),
                        "params": chart_config.get("params", {}).copy(),
                        "title": chart_config.get("title", ""),
                        "description": chart_config.get("description", ""),
                        "index": i,
                        "id": chart_id
                    }
                else:
                    # Fallback to current method if no saved config
                    st.session_state.temp_chart_config = {
                        "type": chart['type'],
                        "params": chart['params'].copy(),
                        "index": i,
                        "id": chart_id
                    }
                
                st.session_state.dashboard_edit_mode = True
                st.rerun()
        with col3:
            if st.button("🗑️", key=f"remove_chart_{i}", help="Remove"):
                st.session_state.configured_charts.pop(i)
                if st.session_state.temp_chart_config.get("index") == i:
                    reset_temp_config()
                    st.session_state.dashboard_edit_mode = False
                elif (st.session_state.temp_chart_config.get("index") is not None and
                      st.session_state.temp_chart_config.get("index") > i):
                     st.session_state.temp_chart_config["index"] -= 1
                st.rerun()

# --- 6. Generate Dashboard Button ---
st.sidebar.markdown("---")
if st.sidebar.button("🛠️ Generate Dashboard", use_container_width=True):
    if st.session_state.configured_charts:
        # Create dashboard structure
        dashboard = {
            "charts": st.session_state.configured_charts,
            "config": {"num_columns": dashboard_num_cols},
            "insights": ""
        }
        st.session_state.dashboard = dashboard
        st.session_state.show_dashboard = True
        st.session_state.dashboard_edit_mode = False # Ensure edit mode is off when generating
        st.rerun()
    else:
        st.sidebar.warning("⚠️ No charts configured. Please add at least one chart.")

# --- 7. Clear All Button (Optional) ---
if st.session_state.configured_charts and st.sidebar.button("🗑️ Clear All Charts", use_container_width=True):
    st.session_state.configured_charts = []
    reset_temp_config()
    st.session_state.dashboard_edit_mode = False
    st.rerun()

# ===== FUNCTION DEFINITIONS =====

def parse_chart_requirements(prompt):
    """Parse user request for specific chart count and types"""
    import re
    requirements = {
        'chart_count': None,
        'specific_types': [],
        'chart_count_specified': False
    }
    
    prompt_lower = prompt.lower()
    
    # Detect specific chart count requests
    count_patterns = [
        r'(\d+)\s+charts?',
        r'(\d+)\s+visualizations?',
        r'(\d+)\s+graphs?',
        r'create\s+(\d+)',
        r'make\s+(\d+)',
        r'show\s+me\s+(\d+)'
    ]
    
    for pattern in count_patterns:
        match = re.search(pattern, prompt_lower)
        if match:
            requirements['chart_count'] = int(match.group(1))
            requirements['chart_count_specified'] = True
            break
    
    # Detect specific chart types mentioned
    chart_type_keywords = {
        'bar': ['bar chart', 'bar graph', 'column chart'],
        'line': ['line chart', 'line graph', 'trend chart'],
        'scatter': ['scatter plot', 'scatter chart', 'correlation plot'],
        'pie': ['pie chart', 'pie graph', 'donut chart'],
        'histogram': ['histogram', 'distribution chart', 'frequency chart'],
        'box': ['box plot', 'box chart', 'boxplot']
    }
    
    for chart_type, keywords in chart_type_keywords.items():
        for keyword in keywords:
            if keyword in prompt_lower:
                requirements['specific_types'].append(chart_type)
                break
    
    # Remove duplicates
    requirements['specific_types'] = list(set(requirements['specific_types']))
    
    return requirements

def analyze_user_request(prompt, df):
    """Analyze user request to provide better context for AI chart generation"""
    analysis = []
    
    # Convert prompt to lowercase for analysis
    prompt_lower = prompt.lower()
    
    # Detect request type
    if any(word in prompt_lower for word in ['trend', 'over time', 'timeline', 'change', 'growth']):
        analysis.append("- User wants to see TRENDS or TIME-BASED analysis")
        if df is not None:
            date_cols = df.select_dtypes(include=['datetime64', 'object']).columns
            for col in date_cols:
                if any(date_word in col.lower() for date_word in ['date', 'time', 'year', 'month']):
                    analysis.append(f"- Suggested time column: {col}")
    
    if any(word in prompt_lower for word in ['compare', 'comparison', 'vs', 'versus', 'difference']):
        analysis.append("- User wants to COMPARE different categories or groups")
        if df is not None:
            cat_cols = df.select_dtypes(include=['object', 'category']).columns[:3]
            analysis.append(f"- Suggested grouping columns: {', '.join(cat_cols)}")
    
    if any(word in prompt_lower for word in ['relationship', 'correlation', 'related', 'connection']):
        analysis.append("- User wants to see RELATIONSHIPS between variables")
        if df is not None:
            num_cols = df.select_dtypes(include=['number']).columns[:4]
            analysis.append(f"- Suggested numeric columns for correlation: {', '.join(num_cols)}")
    
    if any(word in prompt_lower for word in ['distribution', 'spread', 'histogram', 'frequency']):
        analysis.append("- User wants to see DATA DISTRIBUTION")
        if df is not None:
            num_cols = df.select_dtypes(include=['number']).columns[:2]
            analysis.append(f"- Suggested columns for distribution: {', '.join(num_cols)}")
    
    if any(word in prompt_lower for word in ['top', 'bottom', 'highest', 'lowest', 'best', 'worst']):
        analysis.append("- User wants to see RANKINGS or TOP/BOTTOM performers")
        if df is not None:
            num_cols = df.select_dtypes(include=['number']).columns[:2]
            cat_cols = df.select_dtypes(include=['object', 'category']).columns[:2]
            analysis.append(f"- Suggested ranking columns: {', '.join(num_cols)}")
            analysis.append(f"- Suggested category columns: {', '.join(cat_cols)}")
    
    if any(word in prompt_lower for word in ['proportion', 'percentage', 'share', 'breakdown']):
        analysis.append("- User wants to see PROPORTIONS or BREAKDOWNS")
        if df is not None:
            cat_cols = df.select_dtypes(include=['object', 'category']).columns[:2]
            analysis.append(f"- Suggested categorical columns: {', '.join(cat_cols)}")
    
    if any(word in prompt_lower for word in ['summary', 'overview', 'general', 'all']):
        analysis.append("- User wants a COMPREHENSIVE OVERVIEW")
        analysis.append("- Suggest multiple chart types for complete picture")
    
    # Detect specific column mentions
    if df is not None:
        for col in df.columns:
            if col.lower() in prompt_lower or col.replace('_', ' ').lower() in prompt_lower:
                analysis.append(f"- User specifically mentioned column: {col}")
    
    # Default analysis if nothing specific detected
    if not analysis:
        analysis.append("- General dashboard request - create diverse visualizations")
        if df is not None:
            analysis.append(f"- Available numeric columns: {', '.join(df.select_dtypes(include=['number']).columns[:3])}")
            analysis.append(f"- Available categorical columns: {', '.join(df.select_dtypes(include=['object', 'category']).columns[:3])}")
    
    return '\n'.join(analysis)

def create_intelligent_fallback_dashboard(prompt, df):
    """Create intelligent fallback dashboard based on prompt analysis"""
    if df is None:
        return None
    
    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
    datetime_cols = df.select_dtypes(include=['datetime64']).columns.tolist()
    
    # Analyze user prompt for intent
    prompt_lower = prompt.lower()
    
    # Determine chart type based on user request
    chart_type = "bar"  # default
    if any(word in prompt_lower for word in ['histogram', 'distribution', 'spread']):
        chart_type = "histogram"
    elif any(word in prompt_lower for word in ['scatter', 'correlation', 'relationship']):
        chart_type = "scatter"
    elif any(word in prompt_lower for word in ['pie', 'proportion', 'percentage', 'share']):
        chart_type = "pie"
    elif any(word in prompt_lower for word in ['line', 'trend', 'over time', 'timeline']):
        chart_type = "line"
    elif any(word in prompt_lower for word in ['box', 'boxplot', 'quartile']):
        chart_type = "box"
    
    # Extract column names mentioned in prompt
    mentioned_cols = []
    for col in df.columns:
        if col.lower() in prompt_lower or col.replace('_', ' ').lower() in prompt_lower:
            mentioned_cols.append(col)
    
    # Create chart configuration based on analysis
    chart_config = {
        "id": "chart1",
        "title": f"Analysis: {prompt[:50]}..." if len(prompt) > 50 else prompt,
        "type": chart_type,
        "params": {},
        "description": f"Visualization based on your request: {prompt[:100]}..."
    }
    
    # Configure parameters based on chart type and available data
    if chart_type == "bar":
        if mentioned_cols:
            # Use mentioned columns if available
            cat_mentioned = [col for col in mentioned_cols if col in categorical_cols]
            num_mentioned = [col for col in mentioned_cols if col in numeric_cols]
            chart_config["params"]["x_col"] = cat_mentioned[0] if cat_mentioned else (categorical_cols[0] if categorical_cols else df.columns[0])
            chart_config["params"]["y_col"] = num_mentioned[0] if num_mentioned else (numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0])
        else:
            chart_config["params"]["x_col"] = categorical_cols[0] if categorical_cols else df.columns[0]
            chart_config["params"]["y_col"] = numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0]
    
    elif chart_type == "histogram":
        target_col = mentioned_cols[0] if mentioned_cols and mentioned_cols[0] in numeric_cols else (numeric_cols[0] if numeric_cols else df.columns[0])
        chart_config["params"]["column"] = target_col
        chart_config["params"]["bins"] = 20
    
    elif chart_type == "scatter":
        if len(mentioned_cols) >= 2:
            chart_config["params"]["x_col"] = mentioned_cols[0]
            chart_config["params"]["y_col"] = mentioned_cols[1]
        else:
            chart_config["params"]["x_col"] = numeric_cols[0] if len(numeric_cols) > 0 else df.columns[0]
            chart_config["params"]["y_col"] = numeric_cols[1] if len(numeric_cols) > 1 else (numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0])
        
        if categorical_cols:
            chart_config["params"]["color_col"] = categorical_cols[0]
    
    elif chart_type == "pie":
        target_col = mentioned_cols[0] if mentioned_cols and mentioned_cols[0] in categorical_cols else (categorical_cols[0] if categorical_cols else df.columns[0])
        chart_config["params"]["column"] = target_col
    
    elif chart_type == "line":
        x_col = datetime_cols[0] if datetime_cols else (categorical_cols[0] if categorical_cols else df.columns[0])
        y_col = mentioned_cols[0] if mentioned_cols and mentioned_cols[0] in numeric_cols else (numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0])
        chart_config["params"]["x_col"] = x_col
        chart_config["params"]["y_col"] = y_col
    
    elif chart_type == "box":
        chart_config["params"]["y_col"] = numeric_cols[0] if numeric_cols else df.columns[0]
        if categorical_cols:
            chart_config["params"]["x_col"] = categorical_cols[0]
    
    return {
        "title": f"Dashboard: {prompt[:30]}..." if len(prompt) > 30 else f"Dashboard: {prompt}",
        "description": f"Intelligent analysis based on your request: {prompt}",
        "charts": [chart_config],
        "kpis": []
    }


# ==== OPTIMIZED AI DASHBOARD FUNCTION =====
@st.cache_resource
def get_cached_dashboard_llm(api_key, provider, model_name):
    """Cache LLM for dashboard generation to avoid recreation"""
    return get_validated_dashboard_llm()

def create_optimized_ai_dashboard(prompt, df, conversation_history=None):
    """Optimized dashboard generation with caching and conversation memory"""
    
    # Manage conversation history
    manage_dashboard_conversation_history(max_turns=6)
    
    try:
        # Use cached LLM
        llm = get_validated_dashboard_llm()
        
        # Build optimized context
        df_hash = hash(str(df.columns.tolist()))
        data_context = build_optimized_dashboard_context(
            df_hash, 
            df.columns.tolist(),
            len(df)
        )
        
        # Get cached data summary
        summary_context = get_dashboard_data_summary(df_hash)
        
        # Build optimized dataset information (limited to essential info)
        dataset_info = f"""Dataset Overview:
- Shape: {summary_context.get('total_rows', 0):,} rows × {data_context['column_count']} columns
- Numeric columns: {', '.join(summary_context.get('numeric_columns', []))}
- Categorical columns: {', '.join(summary_context.get('categorical_columns', []))}
- Missing data: {summary_context.get('missing_data', 0)} null values"""

        # Construct optimized system message with conversation context
        conversation_context = ""
        if conversation_history:
            recent_history = conversation_history[-4:]  # Last 2 exchanges
            conversation_context = f"\nRecent conversation:\n{recent_history}"
        
        system_message = f"""You are an expert data visualization specialist for InsightNav AI.

DATASET CONTEXT:
{dataset_info}

USER REQUEST: {prompt}
{conversation_context}

Create a dashboard with 3-6 optimized charts. Return JSON format:

    {{
      "charts": [
        {{
          "type": "chart_type",
          "params": {{ ... }}   // parameters for the chart (only include the ones that are needed and available in the dataset)
        }},
        ... // more charts
      ],
      "insights": "Your insights about the dashboard and the data, in markdown format. Explain why you chose these charts and what insights they provide. Be concise (3-5 sentences)."
    }}

    COMPREHENSIVE CHART LIBRARY - Choose the most appropriate charts based on data characteristics and user intent:

    === BASIC CHART TYPES ===
    
    - Scatter Plot (type: "scatter") 🔍 - Best for: Correlation analysis, outlier detection, pattern identification
        * x_col (required, numerical column name)
        * y_col (required, numerical column name)
        * color_col (optional, categorical column name) - adds grouping dimension
        * size_col (optional, numerical column name) - creates bubble chart
        * hover_data (optional, list of column names)
        * color_palette (optional, color scheme)
        Use when: Exploring relationships between 2+ numeric variables, detecting outliers, clustering analysis

    - Bar Chart (type: "bar") 📊 - Best for: Category comparison, performance metrics, rankings
        * x_col (required, categorical column name)
        * y_col (optional, numerical column name; if not provided, counts are used)
        * color_col (optional, categorical column name) - adds grouping
        * agg_func (optional, one of ["count", "sum", "mean", "min", "max"], default "count")
        * barmode (optional, "group" or "stack")
        * color_palette (optional, color scheme)
        Use when: Comparing values across categories, showing rankings, analyzing performance by groups

    - Stacked Bar Chart (type: "stacked_bar") 📊 - Best for: Part-to-whole relationships, composition analysis
        * x_col (required, categorical column name)
        * y_col (required, numerical column name)
        * color_col (required, categorical column name) - for stacking
        * agg_func (optional, one of ["mean", "sum", "count", "min", "max"], default "sum")
        * color_palette (optional, color scheme)
        Use when: Showing composition within categories, budget breakdowns, market share analysis

    - Line Chart (type: "line") 📈 - Best for: Time series, trends, sequential data
        * x_col (required, datetime or sequential column name)
        * y_col (required, numerical column name)
        * color_col_primary (optional, categorical column name) - for multiple series
        * agg_func_primary (optional, one of ["mean", "sum", "count", "min", "max"], default "mean")
        * show_markers (optional, boolean)
        * color_palette (optional, color scheme)
        Use when: Tracking changes over time, showing trends, comparing time series

    - Histogram (type: "histogram") 📏 - Best for: Distribution analysis, data quality assessment
        * col (required, numerical column name)
        * bins (optional, integer, default 20)
        * color_col (optional, categorical column name) - for grouped histograms
        * marginal (optional, one of [None, "rug", "box", "violin"])
        * cumulative (optional, boolean)
        * color_palette (optional, color scheme)
        Use when: Understanding data distribution, identifying skewness, quality control

    - Pie Chart (type: "pie") 🥧 - Best for: Proportional relationships, market share, composition
        * col (required, categorical column name)
        * max_categories (optional, integer, default 8)
        * hole (optional, float 0-1, for donut chart)
        * color_palette (optional, color scheme)
        Use when: Showing proportions, market share, survey results (limit to 3-8 categories)

    - Box Plot (type: "box") 📦 - Best for: Statistical analysis, outlier detection, group comparison
        * num_col (required, numerical column name)
        * cat_col (required, categorical column name)
        * color_col (optional, categorical column name)
        * orientation (optional, 'v' or 'h')
        * color_palette (optional, color scheme)
        Use when: Comparing distributions across groups, identifying outliers, statistical analysis

    - Correlation Heatmap (type: "correlation_heatmap") 🧪 - Best for: Feature relationships, multicollinearity
        * columns (required, list of numerical column names, minimum 3)
        * color_scale (optional, color scheme like "RdBu")
        * text_auto (optional, boolean)
        Use when: Analyzing relationships between multiple numeric variables, feature selection

    === INDUSTRY-STANDARD CHART TYPES ===

    - Waterfall Chart (type: "waterfall") 💹 - Best for: Financial analysis, cumulative impact, variance analysis
        * category_col (required, categorical column name) - sequential categories
        * value_col (required, numerical column name) - positive/negative values
        * show_total (optional, boolean, default true)
        * color_palette (optional, color scheme)
        Use when: Showing profit/loss breakdown, budget variance, sequential impact analysis

    - Gauge Chart (type: "gauge") 🎯 - Best for: KPI monitoring, performance against targets
        * value_col (required, numerical column name) - single metric
        * min_val (optional, number, default 0)
        * max_val (optional, number, default 100)
        * target_val (optional, number, default 75)
        * gauge_type (optional, "Semi-Circle", "Full Circle", or "Linear")
        Use when: Displaying single KPI, showing performance against targets, status indicators

    - Funnel Chart (type: "funnel") 📊 - Best for: Process analysis, conversion tracking, sales funnel
        * stage_col (required, categorical column name) - sequential stages
        * value_col (required, numerical column name) - values at each stage
        * show_percentages (optional, boolean, default true)
        * color_palette (optional, color scheme)
        Use when: Analyzing conversion rates, sales process, customer journey, process optimization

    - Area Chart (type: "area") 🌊 - Best for: Trend visualization, volume tracking, cumulative metrics
        * x_col (required, datetime or sequential column name)
        * y_col (required, numerical column name)
        * color_col (optional, categorical column name) - for multiple series
        * fill_mode (optional, "tozeroy", "tonexty", "tonext")
        * color_palette (optional, color scheme)
        Use when: Showing trends with emphasis on volume, cumulative growth, time series comparison

    - Treemap (type: "treemap") 🌳 - Best for: Hierarchical data, portfolio analysis, nested categories
        * labels_col (required, categorical column name) - categories
        * values_col (required, numerical column name) - sizes
        * parents_col (optional, categorical column name) - for hierarchy
        * color_palette (optional, color scheme)
        Use when: Visualizing hierarchical data, portfolio composition, nested category analysis

    - Violin Plot (type: "violin") 🔔 - Best for: Distribution comparison, statistical analysis
        * x_col (required, categorical column name) - groups
        * y_col (required, numerical column name) - values
        * color_col (optional, categorical column name)
        * color_palette (optional, color scheme)
        Use when: Comparing distribution shapes across groups, advanced statistical analysis

    - Bullet Chart (type: "bullet") 🎯 - Best for: Performance vs targets, benchmark comparison
        * value_col (required, numerical column name) - current performance
        * target_col (optional, numerical column name) - target values
        Use when: Comparing performance against targets, KPI dashboards, benchmark analysis

    - Grouped Bar Chart (type: "grouped_bar") 📊 - Best for: Multi-category comparison, side-by-side analysis
        * x_col (required, categorical column name)
        * y_col (required, numerical column name)
        * color_col (required, categorical column name) - for grouping
        * agg_func (optional, one of ["mean", "sum", "count", "min", "max"], default "mean")
        * color_palette (optional, color scheme)
        Use when: Comparing values across multiple categories with side-by-side bars

    - Dual-Axis Chart (type: "dual_axis") 📈 - Best for: Comparing two different metrics with different scales
        * x_col (required, shared dimension column name)
        * y_col_primary (required, numerical column name) - first metric
        * y_col_secondary (required, numerical column name) - second metric
        * color_col_primary (optional, categorical column name) - for primary series
        * color_col_secondary (optional, categorical column name) - for secondary series
        Use when: Showing relationship between two different metrics with different scales

    - Pareto Chart (type: "pareto") 📊 - Best for: Identifying the most significant factors (80/20 rule)
        * category_col (required, categorical column name)
        * value_col (required, numerical column name)
        * cumulative_percentage (optional, boolean, default true) - show cumulative line
        * color_palette (optional, color scheme)
        Use when: Identifying the most important factors, prioritizing issues, quality control

    - Sunburst Chart (type: "sunburst") 🌐 - Best for: Hierarchical data, multi-level composition
        * labels_col (required, categorical column name) - categories for the rings
        * values_col (required, numerical column name) - sizes of the segments
        * parents_col (optional, categorical column name) - for hierarchy levels
        * max_depth (optional, integer, default 2) - maximum hierarchy depth to show
        * branchvalues (optional, "total" or "remainder") - how to calculate values
        * color_palette (optional, color scheme)
        Use when: Showing hierarchical relationships, organizational structures, multi-level compositions

    - Sankey Diagram (type: "sankey") 📊 - Best for: Flow analysis, relationship mapping
        * source_col (required, categorical column name) - source entities
        * target_col (required, categorical column name) - target entities
        * value_col (required, numerical column name) - flow values
        * node_pad (optional, integer, default 15) - padding between nodes
        * node_thickness (optional, integer, default 30) - thickness of nodes
        * color_palette (optional, color scheme)
        Use when: Showing flows between entities, process mapping, resource allocation

    - Strip Chart (type: "strip") 📊 - Best for: Individual data point visualization
        * x_col (required, categorical column name) - categories
        * y_col (required, numerical column name) - values
        * color_col (optional, categorical column name) - for additional grouping
        * jitter (optional, float 0-1, default 0.3) - amount of jitter to avoid overplotting
        * stripmode (optional, "overlay" or "group") - how to display strips
        * color_palette (optional, color scheme)
        Use when: Showing individual data points, identifying outliers, visualizing distribution

    - Q-Q Plot (type: "qq_plot") 📈 - Best for: Normality testing, distribution comparison
        * data_col (required, numerical column name) - data to test
        * dist (optional, one of ["norm", "uniform", "expon", "logistic"], default "norm") - theoretical distribution
        * line (optional, one of ["45", "s", "r", "q", "none"], default "45") - reference line type
        Use when: Testing for normality, comparing to theoretical distributions, statistical validation

    - Density Plot (type: "density") 📊 - Best for: Probability distribution visualization
        * data_col (required, numerical column name) - data to visualize
        * group_col (optional, categorical column name) - for comparative densities
        * bandwidth (optional, float, default 0.5) - smoothing parameter
        * cumulative (optional, boolean) - show cumulative distribution
        * fill (optional, boolean, default true) - fill area under curve
        * color_palette (optional, color scheme)
        Use when: Visualizing probability distributions, comparing data smoothness, analyzing continuous data

    - Ridge Plot (type: "ridge") 📊 - Best for: Multi-distribution comparison
        * data_col (required, numerical column name) - data to visualize
        * category_col (required, categorical column name) - categories for stacking
        * overlap (optional, float, default 0.8) - amount of overlap between plots
        * bandwidth (optional, float, default 0.5) - smoothing parameter
        * color_palette (optional, color scheme)
        Use when: Comparing distributions across categories, visualizing multi-category data, statistical overview

    - Time Series Chart (type: "timeseries") 📈 - Best for: Temporal data analysis
        * date_col (required, datetime column name) - time dimension
        * value_col (required, numerical column name) - values to track
        * group_col (optional, categorical column name) - for multiple series
        * agg_func (optional, one of ["mean", "sum", "count", "min", "max", "median"], default "mean") - aggregation
        * show_confidence (optional, boolean) - show confidence interval
        * trendline (optional, one of [None, "linear", "lowess", "expanding", "rolling"]) - trendline type
        * color_palette (optional, color scheme)
        Use when: Analyzing trends over time, detecting seasonal patterns, preparing for forecasting

    - Forecast Chart (type: "forecast") 📈 - Best for: Future predictions, trend forecasting
        * date_col (required, datetime column name)
        * value_col (required, numerical column name)
        * periods (optional, integer, default 12) - number of periods to forecast
        * model_type (optional, one of ["linear", "exponential", "additive", "multiplicative"], default "linear")
        * show_confidence (optional, boolean, default true) - show confidence intervals
        Use when: Predicting future values, forecasting trends, planning based on historical patterns

    - Moving Average Chart (type: "moving_average") 📊 - Best for: Trend smoothing, noise reduction
        * date_col (required, datetime column name)
        * value_col (required, numerical column name)
        * window (optional, integer, default 7) - window size for moving average
        * ma_type (optional, "simple" or "exponential", default "simple")
        * show_original (optional, boolean, default true) - show original data alongside moving average
        Use when: Identifying underlying trends, reducing noise in time series data, technical analysis

    - Ranking Chart (type: "ranking") 🏆 - Best for: Performance ranking, top/bottom analysis
        * category_col (required, categorical column name)
        * value_col (required, numerical column name)
        * top_n (optional, integer, default 10) - number of top items to show
        * sort_order (optional, "Descending" or "Ascending", default "Descending")
        * orientation (optional, "Horizontal" or "Vertical", default "Horizontal")
        * color_palette (optional, color scheme)
        Use when: Ranking items by performance, identifying top/bottom performers, comparative analysis 

    - Seasonal Decomposition (type: \"seasonal\") 📊 - Best for: Analyzing seasonality, trend, and residual components
        * date_col (required, datetime column name)
        * value_col (required, numerical column name)
        * model_type (optional, \"additive\" or \"multiplicative\", default \"additive\")
        * period (optional, integer, default 12) - seasonal period
        Use when: Decomposing time series into components, identifying seasonal patterns, analyzing trends

    # In Comparative section:
    - Comparison Chart (type: \"comparison\") 📈 - Best for: Comparing multiple metrics across categories
        * category_col (required, categorical column name)
        * value_cols (required, list of numerical column names) - metrics to compare
        * chart_type (optional, \"Bar\", \"Line\", or \"Area\", default \"Bar\")
        * normalize (optional, boolean) - normalize values for comparison
        Use when: Comparing multiple KPIs, benchmarking performance, analyzing relative strengths

    - Slope Chart (type: \"slope\") 📉 - Best for: Showing changes between time points
        * category_col (required, categorical column name) - items to track
        * time_col (required, categorical column name) - time periods
        * value_col (required, numerical column name) - values to track
        Use when: Tracking progress over time, comparing changes, before-after analysis

    - Dot Plot (type: \"dot_plot\") 🔵 - Best for: Showing distributions with individual points
        * category_col (required, categorical column name)
        * value_col (required, numerical column name)
        * group_col (optional, categorical column name) - for grouping dots
        * orientation (optional, \"Horizontal\" or \"Vertical\", default \"Horizontal\")
        Use when: Visualizing distributions, comparing groups, showing individual data points       

    - Traffic Light Dashboard (type: "traffic_light") 🚦 - Best for: Multi-KPI monitoring, status reporting.
        * kpi_metrics (required, list of numerical column names): The metrics to monitor.
        * {metric}_red_threshold (optional, number): Red threshold for a specific metric.
        * {metric}_yellow_threshold (optional, number): Yellow threshold for a specific metric.
        * {metric}_green_threshold (optional, number): Green threshold for a specific metric.
        * layout_style (optional, "grid", "circular", "linear", "dashboard"): Visual layout of the lights.
        * show_sparklines (optional, boolean): Show a small trend line for each KPI.
        Use when: Monitoring multiple KPIs with traffic light status, executive dashboards, performance reporting.

    - Choropleth Map (type: "choropleth") 🗺️ - Best for: Geographic analysis, regional distribution.
        * location_col (required, categorical column with geographic names/codes like country names, ISO codes, or state abbreviations).
        * value_col (required, numerical column name): The metric to color the regions by.
        * map_style (optional, "open-street-map", "carto-positron", etc.): The base map style.
        * color_scale (optional, color scheme like "Viridis", "Plasma"): The color gradient for the values.
        * show_hover_data (optional, boolean): Show details on hover.
        * animation_duration (optional, milliseconds): Duration for map animations.
        Use when: Analyzing geographic patterns, regional sales, population density, or any location-based data.

    - Multi-Layer Donut Chart (type: "donut") 🍩 - Best for: Hierarchical data, multi-level composition.
        * levels (required, integer): Number of hierarchical levels (rings).
        * level_{X}_col (required for each level, categorical column name): Category for level X (e.g., level_0_col).
        * level_{X}_value (required for each level, numerical column name): Value for level X (e.g., level_0_value).
        * center_metric (optional, numerical column name): A key metric to display in the center of the donut.
        * hole_size (optional, float 0-1): The size of the central hole.
        * show_percentages (optional, boolean): Display percentages on slices.
        Use when: Showing hierarchical relationships (e.g., Continent -> Country), multi-level budget allocation, or complex compositions.
   
    
    

    === SPECIAL COMPONENTS ===

    - KPI Card (type: "kpi") ✅ - Best for: Key metrics, summary statistics, highlights
        * value_col (required, numerical column name)
        * agg_func (required, one of ["mean", "sum", "count", "min", "max", "median"])
        * title (optional, string)
        * prefix (optional, string like "$" or "%")
        * suffix (optional, string)
        * decimals (optional, integer, default 2)
        * Create comprehensive performance dashboards with multiple metrics
        * Include comparisons against targets, previous periods, or benchmarks
        * Use appropriate visual styles for executive reporting
        * Add trend indicators and performance context
        Use when: Highlighting key metrics, showing summary statistics, dashboard highlights

    - Data Table (type: "table") 📋 - Best for: Detailed data view, record listing, reference
        * columns (required, list of column names)
        * max_rows (optional, integer, default 10)
        Use when: Showing detailed records, data exploration, providing reference information

    CHART SELECTION GUIDELINES:
    - For time series data: Use line or area charts
    - For categorical comparisons: Use bar or stacked bar charts
    - For distributions: Use histogram, box, or violin plots
    - For relationships: Use scatter plots or correlation heatmaps
    - For proportions: Use pie charts (limit categories) or treemaps
    - For processes: Use funnel or waterfall charts
    - For KPIs: Use gauge, bullet, or KPI cards
    - For hierarchical data: Use treemaps
    - For statistical analysis: Use box or violin plots

    BEST PRACTICES:
    - Choose 3-6 complementary charts that tell a cohesive story
    - Consider data size: histograms need 20+ rows, correlations need 30+ rows
    - Limit pie chart categories to 3-8 for readability
    - Use color_col strategically to add analytical dimensions
    - Include KPI cards for key metrics
    - Add tables for detailed data when needed

    Note: For all charts, you can also include a 'custom_title' in the params to set a title for the chart.

    IMPORTANT: Only include parameters that are required. Do not include optional parameters if they are not set.

    Output only the JSON. Do not include any other text.
    """

    
    # Append the new chart types to the existing prompt
        full_prompt = prompt + prompt_addition
        
        return full_prompt

        # Prepare messages with conversation history
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": "Generate the dashboard configuration JSON."}
        ]
        
        # Add recent conversation history if available
        if st.session_state.dashboard_chat_history:
            recent_history = st.session_state.dashboard_chat_history[-4:]  # Last 2 exchanges
            messages = [messages[0]] + recent_history + [messages[1]]
        
        # Use LangChain LLM instead of direct OpenAI client
        response = llm.invoke(messages).content
        
        # Parse the JSON response
        dashboard_config = json.loads(response)
        
        # Validate the structure
        if "charts" not in dashboard_config or "insights" not in dashboard_config:
            st.error("❌ The AI response does not have the expected structure (missing 'charts' or 'insights').")
            return None
        
        # Add layout configuration
        dashboard_config["config"] = {"num_columns": 2}
        
        # Store in conversation history for context
        st.session_state.dashboard_chat_history.append({
            "role": "user", 
            "content": f"Create dashboard: {prompt}"
        })
        st.session_state.dashboard_chat_history.append({
            "role": "assistant", 
            "content": f"Created dashboard with {len(dashboard_config.get('charts', []))} charts"
        })
        
        return dashboard_config
        
    except ValueError as e:
        return f"❌ API key error: {str(e)}"
    except json.JSONDecodeError as e:
        st.error(f"❌ Error parsing AI response as JSON: {str(e)}")
        return None
    except Exception as e:
        st.error(f"❌ Error generating AI dashboard: {str(e)}")
        return None
# ==== END OPTIMIZED AI DASHBOARD FUNCTION =====

# ===== CHART METADATA MANAGEMENT FUNCTIONS =====

def bulk_update_chart_metadata(metadata_updates):
    """Bulk update multiple chart metadata"""
    if 'dashboard' not in st.session_state or not st.session_state.dashboard:
        return False
    
    charts = st.session_state.dashboard.get('charts', [])
    updated_count = 0
    
    for chart_index, metadata in metadata_updates.items():
        if 0 <= chart_index < len(charts):
            if 'params' not in charts[chart_index]:
                charts[chart_index]['params'] = {}
            
            if 'title' in metadata:
                charts[chart_index]['params']['custom_title'] = metadata['title']
            if 'description' in metadata:
                charts[chart_index]['params']['custom_description'] = metadata['description']
            
            updated_count += 1
    
    return updated_count

# ===== END CHART METADATA MANAGEMENT FUNCTIONS =====

def render_traffic_light_layout(params, df, custom_title):
    """Render traffic light dashboard with different layout styles"""
    kpi_metrics = params.get("kpi_metrics", [])
    layout_style = params.get("layout_style", "grid")
    show_sparklines = params.get("show_sparklines", True)
    show_values = params.get("show_values", True)

    if not kpi_metrics:
        st.warning("Please select at least one KPI metric for the Traffic Light chart.")
        return

    # Prepare data for each metric
    metrics_data = []
    for metric in kpi_metrics:
        value = df[metric].mean()
        red_thresh = params.get(f"{metric}_red_threshold", 0)
        yellow_thresh = params.get(f"{metric}_yellow_threshold", 0)
        
        if value <= red_thresh:
            color = "#FF4B4B" # Streamlit's error red
            status_text = "Critical"
        elif value <= yellow_thresh:
            color = "#FFC300" # A nice yellow
            status_text = "Warning"
        else:
            color = "#28A745" # Streamlit's success green
            status_text = "Good"
            
        metrics_data.append({
            "name": metric,
            "value": value,
            "color": color,
            "status": status_text,
            "thresholds": [red_thresh, yellow_thresh, df[metric].max()]
        })

    # --- CSS for different layouts ---
    st.markdown("""
    <style>
    /* Base container for all traffic light styles */
    .tl-container {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 1.2rem;
        border-radius: 12px;
        background-color: #f8f9fa;
        border: 1px solid #e9ecef;
        height: 100%;
        min-height: 220px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        transition: all 0.3s ease;
    }
    .tl-container:hover {
        transform: translateY(-4px);
        box-shadow: 0 8px 20px rgba(0,0,0,0.08);
    }
    .tl-label {
        font-weight: 600;
        font-size: 1.1rem;
        color: #31333F;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .tl-value {
        font-size: 2rem;
        font-weight: 700;
        color: #111;
    }
    .tl-status {
        font-size: 0.9rem;
        font-weight: 500;
        padding: 0.2rem 0.6rem;
        border-radius: 15px;
        margin-top: 0.75rem;
    }
    /* Grid/Linear specific circle */
    .tl-circle {
        width: 60px;
        height: 60px;
        border-radius: 50%;
        margin-bottom: 1rem;
        border: 4px solid rgba(0,0,0,0.08);
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.1), 0 4px 8px rgba(0,0,0,0.15);
    }
    /* Dashboard layout specific styles */
    .tl-dashboard-card {
        padding: 1rem;
        border-radius: 12px;
        background-color: #ffffff;
        border: 1px solid #e9ecef;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }
    .tl-dashboard-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 1rem;
    }
    .tl-dashboard-title {
        font-size: 1.1rem;
        font-weight: 600;
        color: #31333F;
    }
    .tl-dashboard-status-dot {
        width: 12px;
        height: 12px;
        border-radius: 50%;
    }
    .tl-dashboard-value {
        font-size: 2.2rem;
        font-weight: 700;
        color: #111;
        text-align: center;
    }
    </style>
    """, unsafe_allow_html=True)

    if custom_title:
        st.subheader(custom_title)

    if layout_style == "circular":
        fig = go.Figure()
        num_metrics = len(metrics_data)
        radius = 0.4
        
        for i, metric_data in enumerate(metrics_data):
            angle = (2 * np.pi * i / num_metrics) - (np.pi / 2) # Start from top
            x_pos = 0.5 + radius * np.cos(angle)
            y_pos = 0.5 + radius * np.sin(angle)
            
            fig.add_trace(go.Scatter(
                x=[x_pos], y=[y_pos],
                mode='markers',
                marker=dict(
                    color=metric_data['color'],
                    size=100,
                    line=dict(width=4, color='rgba(0,0,0,0.2)')
                ),
                hoverinfo='text',
                text=f"<b>{metric_data['name']}</b><br>Value: {metric_data['value']:.2f}<br>Status: {metric_data['status']}"
            ))
            
            fig.add_annotation(
                x=x_pos, y=y_pos,
                text=f"<b>{metric_data['name']}</b><br>{metric_data['value']:.2f}" if show_values else f"<b>{metric_data['name']}</b>",
                showarrow=False,
                font=dict(color='white' if metric_data['color'] != '#FFC300' else 'black', size=14)
            )
        
        fig.update_layout(
            showlegend=False,
            xaxis=dict(visible=False, range=[0, 1]),
            yaxis=dict(visible=False, range=[0, 1], scaleanchor="x", scaleratio=1),
            plot_bgcolor='rgba(0,0,0,0)',
            height=600,
            margin=dict(l=10, r=10, t=40, b=10),
            title="Circular KPI Status"
        )
        st.plotly_chart(fig, use_container_width=True)

    elif layout_style == "dashboard":
        num_cols = min(len(metrics_data), 3) if metrics_data else 1
        cols = st.columns(num_cols)
        for i, metric_data in enumerate(metrics_data):
            with cols[i % num_cols]:
                st.markdown(f"""
                <div class="tl-dashboard-card">
                    <div class="tl-dashboard-header">
                        <div class="tl-dashboard-title">{metric_data['name']}</div>
                        <div class="tl-dashboard-status-dot" style="background-color: {metric_data['color']};"></div>
                    </div>
                    {'<div class="tl-dashboard-value">{:.2f}</div>'.format(metric_data['value']) if show_values else ''}
                </div>
                """, unsafe_allow_html=True)
                if show_sparklines:
                    spark_fig = go.Figure(go.Scatter(
                        y=df[metric_data['name']].dropna().tail(30),
                        mode='lines',
                        line=dict(color=metric_data['color'], width=3),
                        fill='tozeroy',
                        fillcolor=metric_data['color'].replace(')', ', 0.2)').replace('rgb', 'rgba')
                    ))
                    spark_fig.update_layout(
                        height=80, showlegend=False,
                        xaxis=dict(visible=False), yaxis=dict(visible=False),
                        margin=dict(l=0, r=0, t=5, b=0), plot_bgcolor='rgba(0,0,0,0)'
                    )
                    st.plotly_chart(spark_fig, use_container_width=True)

    else: # grid or linear
        if layout_style == "linear":
            num_cols = len(metrics_data) if metrics_data else 1
        else: # grid
            num_cols = min(len(metrics_data), 4) if metrics_data else 1
        
        cols = st.columns(num_cols)
        for i, metric_data in enumerate(metrics_data):
            with cols[i % num_cols]:
                status_color = metric_data['color'].replace(')', ', 0.2)').replace('rgb', 'rgba')
                status_text_color = metric_data['color']
                
                st.markdown(f"""
                <div class="tl-container">
                    <div class="tl-circle" style="background-color: {metric_data['color']};"></div>
                    <div class="tl-label">{metric_data['name']}</div>
                    {'<div class="tl-value">{:.2f}</div>'.format(metric_data['value']) if show_values else ''}
                    <div class="tl-status" style="background-color: {status_color}; color: {status_text_color}; border: 1px solid {status_text_color};">{metric_data['status']}</div>
                </div>
                """, unsafe_allow_html=True)


def render_line_chart(params, df, custom_title):
    """
    Renders an advanced, industry-standard line chart with optional dual-axis
    and time aggregation capabilities.
    """
    # 1. Extract and validate parameters
    x_col = params.get("x_col")
    y_col = params.get("y_col")

    if not x_col or not y_col:
        st.warning("Line chart requires at least an X-axis and a Primary Y-axis.")
        return
    if x_col not in df.columns or y_col not in df.columns:
        st.warning(f"One or more selected columns ('{x_col}', '{y_col}') not found in the data.")
        return

    y_col_secondary = params.get("y_col_secondary")
    if y_col_secondary and y_col_secondary not in df.columns:
        st.warning(f"Secondary Y-axis column '{y_col_secondary}' not found. Ignoring.")
        y_col_secondary = None

    color_col = params.get("color_col")
    color_col_secondary = params.get("color_col_secondary")

    agg_func = params.get("agg_func", "mean")
    agg_func_secondary = params.get("agg_func_secondary", "mean")

    time_agg = params.get("time_agg")
    show_markers = params.get("show_markers", False)

    plot_df = df.copy()

    # 2. Prepare data with aggregation
    is_datetime = pd.api.types.is_datetime64_any_dtype(plot_df[x_col])

    # Define aggregation dictionary
    agg_dict = {y_col: agg_func}
    if y_col_secondary:
        agg_dict[y_col_secondary] = agg_func_secondary

    # Define grouping columns
    group_by_cols = []
    if color_col: group_by_cols.append(color_col)
    if y_col_secondary and color_col_secondary: group_by_cols.append(color_col_secondary)

    try:
        if is_datetime and time_agg:
            # Time-based aggregation
            if group_by_cols:
                plot_df = plot_df.groupby(group_by_cols).resample(time_agg, on=x_col).agg(agg_dict).reset_index()
            else:
                plot_df = plot_df.set_index(x_col).resample(time_agg).agg(agg_dict).reset_index()
        else:
            # Categorical aggregation
            group_by_cols.insert(0, x_col)
            plot_df = plot_df.groupby(list(set(group_by_cols))).agg(agg_dict).reset_index()
    except Exception as e:
        st.error(f"Data aggregation failed: {e}")
        return

    if plot_df.empty:
        st.warning("No data available after aggregation.")
        return

    plot_df = plot_df.sort_values(by=x_col)

    # 3. Create Figure with dual-axis if needed
    fig = make_subplots(specs=[[{"secondary_y": True}]]) if y_col_secondary else go.Figure()

    # 4. Add Traces
    mode = 'lines+markers' if show_markers else 'lines'

    # Primary Y-axis
    if color_col:
        for group_name, group_df in plot_df.groupby(color_col):
            fig.add_trace(go.Scatter(x=group_df[x_col], y=group_df[y_col], name=f"{group_name} ({y_col})", mode=mode, line=dict(width=2.5)), secondary_y=False)
    else:
        fig.add_trace(go.Scatter(x=plot_df[x_col], y=plot_df[y_col], name=y_col, mode=mode, line=dict(color='#1f77b4', width=2.5)), secondary_y=False)

    # Secondary Y-axis
    if y_col_secondary:
        if color_col_secondary:
            for group_name, group_df in plot_df.groupby(color_col_secondary):
                fig.add_trace(go.Scatter(x=group_df[x_col], y=group_df[y_col_secondary], name=f"{group_name} ({y_col_secondary})", mode=mode, line=dict(dash='dash', width=2.5)), secondary_y=True)
        else:
            fig.add_trace(go.Scatter(x=plot_df[x_col], y=plot_df[y_col_secondary], name=y_col_secondary, mode=mode, line=dict(color='#d62728', dash='dash', width=2.5)), secondary_y=True)

    # 5. Update Layout for a professional look
    title_text = custom_title or f"Line Chart: {y_col}" + (f" vs {y_col_secondary}" if y_col_secondary else "")
    fig.update_layout(
        title_text=title_text,
        xaxis_title=x_col,
        yaxis_title=y_col,
        legend_title_text="Legend",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        template="plotly_white",
        hovermode="x unified",
        margin=dict(l=40, r=40, t=80, b=40)
    )

    if y_col_secondary:
        fig.update_yaxes(title_text=y_col_secondary, secondary_y=True)

    # 6. Render chart
    st.plotly_chart(fig, use_container_width=True)    

def render_choropleth_chart(params, df, custom_title):
    """Render choropleth map with advanced features, aggregation, and robust error handling."""
    location_col = params.get("location_col")
    value_col = params.get("value_col") # This is the original value column
    map_type = params.get("map_type", "Standard")

    # 1. Initial validation
    if not location_col or not value_col:
        st.warning("Please select a location and value column for the Choropleth map.")
        return

    if location_col not in df.columns or value_col not in df.columns:
        st.error(f"Columns '{location_col}' or '{value_col}' not found in the data.")
        return

    # 2. Data Preparation (with Aggregation)
    plot_df = None
    agg_value_col = value_col # This will be the column used for plotting color

    if params.get("enable_aggregation"):
        agg_func = params.get("agg_func", "sum")
        st.info(f"Aggregating data by '{location_col}' using '{agg_func}' on '{value_col}'.")
        
        try:
            if agg_func == 'count':
                plot_df = df.groupby(location_col).size().reset_index(name='count')
                agg_value_col = 'count' # The new value column is 'count'
            else:
                if value_col not in df.select_dtypes(include=np.number).columns:
                    st.error(f"Aggregation function '{agg_func}' requires a numeric 'Value Column'. '{value_col}' is not numeric.")
                    return
                plot_df = df.groupby(location_col, as_index=False).agg({value_col: agg_func})
                agg_value_col = value_col # It remains the same
        except Exception as e:
            st.error(f"Failed to aggregate data: {e}")
            return
    else:
        # No aggregation, just drop NA from relevant columns
        plot_df = df[[location_col, value_col]].dropna()
        agg_value_col = value_col

    if plot_df.empty:
        st.warning("No data available for the selected columns after processing.")
        return

    # 3. Prepare common chart arguments
    chart_kwargs = {
        "data_frame": plot_df,
        "locations": location_col,
        "color": agg_value_col, # Use the (potentially new) aggregated value column
        "hover_name": location_col,
        "color_continuous_scale": params.get("color_scale", "Viridis"),
        "title": custom_title or f"Choropleth Map of {agg_value_col} by {location_col}"
    }

    try:
        # 4. Handle different map types and data sources
        if map_type == "Mapbox (Advanced)":
            # --- MAPBOX ENGINE LOGIC ---
            mapbox_token = params.get("mapbox_token_input", "") or st.secrets.get("MAPBOX_TOKEN", "")
            if not mapbox_token:
                st.error("Mapbox maps require an access token.")
                st.info("Please enter your token in the chart configuration or add it to your Streamlit secrets as `MAPBOX_TOKEN`.")
                return
            
            px.set_mapbox_access_token(mapbox_token)
            mapbox_kwargs = chart_kwargs.copy()
            mapbox_kwargs.update({"mapbox_style": "carto-positron", "opacity": 0.6})

            if params.get("use_custom_geojson", False):
                geojson_data = params.get("geojson_data")
                featureidkey = params.get("featureidkey")
                if not geojson_data or not featureidkey:
                    st.warning("Please provide both the GeoJSON data and the Feature ID Key.")
                    return
                if isinstance(geojson_data, str):
                    import requests
                    try:
                        response = requests.get(geojson_data)
                        response.raise_for_status()
                        geojson_features = response.json()
                    except Exception as e:
                        st.error(f"Failed to fetch GeoJSON from URL: {e}")
                        return
                else:
                    geojson_features = geojson_data
                mapbox_kwargs["geojson"] = geojson_features
                mapbox_kwargs["featureidkey"] = featureidkey
                from shapely.geometry import shape
                lats, lons = ([shape(f['geometry']).centroid.y for f in geojson_features.get("features", []) if 'geometry' in f],
                              [shape(f['geometry']).centroid.x for f in geojson_features.get("features", []) if 'geometry' in f])
                mapbox_kwargs.update({"zoom": 3, "center": {"lat": np.mean(lats), "lon": np.mean(lons)} if lats else {"lat": 0, "lon": 0}})
                fig = px.choropleth_mapbox(**mapbox_kwargs)
            else:
                mapbox_kwargs["zoom"] = 1
                fig = px.choropleth_mapbox(**mapbox_kwargs)
        else:
            # --- STANDARD ENGINE LOGIC ---
            standard_kwargs = chart_kwargs.copy()
            if params.get("use_custom_geojson", False):
                geojson_data = params.get("geojson_data")
                featureidkey = params.get("featureidkey")
                if not geojson_data or not featureidkey:
                    st.warning("Please provide both the GeoJSON data and the Feature ID Key.")
                    return
                if isinstance(geojson_data, str):
                    import requests
                    try:
                        response = requests.get(geojson_data)
                        response.raise_for_status()
                        geojson_features = response.json()
                    except Exception as e:
                        st.error(f"Failed to fetch GeoJSON from URL: {e}")
                        return
                else:
                    geojson_features = geojson_data
                standard_kwargs["geojson"] = geojson_features
                standard_kwargs["featureidkey"] = featureidkey
                fig = px.choropleth(**standard_kwargs)
                fig.update_geos(fitbounds="locations", visible=False)
            else:
                standard_kwargs["locationmode"] = params.get("locationmode", "country names")
                standard_kwargs["scope"] = params.get("scope", "world")
                fig = px.choropleth(**standard_kwargs)
                fig.update_geos(fitbounds="locations", visible=False)

        # 5. Final layout enhancements
        fig.update_layout(
            margin={"r":0,"t":40,"l":0,"b":0},
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
        
        # 6. Render the chart
        st.plotly_chart(fig, use_container_width=True)

    except Exception as e:
        st.error(f"Failed to create choropleth map: {e}")
        import traceback
        st.code(traceback.format_exc())
        st.info("Please check your configuration. For custom GeoJSON, ensure the 'Feature ID Key' matches a property in your GeoJSON file.")     




def render_chart(chart_config, df):
    """Renders a chart based on the provided configuration."""
    chart_type = chart_config.get("type")
    params = chart_config.get("params", {})
    custom_title = params.get("custom_title", "")  # Get custom title

    # Handle cases where required columns might be None or config is empty
    if not params:
        st.warning(f"⚠️ Chart configuration for {chart_type} is incomplete or missing data.")
        return
    
    if chart_type == "kpi":
        # Get metrics and aggregation functions
        metrics = params.get("metrics", [])
        agg_funcs = params.get("agg_funcs", {})
        layout_style = params.get("layout_style", "grid")
        color_scheme = params.get("color_scheme", "corporate")
        show_sparklines = params.get("show_sparklines", False)
        show_comparison = params.get("show_comparison", False)
        comparison_type = params.get("comparison_type", "period")
        target_values = params.get("target_values", {})
        
        # Create a title for the KPI section
        title = params.get("custom_title", "KPI Dashboard")
        st.subheader(title)
        
        if not metrics:
            st.error("❌ No metrics selected for KPI dashboard.")
            return None
        
        # Determine the number of columns based on layout style
        if layout_style == "grid":
            num_cols = min(4, len(metrics))
        elif layout_style == "cards":
            num_cols = min(3, len(metrics))
        elif layout_style == "minimal":
            num_cols = min(2, len(metrics))
        elif layout_style == "gauges":
            num_cols = min(2, len(metrics))
        else:
            num_cols = min(3, len(metrics))
        
        # Create columns for the KPIs
        cols = st.columns(num_cols)
        
        # Define color schemes
        color_schemes = {
            "corporate": ["#1f77b4", "#ff7f0e", "#2ca02c", "#d62728"],
            "vibrant": ["#FF6B6B", "#4ECDC4", "#FFE66D", "#6A0572"],
            "pastel": ["#FDCFDF", "#B5EAD7", "#FFDAC1", "#C7CEEA"],
            "mono": ["#666666", "#888888", "#AAAAAA", "#CCCCCC"]
        }
        
        colors = color_schemes.get(color_scheme, color_schemes["corporate"])
        
        # Render each KPI
        for i, metric in enumerate(metrics):
            agg_func = agg_funcs.get(metric, "mean")
            current_value = get_aggregated_value(df, metric, agg_func)
            
            # Calculate comparison values if needed
            delta_value = None
            if show_comparison:
                if comparison_type == "period":
                    # Simple implementation - compare with previous period
                    # You might want to implement a more sophisticated comparison
                    delta_value = current_value * 0.1  # Placeholder
                elif comparison_type == "target" and metric in target_values:
                    target = target_values[metric]
                    delta_value = current_value - target
            
            # Determine which column to use
            col_idx = i % num_cols
            with cols[col_idx]:
                # Format the value based on its magnitude
                if abs(current_value) >= 1000000:
                    formatted_value = f"${current_value/1000000:.1f}M"
                elif abs(current_value) >= 1000:
                    formatted_value = f"${current_value/1000:.1f}K"
                else:
                    formatted_value = f"${current_value:.2f}"
                
                # Create a metric card with styling
                st.markdown(f"""
                <div style="
                    background: linear-gradient(135deg, {colors[i % len(colors)]}20, {colors[i % len(colors)]}40);
                    border-radius: 12px;
                    padding: 1rem;
                    margin: 0.5rem 0;
                    border-left: 4px solid {colors[i % len(colors)]};
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                ">
                    <h4 style="margin: 0 0 0.5rem 0; color: #333; font-size: 0.9rem;">
                        {metric.replace('_', ' ').title()}
                    </h4>
                    <h3 style="margin: 0; color: #222; font-size: 1.5rem; font-weight: 700;">
                        {formatted_value}
                    </h3>
                    <p style="margin: 0.25rem 0 0 0; color: #666; font-size: 0.8rem;">
                        {agg_func.title()} value
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                # Display the metric using Streamlit's native component
                st.metric(
                    label=metric.replace('_', ' ').title(),
                    value=formatted_value,
                    delta=delta_value if show_comparison else None,
                    delta_color="normal"
                )
        
        # Add sparklines if requested
        if show_sparklines and len(metrics) > 0:
            st.markdown("---")
            st.subheader("Trends")
            
            # Create a simple line chart for each metric
            for metric in metrics:
                if metric in df.columns:
                    st.line_chart(df[metric].dropna(), use_container_width=True)
        
        return None

    try:
        # Initialize fig to None to prevent UnboundLocalError in complex branches
        fig = None

        # --- Enhanced Chart Rendering with More Properties ---
        if chart_type == "scatter":
            if not params.get("x_col") or not params.get("y_col"):
                st.warning("❌ Scatter plot requires both X and Y axes.")
                return

            # Prepare arguments for px.scatter
            scatter_kwargs = {
                "data_frame": df,
                "x": params["x_col"],
                "y": params["y_col"],
                "color": params.get("color_col") if params.get("color_col") else None,
                "size": params.get("size_col") if params.get("size_col") else None,
                "hover_data": params.get("hover_data", df.columns.tolist()) # Default to all if not specified
            }
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        scatter_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                else:
                    if hasattr(px.colors.sequential, color_palette):
                        if params.get("color_col"):
                            scatter_kwargs["color_continuous_scale"] = color_palette
                        else:
                            scatter_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
            
            # Add title
            default_title = f"{params['x_col']} vs {params['y_col']}"
            scatter_kwargs["title"] = custom_title if custom_title else default_title

            fig = px.scatter(**scatter_kwargs)
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "correlation_heatmap":
            columns = params.get("columns", [])
            if not columns or len(columns) < 2:
                st.warning("❌ Correlation heatmap requires at least 2 numerical columns.")
                return

            corr_matrix = df[columns].corr()
            # Use enhanced properties
            colorscale = params.get("color_scale", 'RdBu')
            text_auto = params.get("text_auto", True)

            fig = px.imshow(
                corr_matrix,
                title=custom_title if custom_title else "Correlation Heatmap",
                color_continuous_scale=colorscale,
                aspect='auto',
                text_auto=text_auto
            )
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "box":
            if not params.get("num_col") or not params.get("cat_col"):
                st.warning("❌ Box plot requires both a numerical and a categorical column.")
                return

            orientation = params.get("orientation", 'v')
            color = params.get("color_col") if params.get("color_col") else None

            fig = px.box(
                df,
                x=params["cat_col"] if orientation == 'v' else params["num_col"],
                y=params["num_col"] if orientation == 'v' else params["cat_col"],
                color=color,
                orientation=orientation,
                title=custom_title if custom_title else f"{params['num_col']} by {params['cat_col']}"
            )
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "line":
            render_line_chart(params, df, custom_title)

        elif chart_type == "bar":
            x_col = params.get("x_col")
            y_col = params.get("y_col")
            if not x_col:
                st.warning("❌ Bar chart requires an X-axis (categorical) column.")
                return

            agg_func = params.get("agg_func", "count")
            color_col = params.get("color_col")
            barmode = params.get("barmode", "group")

            # Validate columns exist before processing
            if params["x_col"] not in df.columns:
                st.warning(f"❌ Column '{params['x_col']}' not found in data.")
                return
            
            # Handle data preparation based on aggregation
            if agg_func == "count" and params["x_col"] == params.get("y_col", ""):
                df_agg = df[params["x_col"]].value_counts().reset_index()
                df_agg.columns = [params["x_col"], "count"]
                y_col = "count"
            else:
                # Validate y_col exists for non-count aggregations
                if params.get("y_col") and params["y_col"] not in df.columns:
                    st.warning(f"❌ Column '{params['y_col']}' not found in data.")
                    return
                
                group_cols = [params["x_col"]]
                if params.get("color_col"):
                    if params["color_col"] not in df.columns:
                        st.warning(f"❌ Color column '{params['color_col']}' not found in data.")
                        return
                    group_cols.append(params["color_col"])
                df_agg = df.groupby(group_cols)[params["y_col"]].agg(agg_func).reset_index()
                y_col = params["y_col"]
            
            # Prepare bar chart arguments
            bar_kwargs = {
                "data_frame": df_agg,
                "x": params["x_col"],
                "y": y_col,
                "color": params.get("color_col"),
                "barmode": params.get("barmode", "group"),
                "title": custom_title if custom_title else f"{agg_func.capitalize()} of {y_col} by {x_col}" +
                        (f" (Grouped by {color_col})" if color_col else "")
            }
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        bar_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                else:
                    if hasattr(px.colors.sequential, color_palette):
                        bar_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
            
            fig = px.bar(**bar_kwargs)
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "stacked_bar":
            if not all([params.get("x_col"), params.get("y_col"), params.get("color_col")]):
                st.warning("❌ Stacked bar chart requires X-axis, Y-axis, and Color columns.")
                return
                
            agg_func = params.get("agg_func", "mean")
            df_agg = df.groupby([params["x_col"], params["color_col"]])[params["y_col"]].agg(agg_func).reset_index()
            
            fig = px.bar(
                df_agg,
                x=params["x_col"],
                y=params["y_col"],
                color=params["color_col"],
                title=custom_title if custom_title else f"{agg_func.capitalize()} of {params['y_col']} by {params['x_col']} (Stacked by {params['color_col']})",
                barmode="stack"
            )
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "histogram":
            col_name = params.get("col")
            if not col_name:
                st.warning("❌ Histogram requires a numerical column.")
                return
            if col_name not in df.columns:
                st.warning(f"❌ Column '{col_name}' not found in dataset.")
                return
            if not pd.api.types.is_numeric_dtype(df[col_name]):
                st.warning(f"❌ Column '{col_name}' is not numeric. Please select a numeric column.")
                return

            # Prepare histogram arguments with all new features
            hist_kwargs = {
                "data_frame": df,
                "x": col_name,
                "nbins": params.get("bins", 20),
                "color": params.get("color_col"),
                "marginal": params.get("marginal"),
                "cumulative": params.get("cumulative", False),
                "facet_row": params.get("facet_row"),
                "facet_col": params.get("facet_col"),
                "histnorm": params.get("histnorm"),
                "barmode": params.get("barmode", "overlay"),
                "log_x": params.get("log_x", False),
                "log_y": params.get("log_y", False),
                "title": custom_title or f"Distribution of {col_name}"
            }
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        hist_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                else:
                    if hasattr(px.colors.sequential, color_palette):
                        hist_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
            
            fig = px.histogram(**hist_kwargs)
            
            # Further customize layout for faceting
            if params.get("facet_col") or params.get("facet_row"):
                fig.update_layout(bargap=0.1)
                fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[-1]))

            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "pie":
            mode = params.get("mode", "single") # Default to single if not set

            if mode == "single":
                # --- Render Single Pie Chart (Existing Enhanced Logic) ---
                slice_col = params.get("col")
                if not slice_col:
                    st.warning("❌ Pie chart requires a column for slices.")
                    return

                value_col = params.get("value_col")
                grouping_col = params.get("grouping_col")
                max_slices = params.get("max_categories", 8)
                hole_size = params.get("hole", 0.0)
                sort_by = params.get("sort", "Value")
                show_legend = params.get("show_legend", True)

                # Prepare data
                if value_col and value_col in df.columns:
                    if grouping_col and grouping_col in df.columns:
                                    plot_data = df.groupby([grouping_col, slice_col], as_index=False)[value_col].sum()
                    else:
                        plot_data = df.groupby(slice_col, as_index=False)[value_col].sum()
                else:
                    if grouping_col and grouping_col in df.columns:
                                    plot_data = df.groupby([grouping_col, slice_col]).size().reset_index(name='count')
                                    value_col = 'count'
                    else:
                        plot_data = df[slice_col].value_counts().reset_index()
                        plot_data.columns = [slice_col, 'count']
                        value_col = 'count'

                # Limit number of slices if needed (simplified logic)
                if max_slices and max_slices > 0:
                    if grouping_col:
                            # For grouped data, limit categories within each group or overall
                            top_categories_overall = df[slice_col].value_counts().head(max_slices).index
                            plot_data = plot_data[plot_data[slice_col].isin(top_categories_overall)]
                            # Handle 'Other' within groups if necessary (complex, simplified here)
                    else:
                        if len(plot_data) > max_slices:
                                top_data = plot_data.nlargest(max_slices, value_col)
                                other_sum = plot_data[~plot_data[slice_col].isin(top_data[slice_col])][value_col].sum()
                                other_row = pd.DataFrame({slice_col: ['Other'], value_col: [other_sum]})
                                plot_data = pd.concat([top_data, other_row], ignore_index=True)

                # Determine sorting
                if sort_by == "Label":
                    plot_data = plot_data.sort_values(by=slice_col)
                elif sort_by == "Value":
                    plot_data = plot_data.sort_values(by=value_col, ascending=False)

                # Create the figure
                # Prepare pie chart arguments
                pie_kwargs = {
                    "data_frame": plot_data,
                    "names": slice_col,
                    "values": value_col,
                    "title": custom_title if custom_title else f"Distribution of {slice_col}" + (f" grouped by {grouping_col}" if grouping_col else ""),
                    "hole": hole_size
                }
                
                # Apply color palette if specified
                color_palette = params.get("color_palette")
                if color_palette and color_palette != "--None--":
                    if color_palette.startswith("Qualitative_"):
                        palette_name = color_palette.replace("Qualitative_", "")
                        if hasattr(px.colors.qualitative, palette_name):
                            pie_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                    else:
                        if hasattr(px.colors.sequential, color_palette):
                            pie_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
                
                if grouping_col and grouping_col in plot_data.columns:
                    # Facetted pie charts
                    pie_kwargs["facet_col"] = grouping_col
                    pie_kwargs["facet_col_wrap"] = 3
                    fig = px.pie(**pie_kwargs)
                else:
                    fig = px.pie(**pie_kwargs)
                    # Sorting handled by data order for single pie

                if not show_legend and not grouping_col:
                    fig.update_layout(showlegend=False)

                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                # --- End Render Single Pie Chart ---

            elif mode == "multi":
                # --- Render Multi-Column Comparison ---
                columns_to_compare = params.get("columns_to_compare", [])
                value_col_multi = params.get("value_col_multi")
                agg_func_multi = params.get("agg_func_multi", "count")
                multi_viz_type = params.get("multi_viz_type", "Facetted Subplots")
                hole_size_multi = params.get("hole_multi", 0.0)
                max_slices_multi = params.get("max_categories_multi", 5)
                show_legend_multi = params.get("show_legend_multi", True)

                if not columns_to_compare or len(columns_to_compare) < 2:
                    st.warning("❌ Multi-Column Comparison mode requires selecting at least 2 columns.")
                    return

                # Prepare data for comparison
                # We want to show the distribution of 'value_col_multi' (or count) for each selected column.
                # This means for each column in columns_to_compare, we group by its categories and aggregate value_col_multi.
                comparison_data_list = []
                for col in columns_to_compare:
                    if value_col_multi and value_col_multi in df.columns:
                        # Aggregate the specified value column
                        temp_df = df.groupby(col, as_index=False)[value_col_multi].agg(agg_func_multi)
                        temp_df.rename(columns={col: 'Category', value_col_multi: 'Aggregated_Value'}, inplace=True)
                    else:
                        # Count occurrences
                        temp_df = df[col].value_counts().reset_index()
                        temp_df.columns = ['Category', 'Aggregated_Value']
                                
                    temp_df['Compared_Column'] = col # Add identifier for which column this data is for
                    comparison_data_list.append(temp_df)
                            
                if not comparison_data_list:
                        st.error("❌ Failed to prepare data for multi-column comparison.")
                        return

                combined_comparison_data = pd.concat(comparison_data_list, ignore_index=True)
                            
                # Limit categories per compared column for clarity
                if max_slices_multi and max_slices_multi > 0:
                    limited_dfs = []
                    for col in columns_to_compare:
                        col_data = combined_comparison_data[combined_comparison_data['Compared_Column'] == col]
                        if len(col_data) > max_slices_multi:
                            top_data = col_data.nlargest(max_slices_multi, 'Aggregated_Value')
                            other_sum = col_data[~col_data['Category'].isin(top_data['Category'])]['Aggregated_Value'].sum()
                            other_row = pd.DataFrame({
                                'Category': ['Other'],
                                'Aggregated_Value': [other_sum],
                                'Compared_Column': [col]
                            })
                            col_data = pd.concat([top_data, other_row], ignore_index=True)
                        limited_dfs.append(col_data)
                    combined_comparison_data = pd.concat(limited_dfs, ignore_index=True)

                # --- Visualization based on type ---
                if multi_viz_type == "Facetted Subplots":
                    # Create a facetted pie chart, one for each Compared_Column
                    pie_multi_kwargs = {
                        "data_frame": combined_comparison_data,
                        "names": 'Category',
                        "values": 'Aggregated_Value',
                        "facet_col": 'Compared_Column',
                        "facet_col_wrap": 3, # Adjust based on preference/screen size
                        "title": custom_title if custom_title else f"Comparison of Distributions for: {', '.join(columns_to_compare)}",
                        "hole": hole_size_multi
                    }
                    
                    # Apply color palette if specified
                    color_palette = params.get("color_palette")
                    if color_palette and color_palette != "--None--":
                        if color_palette.startswith("Qualitative_"):
                            palette_name = color_palette.replace("Qualitative_", "")
                            if hasattr(px.colors.qualitative, palette_name):
                                pie_multi_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                        else:
                            if hasattr(px.colors.sequential, color_palette):
                                pie_multi_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
                    
                    fig = px.pie(**pie_multi_kwargs)
                    if not show_legend_multi:
                        fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

                elif multi_viz_type == "Side-by-Side Bars":
                    # Create a grouped bar chart for comparison
                    bar_multi_kwargs = {
                        "data_frame": combined_comparison_data,
                        "x": 'Compared_Column',
                        "y": 'Aggregated_Value',
                        "color": 'Category', # Color by the original category names
                        "title": custom_title if custom_title else f"Comparison of Distributions for: {', '.join(columns_to_compare)}",
                        "barmode": 'group' # Group bars for each Compared_Column
                    }
                    
                    # Apply color palette if specified
                    color_palette = params.get("color_palette")
                    if color_palette and color_palette != "--None--":
                        if color_palette.startswith("Qualitative_"):
                            palette_name = color_palette.replace("Qualitative_", "")
                            if hasattr(px.colors.qualitative, palette_name):
                                bar_multi_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                        else:
                            if hasattr(px.colors.sequential, color_palette):
                                bar_multi_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
                    
                    fig = px.bar(**bar_multi_kwargs)
                    fig.update_xaxes(title_text="Compared Column")
                    fig.update_yaxes(title_text=f"Aggregated Value ({agg_func_multi})" if value_col_multi else "Count")
                    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                            
                elif multi_viz_type == "Radar Chart (if supported)":
                    st.info("Radar chart visualization is a concept for multi-column comparison. Displaying as Facetted Subplots instead.")
                    # Fallback or implement radar if feasible (more complex)
                    radar_fallback_kwargs = {
                        "data_frame": combined_comparison_data,
                        "names": 'Category',
                        "values": 'Aggregated_Value',
                        "facet_col": 'Compared_Column',
                        "facet_col_wrap": 3,
                        "title": custom_title if custom_title else f"Comparison of Distributions for: {', '.join(columns_to_compare)} (Radar Fallback)",
                        "hole": hole_size_multi
                    }
                    
                    # Apply color palette if specified
                    color_palette = params.get("color_palette")
                    if color_palette and color_palette != "--None--":
                        if color_palette.startswith("Qualitative_"):
                            palette_name = color_palette.replace("Qualitative_", "")
                            if hasattr(px.colors.qualitative, palette_name):
                                radar_fallback_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                        else:
                            if hasattr(px.colors.sequential, color_palette):
                                radar_fallback_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
                    
                    fig = px.pie(**radar_fallback_kwargs)
                    if not show_legend_multi:
                        fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                            
                else:
                    # Fallback
                    st.warning("Unknown multi-column visualization type selected. Displaying as Facetted Subplots.")
                    fallback_kwargs = {
                        "data_frame": combined_comparison_data,
                        "names": 'Category',
                        "values": 'Aggregated_Value',
                        "facet_col": 'Compared_Column',
                        "facet_col_wrap": 3,
                        "title": custom_title if custom_title else f"Comparison of Distributions for: {', '.join(columns_to_compare)}",
                        "hole": hole_size_multi
                    }
                    
                    # Apply color palette if specified
                    color_palette = params.get("color_palette")
                    if color_palette and color_palette != "--None--":
                        if color_palette.startswith("Qualitative_"):
                            palette_name = color_palette.replace("Qualitative_", "")
                            if hasattr(px.colors.qualitative, palette_name):
                                fallback_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                        else:
                            if hasattr(px.colors.sequential, color_palette):
                                fallback_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
                    
                    fig = px.pie(**fallback_kwargs)
                    if not show_legend_multi:
                        fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                # --- End Visualization based on type ---

                # --- End Render Multi-Column Comparison ---
            else:
                st.error(f"❌ Unsupported Pie Chart mode: {mode}")
                # --- END ENHANCED MULTI-COLUMN PIE CHART RENDERING ---

        elif chart_type == "kpi":
            # FIX: Add proper validation for metrics
            metrics = params.get("metrics", [])
            
            # Debug: Check what metrics we're getting
            print(f"KPI Metrics: {metrics}")  # For debugging
            
            if not metrics or (isinstance(metrics, list) and len(metrics) == 0):
                st.error("❌ No metrics selected for KPI dashboard. Please select at least one metric.")
                # Show a button to edit this chart
                if st.button("Edit This Chart", key="edit_kpi_no_metrics"):
                    # Find the chart index and set it for editing
                    for i, chart in enumerate(st.session_state.configured_charts):
                        if chart.get("id") == chart_config.get("id"):
                            st.session_state.temp_chart_config = {
                                "type": chart['type'],
                                "params": chart['params'].copy(),
                                "index": i,
                                "id": chart.get("id")
                            }
                            st.session_state.dashboard_edit_mode = True
                            st.rerun()
                return None
                
               
        elif chart_type == "table":
            columns = params.get("columns", [])
            if not columns:
                st.warning("❌ Table requires at least one column.")
                return

            st.dataframe(
                df[columns].head(params.get("max_rows", 10)),
                use_container_width=True,
                # hide_index=True # Optional: hide index column
            )

        # === INDUSTRY-STANDARD CHART RENDERINGS ===
        
        elif chart_type == "waterfall":
            category_col = params.get("category_col")
            value_col = params.get("value_col")
            show_total = params.get("show_total", True)
            
            if not category_col or not value_col:
                st.warning("❌ Waterfall chart requires both category and value columns.")
                return
            
            # Prepare waterfall data
            waterfall_data = df.groupby(category_col)[value_col].sum().reset_index()
            waterfall_data = waterfall_data.sort_values(value_col, ascending=False)
            
            # Create proper waterfall chart
            categories = waterfall_data[category_col].tolist()
            values = waterfall_data[value_col].tolist()
            
            # Determine measure types - all relative except total
            measures = ["relative"] * len(categories)
            
            if show_total:
                categories.append("Total")
                values.append(sum(values))
                measures.append("total")
            
            # Create proper waterfall chart with go.Waterfall
            fig = go.Figure(go.Waterfall(
                name="",
                orientation="v",
                measure=measures,
                x=categories,
                y=values,
                text=[f"{val:+,.0f}" if val != 0 else f"{sum(values[:-1]):,.0f}" for val in values],
                textposition="outside",
                connector={"line": {"color": "rgb(63, 63, 63)"}},
                increasing={"marker": {"color": "green"}},
                decreasing={"marker": {"color": "red"}},
                totals={"marker": {"color": "blue"}}
            ))
            
            fig.update_layout(
                title=custom_title if custom_title else f"Waterfall Analysis: {value_col} by {category_col}",
                showlegend=False,
                yaxis_title=value_col
            )
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "gauge":
            value_col = params.get("value_col")
            if not value_col:
                st.warning("❌ Gauge chart requires a value column.")
                return
            
            # Calculate the gauge value
            actual_value = df[value_col].mean()
            min_val = params.get("min_value", 0)
            max_val = params.get("max_value", 100)
            target_val = params.get("target_value", 75)
            
            # Create gauge chart
            fig = go.Figure(go.Indicator(
                mode = "gauge+number+delta",
                value = actual_value,
                domain = {'x': [0, 1], 'y': [0, 1]},
                title = {'text': custom_title if custom_title else f"Gauge: {value_col}"},
                delta = {'reference': target_val},
                gauge = {
                    'axis': {'range': [None, max_val]},
                    'bar': {'color': "darkblue"},
                    'steps': [
                        {'range': [min_val, target_val * 0.7], 'color': "lightgray"},
                        {'range': [target_val * 0.7, target_val], 'color': "gray"}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75,
                        'value': target_val
                    }
                }
            ))
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "funnel":
            stage_col = params.get("stage_col")
            value_col = params.get("value_col")
            
            if not stage_col or not value_col:
                st.warning("❌ Funnel chart requires both stage and value columns.")
                return
            
            # Prepare funnel data
            funnel_data = df.groupby(stage_col)[value_col].sum().reset_index()
            funnel_data = funnel_data.sort_values(value_col, ascending=False)
            
            # Create funnel chart
            fig = px.funnel(
                funnel_data,
                x=value_col,
                y=stage_col,
                title=custom_title if custom_title else f"Funnel Analysis: {stage_col}"
            )
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        fig.update_traces(marker_color=getattr(px.colors.qualitative, palette_name))
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "area":
            x_col = params.get("x_col")
            y_col = params.get("y_col")
            
            if not x_col or not y_col:
                st.warning("❌ Area chart requires both X and Y columns.")
                return
            
            # Create area chart
            area_kwargs = {
                "data_frame": df,
                "x": x_col,
                "y": y_col,
                "title": custom_title if custom_title else f"Area Chart: {y_col} over {x_col}"
            }
            
            color_col = params.get("color_col")
            if color_col:
                area_kwargs["color"] = color_col
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        area_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                else:
                    if hasattr(px.colors.sequential, color_palette):
                        area_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
            
            fig = px.area(**area_kwargs)
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "treemap":
            labels_col = params.get("labels_col")
            values_col = params.get("values_col")
            parent_col = params.get("parent_col")
            
            if not labels_col or not values_col:
                st.warning("❌ Treemap requires both labels and values columns.")
                return
            
            # Prepare treemap data with proper parent handling
            if parent_col and parent_col in df.columns:
                # For hierarchical treemap with parent-child relationships
                treemap_data = df.groupby([parent_col, labels_col])[values_col].sum().reset_index()
                
                treemap_kwargs = {
                    "data_frame": treemap_data,
                    "path": [parent_col, labels_col],
                    "values": values_col,
                    "title": custom_title if custom_title else f"Hierarchical Treemap: {values_col} by {parent_col} → {labels_col}"
                }
            else:
                # For simple treemap without hierarchy
                treemap_data = df.groupby(labels_col)[values_col].sum().reset_index()
                
                treemap_kwargs = {
                    "data_frame": treemap_data,
                    "path": [labels_col],
                    "values": values_col,
                    "title": custom_title if custom_title else f"Treemap: {values_col} by {labels_col}"
                }
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        treemap_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
                elif hasattr(px.colors.sequential, color_palette):
                    treemap_kwargs["color_discrete_sequence"] = getattr(px.colors.sequential, color_palette)
            
            fig = px.treemap(**treemap_kwargs)
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "violin":
            x_col = params.get("x_col")
            y_col = params.get("y_col")
            
            if not x_col or not y_col:
                st.warning("❌ Violin plot requires both category and numerical columns.")
                return
            
            # Create violin plot
            violin_kwargs = {
                "data_frame": df,
                "x": x_col,
                "y": y_col,
                "box": params.get("show_box", True),
                "title": custom_title if custom_title else f"Violin Plot: {y_col} by {x_col}"
            }
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        violin_kwargs["color_discrete_sequence"] = getattr(px.colors.qualitative, palette_name)
            
            fig = px.violin(**violin_kwargs)
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "bullet":
            value_col = params.get("value_col")
            target_col = params.get("target_col")
            
            if not value_col or not target_col:
                st.warning("❌ Bullet chart requires both actual and target value columns.")
                return
            
            # Create bullet chart using bar chart with target line
            actual_val = df[value_col].mean()
            target_val = df[target_col].mean()
            
            # Create a simple bullet chart representation
            bullet_data = pd.DataFrame({
                'Metric': ['Actual', 'Target'],
                'Value': [actual_val, target_val],
                'Type': ['Actual', 'Target']
            })
            
            fig = px.bar(
                bullet_data,
                x='Value',
                y='Metric',
                color='Type',
                orientation='h',
                title=custom_title if custom_title else f"Bullet Chart: {value_col} vs {target_col}"
            )
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "grouped_bar":
            if not all([params.get("x_col"), params.get("y_col"), params.get("color_col")]):
                st.warning("❌ Grouped bar chart requires X-axis, Y-axis, and Color columns.")
                return
                
            df_agg = df.groupby([params["x_col"], params["color_col"]])[params["y_col"]].mean().reset_index()
            
            fig = px.bar(
                df_agg,
                x=params["x_col"],
                y=params["y_col"],
                color=params["color_col"],
                title=custom_title if custom_title else f"{params['y_col']} by {params['x_col']} (Grouped by {params['color_col']})",
                barmode="group"
            )
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                fig = apply_color_palette_to_chart(fig, color_palette, chart_has_color_column=True)
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "dual_axis":
            if not all([params.get("x_col"), params.get("y_col_primary"), params.get("y_col_secondary")]):
                st.warning("❌ Dual-axis chart requires X-axis, Primary Y-axis, and Secondary Y-axis columns.")
                return
            
            # Create figure with secondary y-axis
            fig = go.Figure()
            
            # Add primary trace
            fig.add_trace(go.Scatter(
                x=df[params["x_col"]],
                y=df[params["y_col_primary"]],
                name=params["y_col_primary"],
                line=dict(color='blue')
            ))
            
            # Add secondary trace
            fig.add_trace(go.Scatter(
                x=df[params["x_col"]],
                y=df[params["y_col_secondary"]],
                name=params["y_col_secondary"],
                yaxis="y2",
                line=dict(color='red')
            ))
            
            # Create axis objects
            fig.update_layout(
                title=custom_title if custom_title else f"Dual-Axis: {params['y_col_primary']} vs {params['y_col_secondary']}",
                yaxis=dict(title=params["y_col_primary"]),
                yaxis2=dict(
                    title=params["y_col_secondary"],
                    overlaying="y",
                    side="right"
                )
            )
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "pareto":
            if not all([params.get("category_col"), params.get("value_col")]):
                st.warning("❌ Pareto chart requires Category and Value columns.")
                return
                
            # Prepare data for Pareto chart
            pareto_data = df.groupby(params["category_col"])[params["value_col"]].sum().reset_index()
            pareto_data = pareto_data.sort_values(params["value_col"], ascending=False)
            pareto_data['cumulative_percentage'] = (pareto_data[params["value_col"]].cumsum() / 
                                                   pareto_data[params["value_col"]].sum() * 100)
            
            # Create Pareto chart
            fig = go.Figure()
            
            # Add bar chart
            fig.add_trace(go.Bar(
                x=pareto_data[params["category_col"]],
                y=pareto_data[params["value_col"]],
                name="Values"
            ))
            
            # Add cumulative percentage line if requested
            if params.get("cumulative_percentage", True):
                fig.add_trace(go.Scatter(
                    x=pareto_data[params["category_col"]],
                    y=pareto_data['cumulative_percentage'],
                    name="Cumulative %",
                    yaxis="y2"
                ))
                
                fig.update_layout(
                    yaxis2=dict(
                        title="Cumulative Percentage",
                        overlaying="y",
                        side="right",
                        range=[0, 100]
                    )
                )
            
            fig.update_layout(
                title=custom_title if custom_title else f"Pareto Chart: {params['value_col']} by {params['category_col']}"
            )
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "sunburst":
            if not all([params.get("labels_col"), params.get("values_col")]):
                st.warning("❌ Sunburst chart requires Labels and Values columns.")
                return
                
            # Prepare data for sunburst chart
            if params.get("parents_col"):
                # Hierarchical sunburst with parent-child relationships
                sunburst_data = df.groupby([params["parents_col"], params["labels_col"]])[params["values_col"]].sum().reset_index()
                
                fig = px.sunburst(
                    sunburst_data,
                    path=[params["parents_col"], params["labels_col"]],
                    values=params["values_col"],
                    title=custom_title if custom_title else f"Sunburst: {params['values_col']} by {params['parents_col']} → {params['labels_col']}",
                    maxdepth=params.get("max_depth", 2),
                    branchvalues=params.get("branchvalues", "total")
                )
            else:
                # Simple sunburst without hierarchy
                sunburst_data = df.groupby(params["labels_col"])[params["values_col"]].sum().reset_index()
                
                fig = px.sunburst(
                    sunburst_data,
                    path=[params["labels_col"]],
                    values=params["values_col"],
                    title=custom_title if custom_title else f"Sunburst: {params['values_col']} by {params['labels_col']}",
                    maxdepth=params.get("max_depth", 1),
                    branchvalues=params.get("branchvalues", "total")
                )
            
            # branchvalues is now applied directly in px.sunburst() above
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                if color_palette.startswith("Qualitative_"):
                    palette_name = color_palette.replace("Qualitative_", "")
                    if hasattr(px.colors.qualitative, palette_name):
                        fig.update_traces(marker=dict(colors=getattr(px.colors.qualitative, palette_name)))
                elif hasattr(px.colors.sequential, color_palette):
                    fig.update_traces(marker=dict(colors=getattr(px.colors.sequential, color_palette)))
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "sankey":
            if not all([params.get("source_col"), params.get("target_col"), params.get("value_col")]):
                st.warning("❌ Sankey diagram requires Source, Target, and Value columns.")
                return
                
            # Prepare data for sankey diagram
            sankey_data = df.groupby([params["source_col"], params["target_col"]])[params["value_col"]].sum().reset_index()
            
            # Create unique list of all nodes
            all_nodes = pd.concat([sankey_data[params["source_col"]], sankey_data[params["target_col"]]]).unique()
            node_dict = {node: i for i, node in enumerate(all_nodes)}
            
            # Map source and target to indices
            sankey_data['source_idx'] = sankey_data[params["source_col"]].map(node_dict)
            sankey_data['target_idx'] = sankey_data[params["target_col"]].map(node_dict)
            
            # Create the sankey diagram
            fig = go.Figure(data=[go.Sankey(
                node=dict(
                    pad=params.get("node_pad", 15),
                    thickness=params.get("node_thickness", 30),
                    line=dict(color="black", width=0.5),
                    label=list(node_dict.keys())
                ),
                link=dict(
                    source=sankey_data['source_idx'],
                    target=sankey_data['target_idx'],
                    value=sankey_data[params["value_col"]]
                )
            )])
            
            fig.update_layout(
                title_text=custom_title if custom_title else f"Sankey Diagram: {params['source_col']} to {params['target_col']}",
                font_size=10
            )
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
            
        elif chart_type == "strip":
            if not all([params.get("x_col"), params.get("y_col")]):
                st.warning("❌ Strip chart requires X and Y columns.")
                return
            
            # Create strip chart
            fig = px.strip(
                df,
                x=params["x_col"],
                y=params["y_col"],
                color=params.get("color_col"),
                title=custom_title if custom_title else f"Strip Plot: {params['y_col']} by {params['x_col']}",
                stripmode=params.get("stripmode", "overlay")
            )
            
            # Apply jitter if specified
            if params.get("jitter", 0.3) > 0:
                fig.update_traces(jitter=params["jitter"])
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                fig = apply_color_palette_to_chart(fig, color_palette, chart_has_color_column=bool(params.get("color_col")))
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
            
        elif chart_type == "qq_plot":
            if not params.get("data_col"):
                st.warning("❌ Q-Q plot requires a Data column.")
                return
                
            # Import stats for Q-Q plot
            from scipy import stats
            import statsmodels.api as sm
            
            # Get the data
            data = df[params["data_col"]].dropna()
            
            # Create Q-Q plot
            fig = go.Figure()
            
            # Create theoretical quantiles based on selected distribution
            if params.get("dist", "norm") == "norm":
                theoretical_quantiles = stats.norm.ppf(np.linspace(0.01, 0.99, len(data)))
                dist_name = "Normal"
            elif params.get("dist", "norm") == "uniform":
                theoretical_quantiles = stats.uniform.ppf(np.linspace(0.01, 0.99, len(data)))
                dist_name = "Uniform"
            elif params.get("dist", "norm") == "expon":
                theoretical_quantiles = stats.expon.ppf(np.linspace(0.01, 0.99, len(data)))
                dist_name = "Exponential"
            elif params.get("dist", "norm") == "logistic":
                theoretical_quantiles = stats.logistic.ppf(np.linspace(0.01, 0.99, len(data)))
                dist_name = "Logistic"
            
            # Sort the data for Q-Q plot
            data_sorted = np.sort(data)
            
            # Add the Q-Q points
            fig.add_trace(go.Scatter(
                x=theoretical_quantiles,
                y=data_sorted,
                mode='markers',
                name='Data Points'
            ))
            
            # Add reference line based on selection
            line_type = params.get("line", "45")
            if line_type != "none":
                if line_type == "45":
                    # 45-degree line
                    min_val = min(theoretical_quantiles.min(), data_sorted.min())
                    max_val = max(theoretical_quantiles.max(), data_sorted.max())
                    fig.add_trace(go.Scatter(
                        x=[min_val, max_val],
                        y=[min_val, max_val],
                        mode='lines',
                        name='45° Line',
                        line=dict(color='red', dash='dash')
                    ))
                elif line_type in ["s", "r", "q"]:
                    # Add a trend line (simplified)
                    z = np.polyfit(theoretical_quantiles, data_sorted, 1)
                    p = np.poly1d(z)
                    fig.add_trace(go.Scatter(
                        x=theoretical_quantiles,
                        y=p(theoretical_quantiles),
                        mode='lines',
                        name='Trend Line',
                        line=dict(color='red', dash='dash')
                    ))
            
            fig.update_layout(
                title=custom_title if custom_title else f"Q-Q Plot: {params['data_col']} vs {dist_name} Distribution",
                xaxis_title=f"Theoretical Quantiles ({dist_name})",
                yaxis_title="Sample Quantiles"
            )
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True}) 
        

        elif chart_type == "density":
            if not params.get("data_col"):
                st.warning("❌ Density plot requires a Data column.")
                return
                
            try:
                from scipy.stats import gaussian_kde
                import numpy as np
                
                # Precompute KDE
                data = df[params["data_col"]].dropna()
                kde = gaussian_kde(data, bw_method=params.get("bandwidth"))
                
                # Create evaluation points
                x_range = np.linspace(data.min(), data.max(), 1000)
                y_values = kde(x_range)
                
                # Create figure
                fig = go.Figure()
                fig.add_trace(go.Scatter(
                    x=x_range,
                    y=y_values,
                    fill='tozeroy' if params.get("fill", True) else None,
                    mode='lines',
                    name='Density'
                ))
                
                # Add rug plot if requested
                if params.get("show_rug", True):
                    fig.add_trace(go.Scatter(
                        x=data,
                        y=[0] * len(data),
                        mode='markers',
                        marker=dict(symbol='line-ns-open', size=10),
                        name='Rug'
                    ))
                
                # Update layout
                fig.update_layout(
                    title=custom_title if custom_title else f"Density Plot: {params['data_col']}",
                    xaxis_title=params["data_col"],
                    yaxis_title="Density"
                )
                
                # Apply color palette if specified
                color_palette = params.get("color_palette")
                if color_palette and color_palette != "--None--":
                    fig = apply_color_palette_to_chart(fig, color_palette, False)
                    
            except ImportError:
                st.error("❌ SciPy is required for advanced density plots. Please install it with 'pip install scipy'")
                return
                
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
            
        elif chart_type == "ridge":
            if not all([params.get("data_col"), params.get("category_col")]):
                st.warning("❌ Ridge plot requires Data and Category columns.")
                return
                
            # Create ridge plot using density contours
            fig = px.density_contour(
                df,
                x=params["data_col"],
                y=params["category_col"],
                title=custom_title if custom_title else f"Ridge Plot: {params['data_col']} by {params['category_col']}",
                color=params["category_col"]
            )
            
            # Update for ridge plot style
            fig.update_traces(
                contours_coloring="fill",
                contours_showlabels=False,
                line_width=0
            )
            
            # Apply overlap if specified
            if params.get("overlap"):
                fig.update_layout(barmode="overlay")
            
            # Apply bandwidth if specified
            if params.get("bandwidth"):
                fig.update_traces(contours=dict(size=params["bandwidth"]))
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                fig = apply_color_palette_to_chart(fig, color_palette, chart_has_color_column=True)
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
            
        elif chart_type == "timeseries":
            if not all([params.get("date_col"), params.get("value_col")]):
                st.warning("❌ Time series chart requires Date and Value columns.")
                return
                
            # Prepare data for time series
            if params.get("group_col"):
                # Grouped time series
                df_agg = df.groupby([params["date_col"], params["group_col"]])[params["value_col"]].agg(params.get("agg_func", "mean")).reset_index()
                
                fig = px.line(
                    df_agg,
                    x=params["date_col"],
                    y=params["value_col"],
                    color=params["group_col"],
                    title=custom_title if custom_title else f"Time Series: {params['value_col']} over time by {params['group_col']}"
                )
            else:
                # Simple time series
                df_agg = df.groupby(params["date_col"])[params["value_col"]].agg(params.get("agg_func", "mean")).reset_index()
                
                fig = px.line(
                    df_agg,
                    x=params["date_col"],
                    y=params["value_col"],
                    title=custom_title if custom_title else f"Time Series: {params['value']} over time"
                )
            
            # Add confidence interval if specified
            if params.get("show_confidence", False):
                # Calculate confidence interval (simplified)
                if params.get("group_col"):
                    # For grouped data, calculate CI per group
                    for trace in fig.data:
                        group_name = trace.name
                        group_data = df_agg[df_agg[params["group_col"]] == group_name]
                        mean = group_data[params["value_col"]].mean()
                        std = group_data[params["value_col"]].std()
                        n = len(group_data)
                        ci = 1.96 * std / np.sqrt(n) if n > 1 else 0
                        
                        # Add confidence interval as a filled area
                        fig.add_trace(go.Scatter(
                            x=group_data[params["date_col"]],
                            y=group_data[params["value_col"]] + ci,
                            fill=None,
                            mode='lines',
                            line=dict(width=0),
                            showlegend=False
                        ))
                        fig.add_trace(go.Scatter(
                            x=group_data[params["date_col"]],
                            y=group_data[params["value_col"]] - ci,
                            fill='tonexty',
                            mode='lines',
                            line=dict(width=0),
                            showlegend=False
                        ))
                else:
                    # For simple time series
                    mean = df_agg[params["value_col"]].mean()
                    std = df_agg[params["value_col"]].std()
                    n = len(df_agg)
                    ci = 1.96 * std / np.sqrt(n) if n > 1 else 0
                    
                    # Add confidence interval as a filled area
                    fig.add_trace(go.Scatter(
                        x=df_agg[params["date_col"]],
                        y=df_agg[params["value_col"]] + ci,
                        fill=None,
                        mode='lines',
                        line=dict(width=0),
                        showlegend=False
                    ))
                    fig.add_trace(go.Scatter(
                        x=df_agg[params["date_col"]],
                        y=df_agg[params["value_col"]] - ci,
                        fill='tonexty',
                        mode='lines',
                        line=dict(width=0),
                        showlegend=False
                    ))
            
            # Add trendline if specified
            if params.get("trendline"):
                if params["trendline"] == "linear":
                    # Add linear trendline
                    if params.get("group_col"):
                        for trace in fig.data:
                            if trace.mode == "lines":  # Only for main traces, not CI
                                group_name = trace.name
                                group_data = df_agg[df_agg[params["group_col"]] == group_name]
                                z = np.polyfit(range(len(group_data)), group_data[params["value_col"]], 1)
                                p = np.poly1d(z)
                                fig.add_trace(go.Scatter(
                                    x=group_data[params["date_col"]],
                                    y=p(range(len(group_data))),
                                    mode='lines',
                                    line=dict(dash='dash'),
                                    name=f"{group_name} Trend",
                                    showlegend=True
                                ))
                    else:
                        z = np.polyfit(range(len(df_agg)), df_agg[params["value_col"]], 1)
                        p = np.poly1d(z)
                        fig.add_trace(go.Scatter(
                            x=df_agg[params["date_col"]],
                            y=p(range(len(df_agg))),
                            mode='lines',
                            line=dict(dash='dash'),
                            name="Trend",
                            showlegend=True
                        ))
                # Add other trendline types as needed
            
            # Apply color palette if specified
            color_palette = params.get("color_palette")
            if color_palette and color_palette != "--None--":
                fig = apply_color_palette_to_chart(fig, color_palette, chart_has_color_column=bool(params.get("group_col")))
            
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})

        elif chart_type == "forecast":
            if not params.get("date_col") or not params.get("value_col"):
                st.warning("❌ Forecast chart requires both date and value columns.")
                return
            
            try:
                from statsmodels.tsa.api import ExponentialSmoothing, SimpleExpSmoothing
                from sklearn.linear_model import LinearRegression
                import warnings
                warnings.filterwarnings('ignore')
                
                # Prepare data
                df_ts = df[[params["date_col"], params["value_col"]]].copy()
                df_ts = df_ts.sort_values(params["date_col"])
                df_ts = df_ts.set_index(params["date_col"])
                
                # Create forecast based on model type
                periods = params.get("periods", 12)
                show_confidence = params.get("show_confidence", True)
                
                if params.get("model_type", "linear") == "linear":
                    # Linear regression forecast
                    X = np.array(range(len(df_ts))).reshape(-1, 1)
                    y = df_ts[params["value_col"]].values
                    
                    model = LinearRegression()
                    model.fit(X, y)
                    
                    # Forecast future values
                    future_X = np.array(range(len(df_ts), len(df_ts) + periods)).reshape(-1, 1)
                    forecast = model.predict(future_X)
                    
                    # Create future dates
                    last_date = df_ts.index.max()
                    if isinstance(last_date, pd.Timestamp):
                        freq = pd.infer_freq(df_ts.index)
                        future_dates = pd.date_range(start=last_date, periods=periods+1, freq=freq)[1:]
                    else:
                        future_dates = range(len(df_ts), len(df_ts) + periods)
                    
                    # Create figure
                    fig = go.Figure()
                    
                    # Add historical data
                    fig.add_trace(go.Scatter(
                        x=df_ts.index,
                        y=df_ts[params["value_col"]],
                        mode='lines',
                        name='Historical'
                    ))
                    
                    # Add forecast
                    fig.add_trace(go.Scatter(
                        x=future_dates,
                        y=forecast,
                        mode='lines',
                        name='Forecast',
                        line=dict(dash='dash')
                    ))
                    
                else:
                    # Time series models (exponential smoothing)
                    if params["model_type"] == "exponential":
                        model = ExponentialSmoothing(df_ts[params["value_col"]], trend='add')
                    else:
                        model = SimpleExpSmoothing(df_ts[params["value_col"]])
                        
                    fitted_model = model.fit()
                    forecast_result = fitted_model.forecast(periods)
                    
                    # Create figure
                    fig = go.Figure()
                    
                    # Add historical data
                    fig.add_trace(go.Scatter(
                        x=df_ts.index,
                        y=df_ts[params["value_col"]],
                        mode='lines',
                        name='Historical'
                    ))
                    
                    # Add forecast
                    fig.add_trace(go.Scatter(
                        x=forecast_result.index,
                        y=forecast_result.values,
                        mode='lines',
                        name='Forecast',
                        line=dict(dash='dash')
                    ))
                    
                    if show_confidence and hasattr(fitted_model, 'prediction_intervals'):
                        # Add confidence intervals if available
                        conf_int = fitted_model.prediction_intervals(periods)
                        fig.add_trace(go.Scatter(
                            x=conf_int.index.tolist() + conf_int.index.tolist()[::-1],
                            y=conf_int.iloc[:, 0].tolist() + conf_int.iloc[:, 1].tolist()[::-1],
                            fill='toself',
                            fillcolor='rgba(0,100,80,0.2)',
                            line=dict(color='rgba(255,255,255,0)'),
                            name='Confidence Interval'
                        ))
                
                fig.update_layout(
                    title=custom_title if custom_title else f"{params['value_col']} Forecast",
                    xaxis_title=params["date_col"],
                    yaxis_title=params["value_col"]
                )
                
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                
            except Exception as e:
                st.error(f"Error creating forecast: {str(e)}")

        elif chart_type == "moving_average":
            if not params.get("date_col") or not params.get("value_col"):
                st.warning("❌ Moving average chart requires both date and value columns.")
                return
            
            try:
                # Prepare data
                df_ma = df[[params["date_col"], params["value_col"]]].copy()
                df_ma = df_ma.sort_values(params["date_col"])
                df_ma = df_ma.set_index(params["date_col"])
                
                window = params.get("window", 7)
                ma_type = params.get("ma_type", "simple")
                show_original = params.get("show_original", True)
                
                # Calculate moving average
                if ma_type == "simple":
                    df_ma['ma'] = df_ma[params["value_col"]].rolling(window=window).mean()
                else:  # exponential
                    df_ma['ma'] = df_ma[params["value_col"]].ewm(span=window).mean()
                
                # Create figure
                fig = go.Figure()
                
                if show_original:
                    fig.add_trace(go.Scatter(
                        x=df_ma.index,
                        y=df_ma[params["value_col"]],
                        mode='lines',
                        name='Original',
                        opacity=0.5
                    ))
                
                fig.add_trace(go.Scatter(
                    x=df_ma.index,
                    y=df_ma['ma'],
                    mode='lines',
                    name=f'{window}-Period {ma_type.title()} MA'
                ))
                
                fig.update_layout(
                    title=custom_title if custom_title else f"{params['value_col']} Moving Average",
                    xaxis_title=params["date_col"],
                    yaxis_title=params["value_col"]
                )
                
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                
            except Exception as e:
                st.error(f"Error creating moving average: {str(e)}")

        elif chart_type == "ranking":
            if not params.get("category_col") or not params.get("value_col"):
                st.warning("❌ Ranking chart requires both category and value columns.")
                return
            
            try:
                # Prepare data
                top_n = params.get("top_n", 10)
                sort_order = params.get("sort_order", "Descending")
                orientation = params.get("orientation", "Horizontal")
                
                df_rank = df.groupby(params["category_col"])[params["value_col"]].mean().reset_index()
                
                # Sort data
                ascending = True if sort_order == "Ascending" else False
                df_rank = df_rank.sort_values(params["value_col"], ascending=ascending)
                
                # Get top N items
                df_rank = df_rank.head(top_n)
                
                # Create figure
                if orientation == "Horizontal":
                    fig = px.bar(
                        df_rank,
                        y=params["category_col"],
                        x=params["value_col"],
                        orientation='h',
                        title=custom_title if custom_title else f"Top {top_n} {params['category_col']} by {params['value_col']}"
                    )
                else:
                    fig = px.bar(
                        df_rank,
                        x=params["category_col"],
                        y=params["value_col"],
                        title=custom_title if custom_title else f"Top {top_n} {params['category_col']} by {params['value_col']}"
                    )
                    
                    # Rotate x-axis labels for vertical bars
                    fig.update_layout(xaxis_tickangle=-45)
                
                # Apply color palette if specified
                color_palette = params.get("color_palette")
                if color_palette and color_palette != "--None--":
                    fig = apply_color_palette_to_chart(fig, color_palette)
                
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                
            except Exception as e:
                st.error(f"Error creating ranking chart: {str(e)}")


        elif chart_type == "seasonal":
            if not params.get("date_col") or not params.get("value_col"):
                st.warning("❌ Seasonal decomposition requires both date and value columns.")
                return
            
            try:
                # Perform seasonal decomposition
                from statsmodels.tsa.seasonal import seasonal_decompose
                
                # Prepare time series data
                ts_data = df.set_index(params["date_col"])[params["value_col"]]
                if not pd.api.types.is_datetime64_any_dtype(ts_data.index):
                    ts_data.index = pd.to_datetime(ts_data.index)
                
                # Handle missing values
                ts_data = ts_data.dropna()
                
                # Perform decomposition
                decomposition = seasonal_decompose(
                    ts_data, 
                    model=params.get("model_type", "additive"),
                    period=params.get("period", 12)
                )
                
                # Create subplots
                fig = make_subplots(
                    rows=4, cols=1,
                    subplot_titles=('Observed', 'Trend', 'Seasonal', 'Residual')
                )
                
                # Add traces
                fig.add_trace(go.Scatter(x=ts_data.index, y=ts_data, name='Observed'), row=1, col=1)
                fig.add_trace(go.Scatter(x=ts_data.index, y=decomposition.trend, name='Trend'), row=2, col=1)
                fig.add_trace(go.Scatter(x=ts_data.index, y=decomposition.seasonal, name='Seasonal'), row=3, col=1)
                fig.add_trace(go.Scatter(x=ts_data.index, y=decomposition.resid, name='Residual'), row=4, col=1)
                
                # Update layout
                fig.update_layout(
                    height=800,
                    title_text=custom_title if custom_title else "Seasonal Decomposition",
                    showlegend=False
                )
                
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                
            except Exception as e:
                st.error(f"Error performing seasonal decomposition: {str(e)}")

        elif chart_type == "comparison":
            if not params.get("category_col") or not params.get("value_cols"):
                st.warning("❌ Comparison chart requires category column and at least one value column.")
                return
            
            try:
                # Prepare data for comparison
                comparison_df = df.groupby(params["category_col"])[params["value_cols"]].mean().reset_index()
                
                if params.get("normalize", False):
                    # Normalize the data
                    comparison_df[params["value_cols"]] = comparison_df[params["value_cols"]].apply(
                        lambda x: (x - x.min()) / (x.max() - x.min()) if x.max() != x.min() else 0
                    )
                
                # Melt the dataframe for plotting
                melted_df = comparison_df.melt(
                    id_vars=[params["category_col"]], 
                    value_vars=params["value_cols"],
                    var_name='Metric', 
                    value_name='Value'
                )
                
                # Create the appropriate chart type
                chart_type = params.get("chart_type", "Bar").lower()
                if chart_type == "bar":
                    fig = px.bar(
                        melted_df,
                        x=params["category_col"],
                        y="Value",
                        color="Metric",
                        barmode="group",
                        title=custom_title if custom_title else "Comparison Chart"
                    )
                elif chart_type == "line":
                    fig = px.line(
                        melted_df,
                        x=params["category_col"],
                        y="Value",
                        color="Metric",
                        title=custom_title if custom_title else "Comparison Chart"
                    )
                else:  # area
                    fig = px.area(
                        melted_df,
                        x=params["category_col"],
                        y="Value",
                        color="Metric",
                        title=custom_title if custom_title else "Comparison Chart"
                    )
                    
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                
            except Exception as e:
                st.error(f"Error creating comparison chart: {str(e)}")

        elif chart_type == "slope":
            if not params.get("category_col") or not params.get("time_col") or not params.get("value_col"):
                st.warning("❌ Slope chart requires category, time period, and value columns.")
                return
            
            try:
                # Prepare data for slope chart
                slope_data = df.groupby([params["category_col"], params["time_col"]])[params["value_col"]].mean().reset_index()
                
                # Pivot the data
                pivot_data = slope_data.pivot(
                    index=params["category_col"],
                    columns=params["time_col"],
                    values=params["value_col"]
                ).reset_index()
                
                # Create the slope chart
                fig = go.Figure()
                
                # Add lines for each category
                for _, row in pivot_data.iterrows():
                    category = row[params["category_col"]]
                    values = row.drop(params["category_col"]).values
                    times = pivot_data.columns.drop(params["category_col"])
                    
                    fig.add_trace(go.Scatter(
                        x=times,
                        y=values,
                        name=category,
                        mode='lines+markers'
                    ))
                
                fig.update_layout(
                    title=custom_title if custom_title else "Slope Chart",
                    xaxis_title="Time Period",
                    yaxis_title=params["value_col"]
                )
                
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                
            except Exception as e:
                st.error(f"Error creating slope chart: {str(e)}")

        elif chart_type == "dot_plot":
            if not params.get("category_col") or not params.get("value_col"):
                st.warning("❌ Dot plot requires category and value columns.")
                return
            
            try:
                # Prepare data
                if params.get("group_col"):
                    dot_data = df.groupby([params["category_col"], params["group_col"]])[params["value_col"]].mean().reset_index()
                    
                    fig = px.scatter(
                        dot_data,
                        x=params["value_col"],
                        y=params["category_col"],
                        color=params["group_col"],
                        orientation="h" if params.get("orientation", "Horizontal") == "Horizontal" else "v",
                        title=custom_title if custom_title else "Dot Plot"
                    )
                else:
                    dot_data = df.groupby(params["category_col"])[params["value_col"]].mean().reset_index()
                    
                    fig = px.scatter(
                        dot_data,
                        x=params["value_col"],
                        y=params["category_col"],
                        orientation="h" if params.get("orientation", "Horizontal") == "Horizontal" else "v",
                        title=custom_title if custom_title else "Dot Plot"
                    )
                
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                
            except Exception as e:
                st.error(f"Error creating dot plot: {str(e)}")

        # Enhanced Traffic Light Chart
        elif chart_type == "traffic_light":
            # The layout and rendering are handled by this function
            render_traffic_light_layout(params, df, custom_title)
            return # This function handles its own rendering

        # Enhanced Choropleth Map
        elif chart_type == "choropleth":
            render_choropleth_chart(params, df, custom_title)
            return

        # Enhanced Donut Chart
        elif chart_type == "donut":
            fig = go.Figure()
            
            # Add multiple levels
            for level in range(params.get("levels", 1)):
                level_col = params.get(f"level_{level}_col")
                level_val = params.get(f"level_{level}_value")
                if not level_col or not level_val:
                    st.warning(f"Level {level+1} is not fully configured.")
                    continue

                level_df = df.groupby(level_col)[level_val].sum().reset_index()
                
                fig.add_trace(go.Pie(
                    values=level_df[level_val],
                    labels=level_df[level_col],
                    hole=params.get("hole_size", 0.4),
                    name=f"Level {level+1}",
                    domain={"x": [0.1 * level, 0.9 - 0.1 * level], "y": [0.1 * level, 0.9 - 0.1 * level]}
                ))
            
            # Add center metric if specified
            center_metric = params.get("center_metric")
            if center_metric:
                center_value = df[center_metric].mean()
                fig.add_annotation(
                    text=f"<b>Center KPI</b><br>{center_metric}<br>{center_value:.2f}",
                    x=0.5, y=0.5,
                    showarrow=False,
                    font_size=15
                )
            
            fig.update_traces(textposition='inside', textinfo='percent+label')
            fig.update_layout(title_text=custom_title)
            st.plotly_chart(fig, use_container_width=True)
            return        

        else:
            st.error(f"Unsupported chart type: {chart_type}")

    except Exception as e:
        st.error(f"Error rendering chart: {str(e)}")
        # Show a simplified version as fallback
        if chart_type == "kpi":
            metrics = params.get("metrics", [])
            if metrics:
                st.subheader(params.get("custom_title", "KPI Dashboard"))
                for metric in metrics:
                    value = df[metric].mean() if metric in df.columns else 0
                    st.metric(metric, f"{value:.2f}")



        # Get parameters with fallbacks
        levels = params.get("levels", 1)
        center_metric = params.get("center_metric")
        
        # Create the chart using the saved parameters
        # (implementation details would go here)
        
        return fig
    except Exception as e:
        return create_error_plot(f"Donut Chart Error: {str(e)}")



# Add this function to handle KPI layout rendering
def render_kpi_layout(metrics_data, layout_style, color_scheme):
    """Render KPI metrics in different layout styles"""
    
    if layout_style == "grid":
        # Grid layout - responsive columns
        cols = st.columns(min(4, len(metrics_data)))
        for idx, (metric, value) in enumerate(metrics_data.items()):
            with cols[idx % len(cols)]:
                st.metric(
                    label=metric,
                    value=f"{value:,.2f}",
                    help=f"{metric} value"
                )
    
    elif layout_style == "cards":
        # Cards layout - each metric in a styled card
        cols = st.columns(min(3, len(metrics_data)))
        for idx, (metric, value) in enumerate(metrics_data.items()):
            with cols[idx % len(cols)]:
                st.markdown(f"""
                <div style="
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border-radius: 12px;
                    padding: 1rem;
                    margin: 0.5rem 0;
                    text-align: center;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
                ">
                    <h3 style="margin: 0; font-size: 1rem; opacity: 0.9;">{metric}</h3>
                    <h2 style="margin: 0.5rem 0; font-size: 1.8rem; font-weight: 700;">{value:,.2f}</h2>
                </div>
                """, unsafe_allow_html=True)
    
    elif layout_style == "minimal":
        # Minimal layout - clean and simple
        for metric, value in metrics_data.items():
            st.metric(
                label=metric,
                value=f"{value:,.2f}",
                help=f"{metric} value"
            )
            st.markdown("---")
    
    elif layout_style == "gauges":
        # Gauges layout - circular progress indicators
        cols = st.columns(min(2, len(metrics_data)))
        for idx, (metric, value) in enumerate(metrics_data.items()):
            with cols[idx % len(cols)]:
                # Calculate percentage for gauge (assuming max is 2x current value for demo)
                max_val = value * 2 if value > 0 else 100
                percentage = min(100, (value / max_val) * 100)
                
                st.markdown(f"""
                <div style="
                    text-align: center;
                    margin: 1rem 0;
                    padding: 1rem;
                    background: rgba(255, 255, 255, 0.1);
                    border-radius: 12px;
                    border: 1px solid rgba(255, 255, 255, 0.2);
                ">
                    <h3 style="margin: 0 0 1rem 0; font-size: 1rem; color: #666;">{metric}</h3>
                    <div style="
                        position: relative;
                        width: 100px;
                        height: 100px;
                        margin: 0 auto;
                    ">
                        <div style="
                            position: absolute;
                            width: 100%;
                            height: 100%;
                            border-radius: 50%;
                            background: conic-gradient(#667eea {percentage}%, #e2e8f0 0%);
                        "></div>
                        <div style="
                            position: absolute;
                            top: 10px;
                            left: 10px;
                            width: 80px;
                            height: 80px;
                            border-radius: 50%;
                            background: white;
                            display: flex;
                            align-items: center;
                            justify-content: center;
                            font-weight: bold;
                            font-size: 1.2rem;
                        ">{value:,.0f}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
    
    else:
        # Default layout
        for metric, value in metrics_data.items():
            st.metric(
                label=metric,
                value=f"{value:,.2f}",
                help=f"{metric} value"
            )

def generate_traffic_light_html(params, df):
    """Generate HTML for the traffic light dashboard component for export."""
    kpi_metrics = params.get("kpi_metrics", [])
    layout_style = params.get("layout_style", "grid")
    show_sparklines = params.get("show_sparklines", True)
    show_values = params.get("show_values", True)

    if not kpi_metrics:
        return '<div class="error-card">No KPI metrics selected for Traffic Light chart.</div>'

    # Prepare data for each metric
    metrics_data = []
    for metric in kpi_metrics:
        value = df[metric].mean()
        red_thresh = params.get(f"{metric}_red_threshold", 0)
        yellow_thresh = params.get(f"{metric}_yellow_threshold", 0)
        
        if value <= red_thresh:
            color = "#FF4B4B" # Streamlit's error red
            status_text = "Critical"
        elif value <= yellow_thresh:
            color = "#FFC300" # A nice yellow
            status_text = "Warning"
        else:
            color = "#28A745" # Streamlit's success green
            status_text = "Good"
            
        metrics_data.append({
            "name": metric,
            "value": value,
            "color": color,
            "status": status_text,
            "sparkline_data": df[metric].dropna().tail(30).tolist()
        })

    # For circular layout, generate a Plotly figure
    if layout_style == "circular":
        fig = go.Figure()
        num_metrics = len(metrics_data)
        radius = 0.4
        
        for i, metric_data in enumerate(metrics_data):
            angle = (2 * np.pi * i / num_metrics) - (np.pi / 2)
            x_pos = 0.5 + radius * np.cos(angle)
            y_pos = 0.5 + radius * np.sin(angle)
            
            fig.add_trace(go.Scatter(
                x=[x_pos], y=[y_pos], mode='markers',
                marker=dict(color=metric_data['color'], size=100, line=dict(width=4, color='rgba(0,0,0,0.2)')),
                hoverinfo='text', text=f"<b>{metric_data['name']}</b><br>Value: {metric_data['value']:.2f}<br>Status: {metric_data['status']}"
            ))
            
            fig.add_annotation(
                x=x_pos, y=y_pos,
                text=f"<b>{metric_data['name']}</b><br>{metric_data['value']:.2f}" if show_values else f"<b>{metric_data['name']}</b>",
                showarrow=False, font=dict(color='white' if metric_data['color'] != '#FFC300' else 'black', size=14)
            )
        
        fig.update_layout(showlegend=False, xaxis=dict(visible=False, range=[0, 1]), yaxis=dict(visible=False, range=[0, 1], scaleanchor="x", scaleratio=1), plot_bgcolor='rgba(0,0,0,0)', height=600, margin=dict(l=10, r=10, t=40, b=10), title="Circular KPI Status")
        return fig.to_html(full_html=False, include_plotlyjs=False)

    # For dashboard layout, generate HTML/CSS with sparklines
    elif layout_style == "dashboard":
        num_cols = min(len(metrics_data), 3)
        html = f'<div class="traffic-light-grid" style="display: grid; grid-template-columns: repeat({num_cols}, 1fr); gap: 1rem;">'
        for metric_data in metrics_data:
            value_html = f'<div class="tl-dashboard-value-export">{metric_data["value"]:.2f}</div>' if show_values else ''
            sparkline_html = ""
            if show_sparklines:
                spark_fig = go.Figure(go.Scatter(
                    y=metric_data['sparkline_data'],
                    mode='lines',
                    line=dict(color=metric_data['color'], width=3),
                    fill='tozeroy',
                    fillcolor=metric_data['color'].replace(')', ', 0.2)').replace('rgb', 'rgba')
                ))
                spark_fig.update_layout(
                    height=80, showlegend=False,
                    xaxis=dict(visible=False), yaxis=dict(visible=False),
                    margin=dict(l=0, r=0, t=5, b=0), plot_bgcolor='rgba(0,0,0,0)'
                )
                sparkline_html = spark_fig.to_html(full_html=False, include_plotlyjs=False)

            html += f"""
            <div class="tl-dashboard-card-export">
                <div class="tl-dashboard-header-export">
                    <div class="tl-dashboard-title-export">{metric_data['name']}</div>
                    <div class="tl-dashboard-status-dot-export" style="background-color: {metric_data['color']};"></div>
                </div>
                {value_html}
                {sparkline_html}
            </div>
            """
        html += '</div>'
        # Add CSS for dashboard style
        html += """
        <style>
        .tl-dashboard-card-export { padding: 1rem; border-radius: 12px; background-color: #ffffff; border: 1px solid #e9ecef; box-shadow: 0 4px 12px rgba(0,0,0,0.05); }
        .tl-dashboard-header-export { display: flex; align-items: center; justify-content: space-between; margin-bottom: 1rem; }
        .tl-dashboard-title-export { font-size: 1.1rem; font-weight: 600; color: #31333F; }
        .tl-dashboard-status-dot-export { width: 12px; height: 12px; border-radius: 50%; }
        .tl-dashboard-value-export { font-size: 2.2rem; font-weight: 700; color: #111; text-align: center; }
        </style>
        """
        return html

    # For grid/linear layouts, generate HTML/CSS
    else:
        num_cols = len(metrics_data) if layout_style == "linear" else min(len(metrics_data), 4)
        html = f'<div class="traffic-light-grid" style="display: grid; grid-template-columns: repeat({num_cols}, 1fr); gap: 1rem;">'
        for metric_data in metrics_data:
            value_html = f'<div class="traffic-light-value-export">{metric_data["value"]:.2f}</div>' if show_values else ''
            status_color = metric_data['color'].replace(')', ', 0.2)').replace('rgb', 'rgba')
            status_text_color = metric_data['color']
            html += f"""
            <div class="traffic-light-container-export">
                <div class="traffic-light-circle-export" style="background-color: {metric_data['color']};"></div>
                <div class="traffic-light-label-export">{metric_data['name']}</div>
                {value_html}
                <div class="traffic-light-status-export" style="background-color: {status_color}; color: {status_text_color}; border: 1px solid {status_text_color}; padding: 0.2rem 0.6rem; border-radius: 15px; margin-top: 0.75rem;">{metric_data['status']}</div>
            </div>
            """
        html += '</div>'
        return html
                

def render_dashboard(dashboard, df):
    """Render the dashboard with enhanced layout and controls, with reorder buttons."""
    
    if dashboard and dashboard.get('charts'):
        charts = dashboard['charts']
        num_cols = dashboard.get('config', {}).get('num_columns', 2)
        num_rows = (len(charts) + num_cols - 1) // num_cols
        
        # Check if reorder mode is enabled
        reorder_mode = st.session_state.get('reorder_mode', False)
        
        st.markdown("### 📊 Dashboard")
        
        if reorder_mode:
            st.info("🔀 **Reorder Mode Active** - Use the arrow buttons to rearrange your charts")
        
        # Create the grid using st.columns in a loop for rows
        for row_idx in range(num_rows):
            start_idx = row_idx * num_cols
            end_idx = min(start_idx + num_cols, len(charts))
            row_charts = charts[start_idx:end_idx]
            
            # Create columns for this row
            cols = st.columns(num_cols)
            
            # Populate columns with charts
            for col_idx, chart_config in enumerate(row_charts):
                chart_index = start_idx + col_idx
                with cols[col_idx]:
                    # Add a container for the chart
                    with st.container():
                        # Create buttons for reordering
                        if reorder_mode:
                            st.markdown(f"**Chart {chart_index + 1}**")
                            btn_col1, btn_col2, btn_col3, btn_col4 = st.columns(4)
                            
                            with btn_col1:
                                if st.button("↑", key=f"up_{chart_index}",
                                           disabled=(chart_index == 0),
                                           help="Move up"):
                                    move_chart_up(chart_index)
                            
                            with btn_col2:
                                if st.button("↓", key=f"down_{chart_index}",
                                           disabled=(chart_index == len(charts)-1),
                                           help="Move down"):
                                    move_chart_down(chart_index)
                            
                            with btn_col3:
                                if st.button("←", key=f"left_{chart_index}",
                                           disabled=(chart_index % num_cols == 0),
                                           help="Move left"):
                                    move_chart_left(chart_index, num_cols)
                            
                            with btn_col4:
                                if st.button("→", key=f"right_{chart_index}",
                                           disabled=(chart_index % num_cols == num_cols-1 or chart_index == len(charts)-1),
                                           help="Move right"):
                                    move_chart_right(chart_index, num_cols)
                            
                            st.markdown("---")
                        
                        # Display chart title with enhanced styling
                        chart_title = (chart_config.get("title") or
                                     chart_config.get("params", {}).get("custom_title", "") or
                                     chart_config['type'].replace('_', ' ').title())
                        
                        st.subheader(chart_title, divider='gray')
                        
                        # Display chart description if available
                        chart_description = (chart_config.get("description") or
                                           chart_config.get("params", {}).get("custom_description", ""))
                        
                        if chart_description:
                            st.markdown(f"""
                            <div style="
                                background: rgba(255, 255, 255, 0.1);
                                border-left: 4px solid #667eea;
                                padding: 0.75rem 1rem;
                                margin: 0.5rem 0 1rem 0;
                                border-radius: 0 8px 8px 0;
                                font-style: italic;
                                color: #4a5568;
                                backdrop-filter: blur(5px);
                            ">
                                <small>📄 {chart_description}</small>
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # Render the chart
                        render_chart_with_persistence(chart_config, df)
    elif not st.session_state.get('dashboard_edit_mode', False):
        st.info("No charts configured yet. Add some charts to see your dashboard!")



# ===== AI DASHBOARD DESIGN =====
st.sidebar.header("🤖 AI Dashboard Design")
# Use the global session variables directly
api_key = st.session_state.global_api_key
ai_provider = st.session_state.global_ai_provider
model_name = st.session_state.global_model_name

# Use optimized LLM initialization - only show status when API key is provided
if st.session_state.global_api_key:
    try:
        llm = get_validated_dashboard_llm()
        st.sidebar.success(f"✅ AI Ready ({ai_provider})")
    except Exception as e:
        st.sidebar.error(f"❌ AI Setup Error: {str(e)}")
        llm = None
else:
    # Don't show error by default - only when user tries to use AI
    llm = None

# AI prompt input
ai_prompt = st.sidebar.text_area(
    "💬 Describe the dashboard you want:",
    placeholder="e.g., 'Show sales trends by region over time, compare product categories, and highlight top performers'"
)

def parse_chart_requirements(prompt):
    """Parse user request for specific chart count and types"""
    requirements = {
        'chart_count': None,
        'specific_types': [],
        'chart_count_specified': False
    }
    
    prompt_lower = prompt.lower()
    
    # Detect specific chart count requests
    import re
    
    # Look for numbers followed by chart-related words
    count_patterns = [
        r'(\d+)\s+charts?',
        r'(\d+)\s+visualizations?',
        r'(\d+)\s+graphs?',
        r'create\s+(\d+)',
        r'make\s+(\d+)',
        r'show\s+me\s+(\d+)'
    ]
    
    for pattern in count_patterns:
        match = re.search(pattern, prompt_lower)
        if match:
            requirements['chart_count'] = int(match.group(1))
            requirements['chart_count_specified'] = True
            break
    
    # Detect specific chart types mentioned
    chart_type_keywords = {
        'bar': ['bar chart', 'bar graph', 'column chart'],
        'line': ['line chart', 'line graph', 'trend chart'],
        'scatter': ['scatter plot', 'scatter chart', 'correlation plot'],
        'pie': ['pie chart', 'pie graph', 'donut chart'],
        'histogram': ['histogram', 'distribution chart', 'frequency chart'],
        'box': ['box plot', 'box chart', 'boxplot']
    }
    
    for chart_type, keywords in chart_type_keywords.items():
        for keyword in keywords:
            if keyword in prompt_lower:
                requirements['specific_types'].append(chart_type)
                break
    
    # Remove duplicates
    requirements['specific_types'] = list(set(requirements['specific_types']))
    
    return requirements

def analyze_user_request(prompt, df):
    """Analyze user request to provide better context for AI chart generation"""
    analysis = []
    
    # Convert prompt to lowercase for analysis
    prompt_lower = prompt.lower()
    
    # Detect request type
    if any(word in prompt_lower for word in ['trend', 'over time', 'timeline', 'change', 'growth']):
        analysis.append("- User wants to see TRENDS or TIME-BASED analysis")
        if df is not None:
            date_cols = df.select_dtypes(include=['datetime64', 'object']).columns
            for col in date_cols:
                if any(date_word in col.lower() for date_word in ['date', 'time', 'year', 'month']):
                    analysis.append(f"- Suggested time column: {col}")
    
    if any(word in prompt_lower for word in ['compare', 'comparison', 'vs', 'versus', 'difference']):
        analysis.append("- User wants to COMPARE different categories or groups")
        if df is not None:
            cat_cols = df.select_dtypes(include=['object', 'category']).columns[:3]
            analysis.append(f"- Suggested grouping columns: {', '.join(cat_cols)}")
    
    if any(word in prompt_lower for word in ['relationship', 'correlation', 'related', 'connection']):
        analysis.append("- User wants to see RELATIONSHIPS between variables")
        if df is not None:
            num_cols = df.select_dtypes(include=['number']).columns[:4]
            analysis.append(f"- Suggested numeric columns for correlation: {', '.join(num_cols)}")
    
    if any(word in prompt_lower for word in ['distribution', 'spread', 'histogram', 'frequency']):
        analysis.append("- User wants to see DATA DISTRIBUTION")
        if df is not None:
            num_cols = df.select_dtypes(include=['number']).columns[:2]
            analysis.append(f"- Suggested columns for distribution: {', '.join(num_cols)}")
    
    if any(word in prompt_lower for word in ['top', 'bottom', 'highest', 'lowest', 'best', 'worst']):
        analysis.append("- User wants to see RANKINGS or TOP/BOTTOM performers")
        if df is not None:
            num_cols = df.select_dtypes(include=['number']).columns[:2]
            cat_cols = df.select_dtypes(include=['object', 'category']).columns[:2]
            analysis.append(f"- Suggested ranking columns: {', '.join(num_cols)}")
            analysis.append(f"- Suggested category columns: {', '.join(cat_cols)}")
    
    if any(word in prompt_lower for word in ['proportion', 'percentage', 'share', 'breakdown']):
        analysis.append("- User wants to see PROPORTIONS or BREAKDOWNS")
        if df is not None:
            cat_cols = df.select_dtypes(include=['object', 'category']).columns[:2]
            analysis.append(f"- Suggested categorical columns: {', '.join(cat_cols)}")
    
    if any(word in prompt_lower for word in ['summary', 'overview', 'general', 'all']):
        analysis.append("- User wants a COMPREHENSIVE OVERVIEW")
        analysis.append("- Suggest multiple chart types for complete picture")
    
    # Detect specific column mentions
    if df is not None:
        for col in df.columns:
            if col.lower() in prompt_lower:
                analysis.append(f"- User specifically mentioned column: {col}")
    
    # Default analysis if nothing specific detected
    if not analysis:
        analysis.append("- General dashboard request - create diverse visualizations")
        if df is not None:
            analysis.append(f"- Available numeric columns: {', '.join(df.select_dtypes(include=['number']).columns[:3])}")
            analysis.append(f"- Available categorical columns: {', '.join(df.select_dtypes(include=['object', 'category']).columns[:3])}")
    
    return '\n'.join(analysis)

def generate_comprehensive_dashboard_analysis(dashboard, df, chat_history=None):
    """Generate comprehensive analysis of the dashboard with industry-standard insights"""
    try:
        # Get AI settings from session state
        ai_provider = st.session_state.get('global_ai_provider', 'DeepSeek')
        model_name = st.session_state.get('global_model_name', 'deepseek-chat')
        api_key = st.session_state.get('global_api_key', '')
        
        if not api_key:
            return None
        
        # Import required modules
        from ai_engine import get_model_context_protocol
        from langchain_openai import ChatOpenAI
        from langchain_anthropic import ChatAnthropic
        
        # Get context protocol
        context_protocol = get_model_context_protocol(ai_provider, model_name)
        
        # Initialize LLM based on provider
        if ai_provider == "DeepSeek":
            llm = ChatOpenAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                base_url="https://api.deepseek.com/v1",
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "OpenAI":
            llm = ChatOpenAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "Anthropic":
            llm = ChatAnthropic(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "Google":
            from langchain_google_genai import ChatGoogleGenerativeAI
            llm = ChatGoogleGenerativeAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                google_api_key=api_key,
                max_output_tokens=context_protocol["max_tokens"]
            )
        else:
            return None
        
        # Analyze dashboard structure
        charts = dashboard.get('charts', [])
        dashboard_title = dashboard.get('title', 'Dashboard Analysis')
        dashboard_description = dashboard.get('description', '')
        
        # Analyze data context
        data_context = ""
        if df is not None:
            numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
            categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
            data_context = f"""
DATA CONTEXT:
- Dataset: {df.shape[0]} rows × {df.shape[1]} columns
- Numeric columns: {', '.join(numeric_cols[:5])}{'...' if len(numeric_cols) > 5 else ''}
- Categorical columns: {', '.join(categorical_cols[:5])}{'...' if len(categorical_cols) > 5 else ''}
- Data types: {dict(df.dtypes.value_counts())}
"""
        
        # Analyze each chart
        chart_analysis = []
        for i, chart in enumerate(charts, 1):
            chart_type = chart.get('type', 'unknown')
            chart_title = chart.get('title', f'Chart {i}')
            chart_params = chart.get('params', {})
            
            chart_info = f"""
CHART {i}: {chart_title}
- Type: {chart_type}
- Parameters: {chart_params}
- Purpose: {chart.get('description', 'No description')}
"""
            chart_analysis.append(chart_info)
        
        # Create comprehensive analysis prompt
        analysis_prompt = f"""
DASHBOARD ANALYSIS EXPERT

You are a senior business intelligence analyst with expertise in data visualization and dashboard design. Provide a comprehensive, professional analysis of this dashboard.

DASHBOARD OVERVIEW:
Title: {dashboard_title}
Description: {dashboard_description}
Number of Charts: {len(charts)}

{data_context}

CHART DETAILS:
{''.join(chart_analysis)}

PROVIDE A COMPREHENSIVE ANALYSIS INCLUDING:

## 📊 Executive Summary
Provide a high-level overview of what this dashboard accomplishes and its business value.

## 🔍 Chart-by-Chart Analysis
For each chart, provide:
- **Business Purpose**: What business question does this chart answer?
- **Industry Insights**: What industry-standard insights does this visualization type typically reveal?
- **Key Findings**: What patterns, trends, or insights are likely visible?
- **Actionable Recommendations**: What business actions could be taken based on this chart?

## 📈 Dashboard Design Assessment
- **Visualization Effectiveness**: How well do the chosen chart types serve their purpose?
- **Data Story**: How do the charts work together to tell a cohesive story?
- **Industry Best Practices**: How does this dashboard align with BI best practices?

## 🎯 Strategic Recommendations
- **Optimization Opportunities**: How could this dashboard be improved?
- **Additional Metrics**: What other KPIs or visualizations would enhance insights?
- **Business Impact**: How can stakeholders use this dashboard for decision-making?

## 📋 Technical Assessment
- **Data Quality**: Assessment of data completeness and reliability
- **Scalability**: How well will this dashboard perform with larger datasets?
- **Maintenance**: Recommendations for keeping the dashboard current and relevant

Provide detailed, professional insights that would be valuable to business stakeholders, data analysts, and decision-makers. Use industry terminology and best practices throughout your analysis.

Format your response in clear markdown with proper headings and bullet points for easy reading.
"""
        
        # Generate comprehensive analysis
        response = llm.invoke(analysis_prompt)
        return response.content
        
    except Exception as e:
        print(f"Error generating dashboard analysis: {e}")
        return None

def create_enhanced_dashboard_assistant(dashboard, df):
    """Create enhanced AI dashboard assistant with full dashboard context"""
    
    # Initialize chat history if not exists
    if 'dashboard_chat_history' not in st.session_state:
        st.session_state.dashboard_chat_history = []
    
    # Dashboard context summary for AI
    dashboard_context = ""
    if dashboard and df is not None:
        charts = dashboard.get('charts', [])
        dashboard_context = f"""
DASHBOARD CONTEXT:
- Dashboard Title: {dashboard.get('title', 'Untitled Dashboard')}
- Number of Charts: {len(charts)}
- Dataset: {df.shape[0]} rows × {df.shape[1]} columns
- Available Columns: {', '.join(df.columns.tolist())}
- Chart Types: {', '.join([chart.get('type', 'unknown') for chart in charts])}

CHART DETAILS:
"""
        for i, chart in enumerate(charts, 1):
            chart_context = f"""
Chart {i}: {chart.get('title', f'Chart {i}')}
- Type: {chart.get('type', 'unknown')}
- Parameters: {chart.get('params', {})}
- Description: {chart.get('description', 'No description')}
"""
            dashboard_context += chart_context
    
    # Smart Questions organized by category
    if dashboard and dashboard.get('charts'):
        smart_questions = get_smart_dashboard_questions()
        
        # Enhanced question categories with dashboard context
        dashboard_specific_questions = [
            f"Analyze the {len(dashboard.get('charts', []))} charts in my dashboard",
            "What insights can you derive from my dashboard?",
            "How do the charts in my dashboard work together?",
            "What business recommendations can you make from my dashboard?",
            "Identify patterns and trends across my dashboard charts",
            "What additional charts would complement my current dashboard?",
            "Generate a comprehensive analysis of my dashboard",
            "What story does my dashboard tell about the data?"
        ]
        
        chart_specific_questions = []
        for i, chart in enumerate(dashboard.get('charts', []), 1):
            chart_type = chart.get('type', 'unknown')
            chart_title = chart.get('title', f'Chart {i}')
            chart_specific_questions.extend([
                f"Analyze Chart {i}: {chart_title} ({chart_type})",
                f"What insights does the {chart_type} chart show?",
                f"How should I interpret Chart {i}?"
            ])
        
        # Display questions in organized sections
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 🎯 Dashboard Analysis")
            for i, question in enumerate(dashboard_specific_questions[:4]):
                if st.button(question, key=f"dash_q_{i}", use_container_width=True):
                    with st.spinner("🤖 Analyzing your dashboard..."):
                        # Enhanced response with full dashboard context
                        enhanced_question = f"{question}\n\nDashboard Context:\n{dashboard_context}"
                        response = generate_optimized_dashboard_response(enhanced_question)
                        st.session_state.dashboard_chat_history.append({"role": "user", "content": question})
                        st.session_state.dashboard_chat_history.append({"role": "assistant", "content": response})
                        
                        # Display response in main area
                        st.markdown("#### 🤖 AI Analysis:")
                        st.markdown(response)
        
        with col2:
            st.markdown("### 📊 Chart-Specific Questions")
            for i, question in enumerate(chart_specific_questions[:4]):
                if st.button(question, key=f"chart_spec_q_{i}", use_container_width=True):
                    with st.spinner("🤖 Analyzing chart..."):
                        # Enhanced response with chart context
                        enhanced_question = f"{question}\n\nDashboard Context:\n{dashboard_context}"
                        response = generate_optimized_dashboard_response(enhanced_question)
                        st.session_state.dashboard_chat_history.append({"role": "user", "content": question})
                        st.session_state.dashboard_chat_history.append({"role": "assistant", "content": response})
                        
                        # Display response in main area
                        st.markdown("#### 🤖 AI Analysis:")
                        st.markdown(response)
        
        # More questions in expanders
        if len(dashboard_specific_questions) > 4:
            with st.expander("🔬 More Dashboard Analysis Questions"):
                for i, question in enumerate(dashboard_specific_questions[4:]):
                    if st.button(question, key=f"more_dash_q_{i}", use_container_width=True):
                        with st.spinner("🤖 Analyzing..."):
                            enhanced_question = f"{question}\n\nDashboard Context:\n{dashboard_context}"
                            response = generate_optimized_dashboard_response(enhanced_question)
                            st.session_state.dashboard_chat_history.append({"role": "user", "content": question})
                            st.session_state.dashboard_chat_history.append({"role": "assistant", "content": response})
                            st.markdown("#### 🤖 AI Analysis:")
                            st.markdown(response)
        
        if len(chart_specific_questions) > 4:
            with st.expander("📈 More Chart-Specific Questions"):
                for i, question in enumerate(chart_specific_questions[4:]):
                    if st.button(question, key=f"more_chart_q_{i}", use_container_width=True):
                        with st.spinner("🤖 Analyzing..."):
                            enhanced_question = f"{question}\n\nDashboard Context:\n{dashboard_context}"
                            response = generate_optimized_dashboard_response(enhanced_question)
                            st.session_state.dashboard_chat_history.append({"role": "user", "content": question})
                            st.session_state.dashboard_chat_history.append({"role": "assistant", "content": response})
                            st.markdown("#### 🤖 AI Analysis:")
                            st.markdown(response)
    
    # Custom question input with dashboard context
    st.markdown("### ✍️ Ask Your Own Question")
    custom_question = st.text_area(
        "Ask anything about your dashboard:",
        height=100,
        placeholder="e.g., What trends do you see in my data? How can I improve my dashboard? What insights am I missing?",
        key="dashboard_custom_question_main",
        help="The AI has full access to your dashboard configuration, chart details, and data context."
    )
    
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("🤖 Ask AI Assistant", key="ask_dashboard_ai_main", use_container_width=True, type="primary"):
            if custom_question:
                with st.spinner("🤖 AI is analyzing your dashboard and question..."):
                    # Enhanced response with full dashboard context
                    enhanced_question = f"{custom_question}\n\nDashboard Context:\n{dashboard_context}"
                    response = generate_optimized_dashboard_response(enhanced_question)
                    st.session_state.dashboard_chat_history.append({"role": "user", "content": custom_question})
                    st.session_state.dashboard_chat_history.append({"role": "assistant", "content": response})
                    
                    # Display response in main area
                    st.markdown("#### 🤖 AI Response:")
                    st.markdown(response)
            else:
                st.warning("Please enter a question about your dashboard")
    
    with col2:
        if st.button("🗑️ Clear Chat", key="clear_dashboard_chat", use_container_width=True):
            st.session_state.dashboard_chat_history = []
            st.success("Chat history cleared!")
            st.rerun()
    
    # Show conversation history in main area
    if st.session_state.dashboard_chat_history:
        st.markdown("### 💬 Conversation History")
        with st.expander("View Chat History", expanded=False):
            for i, msg in enumerate(st.session_state.dashboard_chat_history[-10:]):  # Last 5 exchanges
                if msg['role'] == 'user':
                    st.markdown(f"**🙋 You:** {msg['content']}")
                else:
                    st.markdown(f"**🤖 AI:** {msg['content']}")
                if i < len(st.session_state.dashboard_chat_history[-10:]) - 1:
                    st.markdown("---")

# Generate with AI button - moved to top of metadata section
if st.sidebar.button("🚀 Generate with AI", use_container_width=True):
    # Check for API key only when user tries to use AI
    # Define the AI dashboard creation function here before using it
    def create_optimized_ai_dashboard_local(prompt, df, chat_history=None):
        """Create an AI-generated dashboard configuration based on user prompt and data"""
        """Create an AI-generated dashboard configuration based on user prompt and data - ENHANCED"""
        try:
            # Get AI settings from session state
            ai_provider = st.session_state.get('global_ai_provider', 'DeepSeek')
            model_name = st.session_state.get('global_model_name', 'deepseek-chat')
            api_key = st.session_state.get('global_api_key', '')
            
            if not api_key:
                return None
            
            # Import required modules
            from ai_engine import get_model_context_protocol
            from langchain_openai import ChatOpenAI
            from langchain_anthropic import ChatAnthropic
            import json
            import re
            
            # Get context protocol
            context_protocol = get_model_context_protocol(ai_provider, model_name)
            
            # Initialize LLM based on provider
            if ai_provider == "DeepSeek":
                llm = ChatOpenAI(
                    temperature=0.1,  # Lower temperature for more consistent results
                    model=model_name,
                    api_key=api_key,
                    base_url="https://api.deepseek.com/v1",
                    max_tokens=context_protocol["max_tokens"]
                )
            elif ai_provider == "OpenAI":
                llm = ChatOpenAI(
                    temperature=0.1,
                    model=model_name,
                    api_key=api_key,
                    max_tokens=context_protocol["max_tokens"]
                )
            elif ai_provider == "Anthropic":
                llm = ChatAnthropic(
                    temperature=0.1,
                    model=model_name,
                    api_key=api_key,
                    max_tokens=context_protocol["max_tokens"]
                )
            elif ai_provider == "Google":
                from langchain_google_genai import ChatGoogleGenerativeAI
                llm = ChatGoogleGenerativeAI(
                    temperature=0.1,
                    model=model_name,
                    google_api_key=api_key,
                    max_output_tokens=context_protocol["max_tokens"]
                )
            else:
                return None
            
            # Get comprehensive chart knowledge
            chart_knowledge = get_comprehensive_chart_knowledge()
            chart_knowledge_text = "\n".join([f"- **{info['name']} ({chart_type})** {info['icon']}: {info['description']}" for chart_type, info in chart_knowledge.items()])
            
            # Analyze the dataset comprehensively
            if df is not None:
                numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
                categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
                datetime_cols = df.select_dtypes(include=['datetime64']).columns.tolist()
                
                # Get data sample for AI context
                data_sample = df.head(3).to_dict('records')
                data_info = {
                    'shape': df.shape,
                    'numeric_columns': numeric_cols,
                    'categorical_columns': categorical_cols,
                    'datetime_columns': datetime_cols,
                    'sample_data': data_sample,
                    'column_info': {}
                }
                
                # Add detailed column information
                for col in df.columns:
                    col_info = {
                        'type': str(df[col].dtype),
                        'unique_values': df[col].nunique(),
                        'null_count': df[col].isnull().sum()
                    }
                    
                    if col in numeric_cols:
                        col_info.update({
                            'min': float(df[col].min()),
                            'max': float(df[col].max()),
                            'mean': float(df[col].mean())
                        })
                    elif col in categorical_cols:
                        col_info['top_values'] = df[col].value_counts().head(5).to_dict()
                    
                    data_info['column_info'][col] = col_info
                
                # Create comprehensive AI prompt
                enhanced_prompt = f"""You are an expert data visualization specialist. Create a dashboard configuration that EXACTLY matches the user's request.

USER REQUEST: "{prompt}"

DATASET INFORMATION:
- Shape: {data_info['shape'][0]} rows, {data_info['shape'][1]} columns
- Numeric columns: {', '.join(numeric_cols) if numeric_cols else 'None'}
- Categorical columns: {', '.join(categorical_cols) if categorical_cols else 'None'}
- DateTime columns: {', '.join(datetime_cols) if datetime_cols else 'None'}

COLUMN DETAILS:
{json.dumps(data_info['column_info'], indent=2)}

SAMPLE DATA:
{json.dumps(data_sample, indent=2)}

CHART TYPE MAPPING:
- bar: Compare categories (x_col=categorical, y_col=numeric)
- line: Show trends over time (x_col=datetime/categorical, y_col=numeric)
- scatter: Show relationships (x_col=numeric, y_col=numeric, color_col=categorical optional)
- pie: Show proportions (column=categorical, value_col=numeric optional)
- histogram: Show distribution (column=numeric, bins=number)
- box: Show statistical summary (x_col=categorical optional, y_col=numeric)
- treemap: Show hierarchical data (labels_col=categorical, values_col=numeric, parent_col=categorical optional)

INSTRUCTIONS:
1. Analyze the user request carefully to understand their intent
2. Choose the most appropriate chart type based on the request and available data
3. Use EXACT column names from the dataset
4. Create meaningful titles that reflect the user's specific request
5. Ensure the chart configuration is technically valid

RESPONSE FORMAT (JSON only, no additional text):
{{
    "title": "Dashboard title reflecting user request",
    "description": "Brief description of insights this dashboard provides",
    "charts": [
        {{
            "id": "chart1",
            "title": "Specific chart title based on user request",
            "type": "appropriate_chart_type",
            "params": {{
                "x_col": "exact_column_name_from_data",
                "y_col": "exact_column_name_from_data",
                "color_col": "optional_grouping_column"
            }},
            "description": "What specific insights this chart provides"
        }}
    ],
    "kpis": []
}}

CRITICAL REQUIREMENTS:
- Use ONLY column names that exist in the dataset
- Choose chart type that makes sense for the data types
- Create titles that directly address the user's request
- Return ONLY valid JSON, no explanations or additional text

USER REQUEST ANALYSIS:
- If user mentions specific columns, use those exact columns
- If user mentions chart type, use that type if appropriate for the data
- If user asks for comparisons, use bar/line charts
- If user asks for distributions, use histogram/box plots
- If user asks for relationships, use scatter plots
- If user asks for proportions, use pie charts"""

                # Generate AI response
                try:
                    response = llm.invoke(enhanced_prompt)
                    ai_response = response.content.strip()
                    
                    # Clean the response to extract JSON
                    if '```json' in ai_response:
                        ai_response = ai_response.split('```json')[1].split('```')[0].strip()
                    elif '```' in ai_response:
                        ai_response = ai_response.split('```')[1].strip()
                    
                    # Parse JSON response
                    dashboard_config = json.loads(ai_response)
                    
                    # Validate and fix the configuration
                    if 'charts' in dashboard_config:
                        for chart in dashboard_config['charts']:
                            # Ensure required fields exist
                            if 'params' not in chart:
                                chart['params'] = {}
                            
                            # Validate column names exist in dataset
                            params = chart['params']
                            if 'x_col' in params and params['x_col'] not in df.columns:
                                # Try to find a similar column or use first appropriate column
                                if chart['type'] == 'bar' and categorical_cols:
                                    params['x_col'] = categorical_cols[0]
                                elif chart['type'] in ['line', 'scatter'] and numeric_cols:
                                    params['x_col'] = numeric_cols[0]
                            
                            if 'y_col' in params and params['y_col'] not in df.columns:
                                if numeric_cols:
                                    params['y_col'] = numeric_cols[0]
                            
                            if 'column' in params and params['column'] not in df.columns:
                                if chart['type'] == 'histogram' and numeric_cols:
                                    params['column'] = numeric_cols[0]
                                elif chart['type'] == 'pie' and categorical_cols:
                                    params['column'] = categorical_cols[0]
                            
                            # Set default bins for histogram
                            if chart['type'] == 'histogram' and 'bins' not in params:
                                params['bins'] = 20
                    
                    # Ensure required top-level fields
                    if 'title' not in dashboard_config:
                        dashboard_config['title'] = f"AI Dashboard: {prompt[:50]}..."
                    if 'description' not in dashboard_config:
                        dashboard_config['description'] = "Dashboard created with AI assistance"
                    if 'kpis' not in dashboard_config:
                        dashboard_config['kpis'] = []
                    
                    return dashboard_config
                    
                except json.JSONDecodeError as e:
                    print(f"JSON parsing error: {e}")
                    print(f"AI Response: {ai_response[:200]}...")
                    # Fall through to fallback
                except Exception as e:
                    print(f"AI generation error: {e}")
                    # Fall through to fallback
            
            # Intelligent fallback based on user prompt analysis
            return create_intelligent_fallback_dashboard(prompt, df)
            
        except Exception as e:
            print(f"Error in AI dashboard creation: {e}")
            return create_intelligent_fallback_dashboard(prompt, df)

    def create_intelligent_fallback_dashboard(prompt, df):
        """Create intelligent fallback dashboard based on prompt analysis"""
        if df is None:
            return None
        
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        datetime_cols = df.select_dtypes(include=['datetime64']).columns.tolist()
        
        # Analyze user prompt for intent
        prompt_lower = prompt.lower()
        
        # Determine chart type based on user request
        chart_type = "bar"  # default
        if any(word in prompt_lower for word in ['histogram', 'distribution', 'spread']):
            chart_type = "histogram"
        elif any(word in prompt_lower for word in ['scatter', 'correlation', 'relationship']):
            chart_type = "scatter"
        elif any(word in prompt_lower for word in ['pie', 'proportion', 'percentage', 'share']):
            chart_type = "pie"
        elif any(word in prompt_lower for word in ['line', 'trend', 'over time', 'timeline']):
            chart_type = "line"
        elif any(word in prompt_lower for word in ['box', 'boxplot', 'quartile']):
            chart_type = "box"
        
        # Extract column names mentioned in prompt
        mentioned_cols = []
        for col in df.columns:
            if col.lower() in prompt_lower or col.replace('_', ' ').lower() in prompt_lower:
                mentioned_cols.append(col)
        
        # Create chart configuration based on analysis
        chart_config = {
            "id": "chart1",
            "title": f"Analysis: {prompt[:50]}..." if len(prompt) > 50 else prompt,
            "type": chart_type,
            "params": {},
            "description": f"Visualization based on your request: {prompt[:100]}..."
        }
        
        # Configure parameters based on chart type and available data
        if chart_type == "bar":
            if mentioned_cols:
                # Use mentioned columns if available
                cat_mentioned = [col for col in mentioned_cols if col in categorical_cols]
                num_mentioned = [col for col in mentioned_cols if col in numeric_cols]
                chart_config["params"]["x_col"] = cat_mentioned[0] if cat_mentioned else (categorical_cols[0] if categorical_cols else df.columns[0])
                chart_config["params"]["y_col"] = num_mentioned[0] if num_mentioned else (numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0])
            else:
                chart_config["params"]["x_col"] = categorical_cols[0] if categorical_cols else df.columns[0]
                chart_config["params"]["y_col"] = numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0]
        
        elif chart_type == "histogram":
            target_col = mentioned_cols[0] if mentioned_cols and mentioned_cols[0] in numeric_cols else (numeric_cols[0] if numeric_cols else df.columns[0])
            chart_config["params"]["column"] = target_col
            chart_config["params"]["bins"] = 20
        
        elif chart_type == "scatter":
            if len(mentioned_cols) >= 2:
                chart_config["params"]["x_col"] = mentioned_cols[0]
                chart_config["params"]["y_col"] = mentioned_cols[1]
            else:
                chart_config["params"]["x_col"] = numeric_cols[0] if len(numeric_cols) > 0 else df.columns[0]
                chart_config["params"]["y_col"] = numeric_cols[1] if len(numeric_cols) > 1 else (numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0])
            
            if categorical_cols:
                chart_config["params"]["color_col"] = categorical_cols[0]
        
        elif chart_type == "pie":
            target_col = mentioned_cols[0] if mentioned_cols and mentioned_cols[0] in categorical_cols else (categorical_cols[0] if categorical_cols else df.columns[0])
            chart_config["params"]["column"] = target_col
        
        elif chart_type == "line":
            x_col = datetime_cols[0] if datetime_cols else (categorical_cols[0] if categorical_cols else df.columns[0])
            y_col = mentioned_cols[0] if mentioned_cols and mentioned_cols[0] in numeric_cols else (numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0])
            chart_config["params"]["x_col"] = x_col
            chart_config["params"]["y_col"] = y_col
        
        elif chart_type == "box":
            chart_config["params"]["y_col"] = numeric_cols[0] if numeric_cols else df.columns[0]
            if categorical_cols:
                chart_config["params"]["x_col"] = categorical_cols[0]
        
        return {
            "title": f"Dashboard: {prompt[:30]}..." if len(prompt) > 30 else f"Dashboard: {prompt}",
            "description": f"Intelligent analysis based on your request: {prompt}",
            "charts": [chart_config],
            "kpis": []
        }

    if 'global_api_key' not in st.session_state or not st.session_state.global_api_key:
        st.sidebar.error("❌ API key not set! Please configure it in the Analysis Agent page.")
    else:
        with st.spinner("🎨 Designing dashboard with AI..."):
            # Call optimized AI to generate dashboard configuration
            # Call local, enhanced AI to generate dashboard configuration
            dashboard = create_optimized_ai_dashboard_local(
                ai_prompt, 
                df_sample, 
                st.session_state.dashboard_chat_history
            )
            if dashboard:
                st.session_state.dashboard = dashboard
                st.session_state.show_dashboard = True
                st.session_state.dashboard_edit_mode = False
                st.success("✅ Dashboard generated successfully!")
                st.rerun()
            else:
                st.sidebar.error("❌ Failed to generate AI dashboard")

# Chart Metadata Management
if st.session_state.configured_charts:
    if 'configured_charts' in st.session_state and st.session_state.configured_charts:
        st.sidebar.markdown("---")
        st.sidebar.subheader("📝 Chart Metadata Manager")
        
        with st.sidebar.expander("🏷️ Bulk Edit Titles & Descriptions", expanded=False):
            st.markdown("**Quick edit chart metadata:**")
            
            for i, chart in enumerate(st.session_state.configured_charts):
                chart_type = chart.get('type', 'unknown').replace('_', ' ').title()
                current_title = chart.get('params', {}).get('custom_title', '')
                current_desc = chart.get('params', {}).get('custom_description', '')
                
                st.markdown(f"**Chart {i+1}: {chart_type}**")
                
                # Title input
                new_title = st.text_input(
                    f"Title for Chart {i+1}",
                    value=current_title,
                    key=f"bulk_title_{i}",
                    placeholder="Enter chart title..."
                )
                
                # Description input
                new_desc = st.text_area(
                    f"Description for Chart {i+1}",
                    value=current_desc,
                    key=f"bulk_desc_{i}",
                    placeholder="Enter chart description...",
                    height=80
                )
                
                # Update if changed
                if new_title != current_title or new_desc != current_desc:
                    if 'params' not in chart:
                        chart['params'] = {}
                    chart['params']['custom_title'] = new_title
                    chart['params']['custom_description'] = new_desc
                    
                    # Add timestamp
                    from datetime import datetime
                    chart['params']['last_modified'] = datetime.now().isoformat()
                
                st.markdown("---")
            
            # Bulk actions
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Save All Metadata", key="save_metadata"):
                    st.success("✅ Metadata saved!")
                    st.rerun()
            
            with col2:
                if st.button("🗑️ Clear All Metadata", key="clear_metadata"):
                    for chart in st.session_state.configured_charts:
                        if 'params' in chart:
                            chart['params']['custom_title'] = ''
                            chart['params']['custom_description'] = ''
                    st.success("✅ Metadata cleared!")
                    st.rerun()
        
                st.markdown("---")

def create_dashboard_report_download(report_content, dashboard, df, key_suffix=""):
    """Create Word document download for dashboard report"""
    """Create Word document download for dashboard report - ENHANCED"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from datetime import datetime
        # Ensure report_content is a string
        if not isinstance(report_content, str):
            st.error("Report content is not in a valid format for download.")
            return
        
        # Create Word document
        doc = Document()
        
        # Add title and metadata
        doc.add_heading('📊 InsightNav AI - Comprehensive Dashboard Analysis Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        
        # Add dashboard overview
        doc.add_heading('📈 Dashboard Overview', level=1)
        dashboard_title = dashboard.get('title', 'Untitled Dashboard')
        dashboard_description = dashboard.get('description', 'No description provided')
        charts = dashboard.get('charts', [])
        
        doc.add_paragraph(f"Dashboard Title: {dashboard_title}")
        doc.add_paragraph(f"Description: {dashboard_description}")
        doc.add_paragraph(f"Number of Charts: {len(charts)}")
        
        if df is not None:
            doc.add_paragraph(f"Dataset: {df.shape[0]} rows × {df.shape[1]} columns")
            doc.add_paragraph(f"Columns: {', '.join(df.columns.tolist())}")
        
        doc.add_paragraph("")
        
        # Add chart details
        doc.add_heading('📊 Chart Configuration', level=1)
        for i, chart in enumerate(charts, 1):
            doc.add_paragraph(f"Chart {i}: {chart.get('title', f'Chart {i}')}")
            doc.add_paragraph(f"  Type: {chart.get('type', 'unknown')}")
            doc.add_paragraph(f"  Parameters: {chart.get('params', {})}")
            doc.add_paragraph(f"  Description: {chart.get('description', 'No description')}")
            doc.add_paragraph("")
        
        # Add AI analysis
        doc.add_heading('🤖 AI Analysis Report', level=1)
        # Handle markdown conversion more gracefully
        for line in report_content.split('\n'):
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            else:
                doc.add_paragraph(line)
        
        # Convert markdown to plain text for Word document
        import re
        # Remove markdown formatting for Word document
        plain_text = re.sub(r'#{1,6}\s*', '', report_content)  # Remove headers
        plain_text = re.sub(r'\*\*(.*?)\*\*', r'\1', plain_text)  # Remove bold
        plain_text = re.sub(r'\*(.*?)\*', r'\1', plain_text)  # Remove italic
        plain_text = re.sub(r'`(.*?)`', r'\1', plain_text)  # Remove code formatting
        
        doc.add_paragraph(plain_text)
        doc.add_paragraph("")
        
        # Add footer
        doc.add_paragraph("---")
        doc.add_paragraph("Report generated by InsightNav AI Dashboard Analysis System")
        doc.add_paragraph(f"For more information, visit your dashboard at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Save to buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Create download button with unique key
        unique_key = f"dashboard_report_download_btn{key_suffix}"
        st.download_button(
            label="📄 Download Dashboard Report (Word)",
            data=doc_buffer.getvalue(),
            file_name=f"dashboard_analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="secondary",
            key=unique_key
        )
        
    except Exception as e:
        st.error(f"❌ Error creating Word document: {str(e)}")
        st.info("💡 You can still copy the report text above and paste it into a Word document manually.")

def create_optimized_ai_dashboard(prompt, df, chat_history=None):
    """Create an AI-generated dashboard configuration based on user prompt and data"""
    try:
        # Get AI settings from session state
        ai_provider = st.session_state.get('global_ai_provider', 'DeepSeek')
        model_name = st.session_state.get('global_model_name', 'deepseek-chat')
        api_key = st.session_state.get('global_api_key', '')
        
        if not api_key:
            return None
        
        # Import required modules
        from ai_engine import get_model_context_protocol
        from langchain_openai import ChatOpenAI
        from langchain_anthropic import ChatAnthropic
        
        # Get context protocol
        context_protocol = get_model_context_protocol(ai_provider, model_name)
        
        # Initialize LLM based on provider
        if ai_provider == "DeepSeek":
            llm = ChatOpenAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                base_url="https://api.deepseek.com/v1",
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "OpenAI":
            llm = ChatOpenAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "Anthropic":
            llm = ChatAnthropic(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "Google":
            from langchain_google_genai import ChatGoogleGenerativeAI
            llm = ChatGoogleGenerativeAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                google_api_key=api_key,
                max_output_tokens=context_protocol["max_tokens"]
            )
        else:
            return None
        
        # Analyze the data to provide context
        data_info = ""
        if df is not None:
            numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
            categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
            data_info = f"""
DATA CONTEXT:
- Dataset shape: {df.shape[0]} rows, {df.shape[1]} columns
- Numeric columns: {', '.join(numeric_cols[:5])}{'...' if len(numeric_cols) > 5 else ''}
- Categorical columns: {', '.join(categorical_cols[:5])}{'...' if len(categorical_cols) > 5 else ''}
- Sample data preview: {df.head(2).to_dict('records')}
"""
        
        # Parse user requirements for chart count and types
        chart_requirements = parse_chart_requirements(prompt)
        request_analysis = analyze_user_request(prompt, df)
        
        # Determine chart count and types based on user request
        if chart_requirements['chart_count_specified']:
            chart_count = chart_requirements['chart_count']
            chart_count_instruction = f"Create EXACTLY {chart_count} chart{'s' if chart_count != 1 else ''} as requested by the user"
        else:
            chart_count = 1  # Default to 1 chart if not specified
            chart_count_instruction = "Create 1 chart that best addresses the user's request (unless they specify otherwise)"
        
        # Build chart type instructions
        if chart_requirements['specific_types']:
            type_instruction = f"Use these specific chart types as requested: {', '.join(chart_requirements['specific_types'])}"
        else:
            type_instruction = "Choose the most appropriate chart type based on the data and user's analytical needs"
        
        # Create enhanced prompt for dashboard generation
        enhanced_prompt = f"""
AI DASHBOARD GENERATOR

You are an expert data visualization specialist. Create a dashboard that EXACTLY matches the user's specific request.

{data_info}

USER REQUEST: "{prompt}"

REQUEST ANALYSIS:
{request_analysis}

CHART REQUIREMENTS:
- Number of charts: {chart_count} {'(user specified)' if chart_requirements['chart_count_specified'] else '(default)'}
- Chart types: {', '.join(chart_requirements['specific_types']) if chart_requirements['specific_types'] else 'AI selected based on data'}

INSTRUCTIONS:
1. {chart_count_instruction}
2. {type_instruction}
3. Use actual column names from the dataset
4. Create meaningful titles that reflect the user's intent
5. Only add KPIs if the user specifically requests them

CHART TYPE GUIDELINES:
- bar: For comparing categories, showing distributions
- line: For trends over time, continuous data
- scatter: For relationships between two numeric variables
- pie: For showing proportions (max 6 categories)
- histogram: For data distributions
- box: For statistical summaries and outliers

RESPONSE FORMAT (JSON only):
{{
    "title": "Dashboard title reflecting user request",
    "description": "Brief description of what this dashboard shows",
    "charts": [
        {{
            "id": "chart1",
            "title": "Chart title addressing user request",
            "type": "{'|'.join(chart_requirements['specific_types']) if chart_requirements['specific_types'] else 'most_appropriate_chart_type'}",
            "params": {{
                "x_col": "actual_column_name_from_data",
                "y_col": "actual_column_name_from_data",
                "color_col": "optional_grouping_column"
            }},
            "description": "What insights this chart provides"
        }}{"," if chart_count > 1 else ""}
        {"// Add more charts only if user requested multiple charts" if chart_count > 1 else ""}
    ],
    "kpis": [
        {"// Only include KPIs if user specifically requested them" if not any(word in prompt.lower() for word in ['kpi', 'metric', 'summary', 'total', 'average']) else ""}
    ]
}}

CRITICAL: Create EXACTLY {chart_count} chart{'s' if chart_count != 1 else ''}. Do not create more or fewer charts than requested.
IMPORTANT: Return ONLY the JSON configuration. No additional text or explanations.
"""
        
        # Generate response
        response = llm.invoke(enhanced_prompt)
        
        # Try to parse the JSON response
        import json
        try:
            dashboard_config = json.loads(response.content)
            
            # Convert old format to new format if needed
            if 'charts' in dashboard_config:
                for chart in dashboard_config['charts']:
                    if 'params' not in chart and ('x_column' in chart or 'y_column' in chart):
                        # Convert old format to new format
                        chart['params'] = {}
                        if 'x_column' in chart:
                            chart['params']['x_col'] = chart.pop('x_column')
                        if 'y_column' in chart:
                            chart['params']['y_col'] = chart.pop('y_column')
                        if 'color_column' in chart:
                            chart['params']['color_col'] = chart.pop('color_column')
            
            # Validate the configuration has required fields
            if 'charts' not in dashboard_config:
                dashboard_config['charts'] = []
            if 'kpis' not in dashboard_config:
                dashboard_config['kpis'] = []
            if 'title' not in dashboard_config:
                dashboard_config['title'] = "AI Generated Dashboard"
            if 'description' not in dashboard_config:
                dashboard_config['description'] = "Dashboard created with AI assistance"
            
            return dashboard_config
            
        except json.JSONDecodeError:
            # If JSON parsing fails, create a fallback dashboard respecting user requirements
            chart_requirements = parse_chart_requirements(prompt)
            fallback_charts = []
            
            # Determine how many charts to create
            chart_count = chart_requirements['chart_count'] if chart_requirements['chart_count_specified'] else 1
            
            if df is not None:
                numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
                categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
                
                # Create charts based on user requirements
                for i in range(min(chart_count, 3)):  # Max 3 fallback charts
                    chart_id = f"chart{i+1}"
                    
                    if i == 0:
                        # First chart: Use user-specified type or best available
                        if chart_requirements['specific_types']:
                            chart_type = chart_requirements['specific_types'][0]
                        elif len(categorical_cols) > 0 and len(numeric_cols) > 0:
                            chart_type = "bar"
                        else:
                            chart_type = "histogram" if len(numeric_cols) > 0 else "bar"
                        
                        if chart_type == "bar" and len(categorical_cols) > 0 and len(numeric_cols) > 0:
                            fallback_charts.append({
                                "id": chart_id,
                                "title": f"{categorical_cols[0]} vs {numeric_cols[0]}",
                                "type": "bar",
                                "params": {
                                    "x_col": categorical_cols[0],
                                    "y_col": numeric_cols[0]
                                },
                                "description": f"Comparison of {numeric_cols[0]} across {categorical_cols[0]}"
                            })
                        elif chart_type == "histogram" and len(numeric_cols) > 0:
                            fallback_charts.append({
                                "id": chart_id,
                                "title": f"Distribution of {numeric_cols[0]}",
                                "type": "histogram",
                                "params": {
                                    "col": numeric_cols[0],
                                    "bins": 20
                                },
                                "description": f"Distribution pattern of {numeric_cols[0]}"
                            })
                    
                    elif i == 1 and chart_count > 1:
                        # Second chart: Different type or next specified type
                        if len(chart_requirements['specific_types']) > 1:
                            chart_type = chart_requirements['specific_types'][1]
                        elif len(numeric_cols) > 0:
                            chart_type = "histogram"
                        else:
                            continue
                        
                        if chart_type == "histogram" and len(numeric_cols) > 0:
                            fallback_charts.append({
                                "id": chart_id,
                                "title": f"Distribution of {numeric_cols[0]}",
                                "type": "histogram",
                                "params": {
                                    "col": numeric_cols[0],
                                    "bins": 20
                                },
                                "description": f"Distribution pattern of {numeric_cols[0]}"
                            })
                    
                    elif i == 2 and chart_count > 2:
                        # Third chart: Scatter plot if possible
                        if len(chart_requirements['specific_types']) > 2:
                            chart_type = chart_requirements['specific_types'][2]
                        elif len(numeric_cols) >= 2:
                            chart_type = "scatter"
                        else:
                            continue
                        
                        if chart_type == "scatter" and len(numeric_cols) >= 2:
                            fallback_charts.append({
                                "id": chart_id,
                                "title": f"{numeric_cols[0]} vs {numeric_cols[1]}",
                                "type": "scatter",
                                "params": {
                                    "x_col": numeric_cols[0],
                                    "y_col": numeric_cols[1],
                                    "color_col": categorical_cols[0] if len(categorical_cols) > 0 else None
                                },
                                "description": f"Relationship between {numeric_cols[0]} and {numeric_cols[1]}"
                            })
            
            # Default single chart if no data analysis possible
            if not fallback_charts:
                fallback_charts = [{
                    "id": "chart1",
                    "title": "Data Overview",
                    "type": chart_requirements['specific_types'][0] if chart_requirements['specific_types'] else "bar",
                    "params": {
                        "x_col": df.columns[0] if df is not None and len(df.columns) > 0 else "x",
                        "y_col": df.select_dtypes(include=['number']).columns[0] if df is not None and len(df.select_dtypes(include=['number']).columns) > 0 else "y"
                    },
                    "description": "Overview of the data"
                }]
            
            return {
                "title": f"Dashboard: {prompt[:50]}..." if len(prompt) > 50 else f"Dashboard: {prompt}",
                "description": "Dashboard created with AI assistance",
                "charts": fallback_charts[:chart_count],  # Limit to requested count
                "kpis": []
            }
        
    except Exception as e:
        print(f"Error creating AI dashboard: {e}")
        return None


# Visualization Settings
st.sidebar.divider()
if False:  # Hidden visualization settings
    st.sidebar.subheader("🎨 Visualization Settings")
try:
    color_sequences = [seq for seq in dir(px.colors.sequential) if not seq.startswith('_') and isinstance(getattr(px.colors.sequential, seq), list)]
    # Add "--None--" option at the beginning
    color_options = ["--None--"] + color_sequences
    default_index = 0  # Default to "--None--"
    color_palette = st.sidebar.selectbox(
        "**Color Palette:**",
        options=color_options,
        index=default_index,
        key="dashboard_color_palette",
        help="Select '--None--' to use default chart colors, or choose a palette to apply to all charts"
    )
except Exception as e:
    color_palette = "--None--"
    st.sidebar.warning(f"Could not load color palettes: {e}")

# ===== AI CHAT INTERFACE IN SIDEBAR =====
st.sidebar.divider()

# create_optimized_ai_dashboard function is defined above before the Generate with AI button

def generate_optimized_dashboard_response(prompt, context=""):
    """Generate AI response for dashboard questions using configured provider"""
    try:
        # Get AI settings from session state (same as analysis_agent.py)
        ai_provider = st.session_state.get('global_ai_provider', 'DeepSeek')
        model_name = st.session_state.get('global_model_name', 'deepseek-chat')
        api_key = st.session_state.get('global_api_key', '')
        
        if not api_key:
            return "⚠️ Please configure your AI API key in the Analysis Agent page first."
        
        # Import required modules
        from ai_engine import get_model_context_protocol
        from langchain_openai import ChatOpenAI
        from langchain_anthropic import ChatAnthropic
        
        # Get context protocol
        context_protocol = get_model_context_protocol(ai_provider, model_name)
        
        # Initialize LLM based on provider
        if ai_provider == "DeepSeek":
            llm = ChatOpenAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                base_url="https://api.deepseek.com/v1",
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "OpenAI":
            llm = ChatOpenAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "Anthropic":
            llm = ChatAnthropic(
                temperature=context_protocol["temperature"],
                model=model_name,
                api_key=api_key,
                max_tokens=context_protocol["max_tokens"]
            )
        elif ai_provider == "Google":
            from langchain_google_genai import ChatGoogleGenerativeAI
            llm = ChatGoogleGenerativeAI(
                temperature=context_protocol["temperature"],
                model=model_name,
                google_api_key=api_key,
                max_output_tokens=context_protocol["max_tokens"]
            )
        else:
            return f"❌ Unsupported AI provider: {ai_provider}"
        
        # Get comprehensive chart knowledge for context
        chart_knowledge = get_comprehensive_chart_knowledge()
        
        # Build chart expertise context
        chart_expertise = []
        for chart_type, info in chart_knowledge.items():
            chart_expertise.append(f"- **{info['name']} ({chart_type})** {info['icon']}: {info['description']}")
            chart_expertise.append(f"  Best for: {', '.join(info['use_cases'][:3])}")
            chart_expertise.append(f"  Business Value: {info['business_value']}")
        
        chart_expertise_text = '\n'.join(chart_expertise[:50])  # Limit for token efficiency
        
        # Create enhanced prompt with comprehensive chart knowledge
        enhanced_prompt = f"""
ADVANCED DASHBOARD AI ASSISTANT

You are a senior business intelligence consultant and dashboard expert with comprehensive knowledge of all visualization types and their strategic applications.

🎯 **YOUR EXPERTISE:**
- 16+ chart types including industry-standard visualizations (waterfall, funnel, gauge, treemap, violin, bullet charts)
- Deep understanding of when and how to use each chart type for maximum business impact
- Statistical analysis and data interpretation across all visualization methods
- Business intelligence strategy and dashboard optimization

📊 **COMPREHENSIVE CHART LIBRARY YOU KNOW:**
{chart_expertise_text}

🔧 **CURRENT CONTEXT:**
- AI Provider: {ai_provider}
- User has access to full dashboard creation suite with all chart types
- Focus on strategic insights and practical business recommendations
- Consider both basic and advanced visualization techniques

{context}

👤 **USER QUESTION:** {prompt}

🎯 **RESPONSE GUIDELINES:**
- Leverage your comprehensive knowledge of all 16+ chart types
- Recommend specific chart types based on data characteristics and business goals
- Provide strategic insights about visualization choices and their business impact
- Consider advanced charts like waterfall (financial analysis), funnel (process optimization), gauge (KPI monitoring)
- Give actionable recommendations with specific implementation guidance
- Use professional business intelligence language
- Reference specific chart capabilities and use cases when relevant

Please provide a detailed, expert-level response with specific recommendations and strategic insights.
"""
        
        # Generate response
        response = llm.invoke(enhanced_prompt)
        return response.content
        
    except Exception as e:
        return f"❌ Error generating AI response: {str(e)}"
    """Enhanced AI response generation for dashboard questions with comprehensive report capability"""
    
    # Check if this is a comprehensive report request
    is_comprehensive_request = any(keyword in prompt.lower() for keyword in [
        "comprehensive", "detailed analysis", "full report", "complete analysis",
        "generate comprehensive dashboard analysis report", "analyze each chart individually"
    ])
    
    # Manage conversation history
    manage_dashboard_conversation_history(max_turns=6)
    
    try:
        # Use unified LLM initialization
        llm = get_validated_dashboard_llm()
        
        # Check if comprehensive report is requested and dashboard exists
        if is_comprehensive_request and st.session_state.get('dashboard') and current_df is not None:
            # Generate comprehensive dashboard report
            try:
                dashboard_report = generate_comprehensive_dashboard_report(
                    current_df, 
                    st.session_state.dashboard,
                    st.session_state.get('dashboard_chat_history', [])
                )
                
                # Store report for download
                st.session_state.comprehensive_dashboard_report = dashboard_report
                
                return f"""# 🔬 Comprehensive Dashboard Analysis Report Generated

{dashboard_report}

---

💾 **Download Available**: A Word document version of this comprehensive report is now available for download. Look for the download button below this chat interface.

🎯 **Key Features Analyzed**: 
- Individual chart insights and interpretations
- Cross-chart relationships and patterns  
- Business intelligence and strategic recommendations
- Data quality and optimization suggestions"""
                
            except Exception as e:
                return f"❌ Error generating comprehensive report: {str(e)}\n\nFalling back to regular dashboard analysis..."
        
        # Regular dashboard response generation
        # Build optimized context
        df_hash = hash(str(current_df.columns.tolist())) if current_df is not None else 0
        data_context = build_optimized_dashboard_context(
            df_hash, 
            current_df.columns.tolist() if current_df is not None else [],
            len(df_sample) if df_sample is not None else 0,
            len(st.session_state.get('dashboard', {}).get('charts', []))
        )
        
        # Get cached data summary
        summary_context = get_dashboard_data_summary(df_hash)
        
        # Get current dashboard context
        dashboard_summary = get_dashboard_summary(st.session_state.get('dashboard', {}))
        
        # Enhanced system message for comprehensive analysis capability
        system_message_content = f"""You are an expert dashboard and data visualization consultant for InsightNav AI with comprehensive analysis capabilities.

CURRENT CONTEXT:
- Dataset: {data_context['column_count']} columns, {summary_context.get('total_rows', 0):,} total rows
- Sample: {data_context['sample_size']:,} rows for visualization
- Dashboard: {dashboard_summary}
- Charts: {data_context['chart_count']} created

ENHANCED CAPABILITIES:
- Chart-by-chart analysis and interpretation
- Cross-chart relationship identification
- Business intelligence and strategic insights
- Data storytelling and narrative creation
- Performance optimization recommendations

{context}

GUIDELINES:
1. Provide actionable dashboard and visualization advice
2. For comprehensive requests, be thorough and detailed
3. For regular questions, be concise but insightful (1-2 paragraphs)
4. Suggest specific chart types and improvements
5. Focus on data storytelling and user experience
6. Relate advice to the user's current dashboard
7. If asked about comprehensive analysis, explain the detailed report capability"""

        # Prepare messages with limited history
        messages = [
            {"role": "system", "content": system_message_content},
            {"role": "user", "content": prompt}
        ]
        
        # Add recent conversation history (limited to prevent token overflow)
        if st.session_state.dashboard_chat_history:
            recent_history = st.session_state.dashboard_chat_history[-4:]  # Only last 2 exchanges
            messages = [messages[0]] + recent_history + [messages[1]]
        
        # Generate response using LangChain
        response = llm.invoke(messages).content.strip()
        
        return response
        
    except ValueError as e:
        return str(e)  # API key validation error
    except Exception as e:
        return f"❌ Error generating AI response: {str(e)}"

# AI Dashboard Assistant moved to main page - see below dashboard rendering
    
    # Performance metrics and controls
    col1, col2 = st.columns(2)
    with col1:
        st.metric("💬 Chats", f"{len(st.session_state.dashboard_chat_history)//2}")
    with col2:
        provider = st.session_state.get('global_ai_provider', 'OpenAI')
        st.metric("🤖 Provider", provider)
    
    # Clear chat button
    if st.button("🧹 Clear Chat", key="clear_dashboard_chat", use_container_width=True):
        st.session_state.dashboard_chat_history = []
        cleanup_dashboard_session_state()
        st.success("✅ Chat cleared!")
        st.rerun()

# ===== SMART CHART RECOMMENDATIONS =====
def show_smart_chart_recommendations():
    """Display intelligent chart recommendations"""
    
    if current_df is None:
        return
    
    st.markdown("---")
    st.subheader("📊 Smart Chart Recommendations")
    
    # Analyze data for recommendations
    analysis = analyze_data_for_chart_recommendations(current_df)
    
    if not analysis['recommendations']:
        st.info("💡 Upload data to get personalized chart recommendations!")
        return
    
    st.write("**🎯 Recommended charts based on your data:**")
    
    # Group recommendations by priority
    high_priority = [r for r in analysis['recommendations'] if r['priority'] == 'high']
    medium_priority = [r for r in analysis['recommendations'] if r['priority'] == 'medium']
    low_priority = [r for r in analysis['recommendations'] if r['priority'] == 'low']
    
    # Show high priority recommendations first
    if high_priority:
        st.write("**🔥 Highly Recommended:**")
        for rec in high_priority:
            with st.expander(f"📈 {rec['title']}"):
                st.write(f"**Why this chart:** {rec['reasoning']}")
                st.write(f"**Suggested columns:** {', '.join(rec['columns_suggested'])}")
                
                if rec.get('warning'):
                    st.warning(f"⚠️ {rec['warning']}")
                
                if st.button(f"Get Advice for {rec['type'].title()}", key=f"advice_chart_{rec['type']}"):
                    advice_prompt = f"How do I create an effective {rec['type']} chart with columns {', '.join(rec['columns_suggested'])}?"
                    with st.spinner("🤖 Getting chart advice..."):
                        advice = generate_optimized_dashboard_response(advice_prompt)
                        st.session_state.dashboard_chat_history.append({"role": "user", "content": advice_prompt})
                        st.session_state.dashboard_chat_history.append({"role": "assistant", "content": advice})
                        st.markdown(f"**Chart Advice:**\n\n{advice}")
    
    # Show medium priority in a collapsed section
    if medium_priority:
        with st.expander("📊 Additional Recommendations"):
            for rec in medium_priority:
                st.write(f"**{rec['title']}:** {rec['reasoning']}")
    
    # Show data insights
    st.write("**📋 Data Analysis Summary:**")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Numeric Columns", len(analysis['numeric_columns']))
    with col2:
        st.metric("Categorical Columns", len(analysis['categorical_columns']))
    with col3:
        st.metric("Data Points", analysis['data_size'])

# ===== DASHBOARD DESIGN ANALYSIS (SIMPLIFIED) =====
def show_dashboard_design_analysis():
    """Simplified dashboard design analysis - removed to reduce clutter"""
    # Function kept for compatibility but no longer displays cluttered feedback
    pass

# Show smart features when appropriate
if current_df is not None and not st.session_state.get('show_dashboard'):
    show_smart_chart_recommendations()

# Dashboard design analysis removed to reduce clutter

# ===== DASHBOARD RENDERING =====
if 'show_dashboard' in st.session_state and st.session_state.show_dashboard:
    if 'dashboard' in st.session_state:
        render_dashboard(st.session_state.dashboard, df_sample)
        
        # Dashboard management buttons
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("🔄 Regenerate", use_container_width=True):
                st.session_state.show_dashboard = False
                st.session_state.dashboard_edit_mode = False # Exit edit mode
                cleanup_dashboard_session_state()
                st.rerun()
        with col2:
            # Export as HTML - Pass the selected palette name
            try:
                palette_to_export = st.session_state.get("dashboard_color_palette", "Viridis")
                dashboard_html = export_dashboard_to_html(st.session_state.dashboard, df_sample) # Palette handled inside
                st.download_button(
                    label="💾 Export HTML",
                    data=dashboard_html,
                    file_name="dashboard.html",
                    mime="text/html",
                    use_container_width=True,
                    key="dashboard_html_export_btn"
                )
            except Exception as e:
                st.error(f"❌ Error exporting dashboard: {str(e)}")
        with col3:
            if st.button("🗑️ Reset All", use_container_width=True):
                cleanup_dashboard_session_state()
                st.session_state.show_dashboard = False
                st.session_state.dashboard_edit_mode = False
                st.rerun()
        with col4:
            if st.button("❌ Close", use_container_width=True):
                st.session_state.show_dashboard = False
                st.session_state.dashboard_edit_mode = False
                st.rerun()
        
        # Add comprehensive dashboard report section
        st.markdown("---")
        st.markdown("## 🤖 AI Dashboard Analysis")
        
        # Generate comprehensive dashboard report button
        if st.button("📊 Generate Comprehensive Dashboard Report", use_container_width=True, type="primary"):
            if not st.session_state.get('global_api_key'):
                st.error("❌ Please configure your API key in the Analysis Agent page first.")
            else:
                with st.spinner("🔍 Analyzing dashboard and generating comprehensive report..."):
                    try:
                        # Generate comprehensive dashboard analysis
                        dashboard_report = generate_comprehensive_dashboard_analysis(
                            st.session_state.dashboard, 
                            df_sample,
                            st.session_state.get('dashboard_chat_history', [])
                        )
                        
                        if dashboard_report:
                            # Display the report in the main area
                            st.markdown("### 📋 Dashboard Analysis Report")
                            st.markdown(dashboard_report)
                            
                            # Store report for download
                            st.session_state.dashboard_report = dashboard_report
                            
                            # Create Word document download
                            create_dashboard_report_download(dashboard_report, st.session_state.dashboard, df_sample, "_new")
                        else:
                            st.error("❌ Failed to generate dashboard report. Please try again.")
                            
                    except Exception as e:
                        st.error(f"❌ Error generating dashboard report: {str(e)}")
        
        # Display existing report if available
        if 'dashboard_report' in st.session_state and st.session_state.dashboard_report:
            st.markdown("### 📋 Latest Dashboard Analysis Report")
            with st.expander("View Report", expanded=False):
                st.markdown(st.session_state.dashboard_report)
            
            # Always show download button if report exists
            create_dashboard_report_download(st.session_state.dashboard_report, st.session_state.dashboard, df_sample, "_existing")
        
        # Enhanced AI Dashboard Assistant - Main Page
        st.markdown("---")
        st.markdown("## 💬 AI Dashboard Assistant")
        st.markdown("Ask questions about your dashboard, get insights, and generate comprehensive analysis reports.")
        
        # Create enhanced dashboard-aware AI assistant
        create_enhanced_dashboard_assistant(st.session_state.dashboard, df_sample)
    else:
        st.warning("⚠️ Dashboard not generated. Please click 'Generate Dashboard' first.")
else:
    st.info("⚙️ Configure your dashboard using the sidebar and click 'Generate Dashboard'.")

    # ===== COPYRIGHT FOOTER =====
st.markdown("<div class='modern-divider'></div>", unsafe_allow_html=True)
st.markdown("""
    <div class="modern-card" style="text-align: center; padding: 20px;">
        <p style="margin: 0; color: #667eea; font-weight: 600;">© 2025 InsightNav AI. All rights reserved.</p>
        <p style="margin: 5px 0 0 0; color: #764ba2; font-style: italic;">Your intelligent data exploration companion</p>
    </div>
""", unsafe_allow_html=True)
